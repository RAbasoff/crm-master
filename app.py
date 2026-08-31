"""
CRM-система для мастерской с производственным модулем
Авторизация · Роли · Карта цеха · Заявки · Уведомления · Отчёты
"""
from flask import Flask, render_template, request, redirect, url_for, flash, jsonify, send_file, session, g
from flask_babel import Babel, gettext as _, get_locale
from flask_login import login_user, logout_user, login_required, current_user
from flask_wtf.csrf import CSRFProtect
from werkzeug.utils import secure_filename
from urllib.parse import urlparse
from datetime import datetime, timedelta
import os, io, json
import qrcode

from config import Config, LANGUAGES, SECTION_KEYS
from models import (db, User, UserSectionAccess, FactorySection, Machine, MachinePart,
                    PartMaintenanceLog, MachineDocument, MaintenanceRecord, MaintenancePhoto,
                    MaintenancePlan, MachineSparePart, ResponsibleGroup, Verantwoordelijke,
                    Monteur, Contractor, ContractorEmployee, WarehouseGroup, VoorraadItem,
                    VoorraadMutatie, Invoice, InvoiceItem, GasCylinder, CylinderLog,
                    CylinderOrder, FaultReport, FaultPhoto, FaultVideo, WorkReport, WorkReportPhoto,
                    PurchaseRequest, WorkSchedule, TimeEntry, Vacation, Message, Notification,
                    Opdracht, AuditLog, TechnicalWorkOrder, TWOPhoto, two_workers,
                    GasSystemComponent, EquipmentRepair, fault_technicians, user_machine, section_responsible,
                    GroupPermission, ResponsibleAuth, ElectricalCabinet, CircuitBreaker, WeekendShift,
                    FaultStatusHistory, WorkReportEntry, ToolWear, MonthlyArchive,
                    TWOChecklistItem, TWOSignature, TWOAssignment,
                    UserActivityLog, SystemLog,
                    MuleMaintenance, MuleMaintenancePart, MulePartOrder, MuleComponent)
from utils import (role_required, user_has_section_access, section_access_required,
                   create_notification, log_audit, genereer_nummer, date_plus_days,
                   save_uploaded_file, translate_text, run_migrations,
                   log_user_activity, log_system)

# ============================================================
# APP CONFIG
# ============================================================

app = Flask(__name__)
app.config.from_object(Config)
app.config['WTF_CSRF_TIME_LIMIT'] = None  # no timeout for long sessions
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(os.path.join(os.path.dirname(os.path.abspath(__file__)), 'instance'), exist_ok=True)

csrf = CSRFProtect(app)
db.init_app(app)
from flask_login import LoginManager

login_manager = LoginManager(app)
login_manager.login_view = 'login'

# Auto-create database tables and seed data on import (for WSGI deployment)
with app.app_context():
    db.create_all()
    run_migrations()
    if User.query.count() == 0:
        import secrets as _secrets
        admin = User(username='admin', display_name='Administrator', role='admin')
        admin.set_password('admin123')
        tech = User(username='tech', display_name='Sergei Petrov', role='technician')
        tech.set_password('tech123')
        user = User(username='user', display_name='Jan de Vries', role='user')
        user.set_password('user123')
        director = User(username='director', display_name='Director', role='director')
        director.set_password('director123')
        db.session.add_all([admin, tech, user, director])
        db.session.commit()

def get_current_locale():
    return session.get('lang', 'ru')

babel = Babel(app, locale_selector=get_current_locale)

@app.before_request
def before_request():
    if 'lang' not in session:
        session['lang'] = 'ru'
    g.lang = session.get('lang', 'ru')
    g.LANGUAGES = LANGUAGES

@app.route('/set_language/<lang>')
def set_language(lang):
    if lang in LANGUAGES:
        session['lang'] = lang
    return redirect(request.referrer or url_for('index'))

# ============================================================
# USER LOADER
# ============================================================

@login_manager.user_loader
def load_user(user_id):
    # Responsible person IDs are prefixed with "r_"
    if str(user_id).startswith('r_'):
        try:
            person_id = int(str(user_id)[2:])
            person = Verantwoordelijke.query.get(person_id)
            if person and person.password_hash and person.is_active:
                return ResponsibleAuth(person)
        except (ValueError, TypeError):
            pass
        return None
    return User.query.get(int(user_id))

# ============================================================
# ERROR HANDLERS
# ============================================================

@app.errorhandler(400)
def bad_request(e):
    if request.accept_mimetypes.accept_json and not request.accept_mimetypes.accept_html:
        return jsonify(error=str(e)), 400
    flash(_('Bad request'), 'error')
    return redirect(request.referrer or url_for('index'))

@app.errorhandler(403)
def forbidden(e):
    if request.accept_mimetypes.accept_json and not request.accept_mimetypes.accept_html:
        return jsonify(error='Forbidden'), 403
    flash(_('Access denied'), 'error')
    return redirect(url_for('index'))

@app.errorhandler(404)
def not_found(e):
    if request.accept_mimetypes.accept_json and not request.accept_mimetypes.accept_html:
        return jsonify(error='Not found'), 404
    flash(_('Page not found'), 'error')
    return redirect(url_for('index'))

@app.errorhandler(500)
def internal_error(e):
    db.session.rollback()
    if request.accept_mimetypes.accept_json and not request.accept_mimetypes.accept_html:
        return jsonify(error='Internal server error'), 500
    flash(_('An error occurred. Please try again.'), 'error')
    return redirect(url_for('index'))

# ============================================================
# ROUTES — AUTH
# ============================================================

@app.route('/login', methods=['GET', 'POST'])
def login():
    if current_user.is_authenticated:
        return redirect(url_for('index'))
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        # Try User first
        user = User.query.filter_by(username=username).first()
        if user and user.check_password(password) and user.is_active_user:
            login_user(user, remember=True)
            log_user_activity('login', page='/login', details=f'User {username} logged in')
            log_system('INFO', 'auth', f'User {username} logged in', source='login')
            next_url = request.args.get('next')
            if next_url:
                parsed = urlparse(next_url)
                if parsed.netloc and parsed.netloc != request.host:
                    next_url = None
            return redirect(next_url or url_for('index'))
        # Try Verantwoordelijke by email or naam
        person = Verantwoordelijke.query.filter(
            (Verantwoordelijke.email == username) | (Verantwoordelijke.naam == username)
        ).first()
        if person and person.check_password(password) and person.is_active:
            auth = ResponsibleAuth(person)
            person.last_login = datetime.utcnow()
            db.session.commit()
            login_user(auth, remember=True)
            log_user_activity('login', page='/login', details=f'Responsible {username} logged in')
            log_system('INFO', 'auth', f'Responsible {username} logged in', source='login')
            next_url = request.args.get('next')
            if next_url:
                parsed = urlparse(next_url)
                if parsed.netloc and parsed.netloc != request.host:
                    next_url = None
            return redirect(next_url or url_for('floor_plan'))
        log_system('WARNING', 'auth', f'Failed login attempt for {username}', source='login')
        flash(_('Invalid credentials'), 'error')
    return render_template('login.html')

@app.route('/logout')
@login_required
def logout():
    username = current_user.username if current_user.is_authenticated else 'unknown'
    log_user_activity('logout', page='/logout', details=f'User {username} logged out')
    log_system('INFO', 'auth', f'User {username} logged out', source='logout')
    logout_user()
    return redirect(url_for('login'))

@app.route('/profile', methods=['GET', 'POST'])
@login_required
def profile():
    if request.method == 'POST':
        new_pass = request.form.get('new_password')
        if new_pass:
            if hasattr(current_user, '_person'):
                current_user._person.set_password(new_pass)
            else:
                current_user.set_password(new_pass)
                current_user.display_name = request.form.get('display_name', current_user.display_name)
            db.session.commit()
            flash(_('Profile updated'), 'success')
    return render_template('profile.html')

# ============================================================
# ROUTES — USER MANAGEMENT (Admin only)
# ============================================================

@app.route('/users')
@login_required
@role_required('admin')
def users_list():
    users = User.query.all()
    return render_template('users.html', users=users)

@app.route('/users/<int:user_id>')
@login_required
def user_cabinet(user_id):
    u = User.query.get_or_404(user_id)
    return render_template('user_cabinet.html', user=u)

@app.route('/users/<int:user_id>/change-password', methods=['POST'])
@login_required
@role_required('admin')
def user_change_password(user_id):
    u = User.query.get_or_404(user_id)
    new_pass = request.form.get('new_password')
    confirm_pass = request.form.get('confirm_password')
    if not new_pass:
        flash(_('Password cannot be empty'), 'error')
    elif new_pass != confirm_pass:
        flash(_('Passwords do not match'), 'error')
    else:
        u.set_password(new_pass)
        db.session.commit()
        flash(_('Password changed for %(username)s', username=u.username), 'success')
    return redirect(url_for('user_cabinet', user_id=u.id))

@app.route('/users/<int:user_id>/cabinet-update', methods=['POST'])
@login_required
@role_required('admin')
def user_cabinet_update(user_id):
    u = User.query.get_or_404(user_id)
    u.first_name = request.form.get('first_name', u.first_name)
    u.last_name = request.form.get('last_name', u.last_name)
    u.display_name = request.form.get('display_name', u.display_name)
    u.role = request.form.get('role', u.role)
    u.access_level = request.form.get('access_level', u.access_level)
    u.is_active_user = 'is_active' in request.form
    u.hire_date = datetime.strptime(request.form['hire_date'], '%Y-%m-%d').date() if request.form.get('hire_date') else u.hire_date
    u.fire_date = datetime.strptime(request.form['fire_date'], '%Y-%m-%d').date() if request.form.get('fire_date') else None
    # Password change (optional)
    new_pass = request.form.get('new_password')
    if new_pass:
        confirm_pass = request.form.get('confirm_password')
        if new_pass != confirm_pass:
            flash(_('Passwords do not match'), 'error')
            return redirect(url_for('user_cabinet', user_id=u.id))
        u.set_password(new_pass)
    db.session.commit()
    flash(_('User updated'), 'success')
    return redirect(url_for('user_cabinet', user_id=u.id))

@app.route('/users/<int:user_id>/delete', methods=['POST'])
@login_required
@role_required('admin')
def user_delete(user_id):
    if user_id == current_user.id:
        flash(_('Cannot delete yourself'), 'error')
        return redirect(url_for('user_cabinet', user_id=user_id))
    u = User.query.get_or_404(user_id)
    username = u.username
    uid = u.id
    # Delete NOT NULL FK rows
    Message.query.filter((Message.sender_id == uid) | (Message.receiver_id == uid)).delete(synchronize_session=False)
    Notification.query.filter_by(user_id=uid).delete()
    TimeEntry.query.filter_by(user_id=uid).delete()
    Vacation.query.filter_by(user_id=uid).delete()
    WorkSchedule.query.filter_by(user_id=uid).delete()
    WorkReport.query.filter_by(technician_id=uid).delete()
    PurchaseRequest.query.filter_by(requester_id=uid).delete()
    # Null out nullable FKs
    FactorySection.query.filter_by(responsible_user_id=uid).update({'responsible_user_id': None})
    Machine.query.filter_by(responsible_user_id=uid).update({'responsible_user_id': None})
    MachinePart.query.filter_by(responsible_user_id=uid).update({'responsible_user_id': None})
    PartMaintenanceLog.query.filter_by(performed_by=uid).update({'performed_by': None})
    MachineDocument.query.filter_by(uploaded_by=uid).update({'uploaded_by': None})
    MaintenanceRecord.query.filter_by(performed_by=uid).update({'performed_by': None})
    MaintenancePlan.query.filter_by(created_by=uid).update({'created_by': None})
    Monteur.query.filter_by(user_id=uid).update({'user_id': None})
    Invoice.query.filter_by(signed_by=uid).update({'signed_by': None})
    Invoice.query.filter_by(created_by=uid).update({'created_by': None})
    CylinderLog.query.filter_by(performed_by=uid).update({'performed_by': None})
    CylinderOrder.query.filter_by(ordered_by=uid).update({'ordered_by': None})
    FaultReport.query.filter_by(reporter_id=uid).delete()
    FaultReport.query.filter_by(technician_id=uid).update({'technician_id': None})
    PurchaseRequest.query.filter_by(reviewer_id=uid).update({'reviewer_id': None})
    TechnicalWorkOrder.query.filter_by(created_by=uid).update({'created_by': None})
    AuditLog.query.filter_by(user_id=uid).update({'user_id': None})
    TimeEntry.query.filter_by(approved_by=uid).update({'approved_by': None})
    Vacation.query.filter_by(approved_by=uid).update({'approved_by': None})
    # Association tables
    db.session.execute(user_machine.delete().where(user_machine.c.user_id == uid))
    db.session.execute(fault_technicians.delete().where(fault_technicians.c.technician_id == uid))
    db.session.execute(two_workers.delete().where(two_workers.c.worker_id == uid))
    # UserSectionAccess with cascade
    UserSectionAccess.query.filter_by(user_id=uid).delete()
    # Delete user
    db.session.delete(u)
    db.session.commit()
    flash(_('User %(username)s deleted', username=username), 'success')
    return redirect(url_for('index'))

@app.route('/users/new', methods=['GET', 'POST'])
@login_required
@role_required('admin')
def user_new():
    if request.method == 'POST':
        u = User(
            username=request.form['username'],
            first_name=request.form.get('first_name', ''),
            last_name=request.form.get('last_name', ''),
            display_name=request.form.get('display_name', ''),
            role=request.form.get('role', 'user'),
            access_level=request.form.get('access_level', 'full'),
            person_id=int(request.form['person_id']) if request.form.get('person_id') else None,
            hire_date=datetime.strptime(request.form['hire_date'], '%Y-%m-%d').date() if request.form.get('hire_date') else None
        )
        u.set_password(request.form['password'])
        db.session.add(u)
        db.session.flush()
        # Save allowed sections
        for key in request.form.getlist('allowed_sections'):
            db.session.add(UserSectionAccess(user_id=u.id, section_key=key))
        # Save assigned machines
        for mid in request.form.getlist('machines'):
            m = Machine.query.get(int(mid))
            if m:
                u.assigned_machines.append(m)
        db.session.commit()
        flash(_('User created'), 'success')
        return redirect(url_for('users_list'))
    verantwoordelijken = Verantwoordelijke.query.order_by(Verantwoordelijke.naam).all()
    return render_template('user_form.html', user=None, machines=Machine.query.all(), section_keys=SECTION_KEYS, verantwoordelijken=verantwoordelijken)

@app.route('/users/<int:user_id>/edit', methods=['GET', 'POST'])
@login_required
@role_required('admin')
def user_edit(user_id):
    u = User.query.get_or_404(user_id)
    if request.method == 'POST':
        u.first_name = request.form.get('first_name', u.first_name)
        u.last_name = request.form.get('last_name', u.last_name)
        u.display_name = request.form.get('display_name', u.display_name)
        u.role = request.form.get('role', u.role)
        u.access_level = request.form.get('access_level', u.access_level)
        u.person_id = int(request.form['person_id']) if request.form.get('person_id') else None
        u.is_active_user = 'is_active' in request.form
        u.hire_date = datetime.strptime(request.form['hire_date'], '%Y-%m-%d').date() if request.form.get('hire_date') else u.hire_date
        u.fire_date = datetime.strptime(request.form['fire_date'], '%Y-%m-%d').date() if request.form.get('fire_date') else None
        new_pass = request.form.get('password')
        if new_pass:
            u.set_password(new_pass)
        # Update allowed sections
        UserSectionAccess.query.filter_by(user_id=u.id).delete()
        for key in request.form.getlist('allowed_sections'):
            db.session.add(UserSectionAccess(user_id=u.id, section_key=key))
        # Update assigned machines
        u.assigned_machines = []
        for mid in request.form.getlist('machines'):
            m = Machine.query.get(int(mid))
            if m:
                u.assigned_machines.append(m)
        db.session.commit()
        flash(_('User updated'), 'success')
        return redirect(url_for('users_list'))
    verantwoordelijken = Verantwoordelijke.query.order_by(Verantwoordelijke.naam).all()
    return render_template('user_form.html', user=u, machines=Machine.query.all(), section_keys=SECTION_KEYS, verantwoordelijken=verantwoordelijken)

@app.route('/monteurs')
@login_required
@role_required('admin')
def monteurs_list():
    monteurs = User.query.filter_by(role='technician').order_by(User.display_name).all()
    machines = Machine.query.order_by(Machine.name).all()
    return render_template('monteurs.html', monteurs=monteurs, machines=machines, section_keys=SECTION_KEYS)

@app.route('/monteurs/<int:user_id>/permissions', methods=['POST'])
@login_required
@role_required('admin')
def monteur_permissions(user_id):
    u = User.query.get_or_404(user_id)
    if u.role != 'technician':
        flash(_('Only technicians can be edited here'), 'error')
        return redirect(url_for('monteurs_list'))
    u.access_level = request.form.get('access_level', 'full')
    UserSectionAccess.query.filter_by(user_id=u.id).delete()
    for key in request.form.getlist('allowed_sections'):
        db.session.add(UserSectionAccess(user_id=u.id, section_key=key))
    u.assigned_machines = []
    for mid in request.form.getlist('machines'):
        m = Machine.query.get(int(mid))
        if m:
            u.assigned_machines.append(m)
    u.is_active_user = 'is_active' in request.form
    new_pass = request.form.get('password')
    if new_pass:
        u.set_password(new_pass)
    db.session.commit()
    flash(_('Permissions updated for') + ' ' + (u.display_name or u.username), 'success')
    return redirect(url_for('monteurs_list'))

# ============================================================
# ROUTES — MACHINES & FACTORY FLOOR
# ============================================================

@app.route('/machines')
@login_required
def machines_list():
    if current_user.has_role('admin', 'director'):
        machines = Machine.query.all()
    elif current_user.has_role('technician'):
        machines = Machine.query.all()  # Technicians see all
    else:
        machines = current_user.assigned_machines
    return render_template('machines.html', machines=machines)

@app.route('/machines/new', methods=['GET', 'POST'])
@login_required
@role_required('admin')
def machine_new():
    if request.method == 'POST':
        m = Machine(
            name=request.form['name'],
            description=request.form.get('description', ''),
            serial_number=request.form.get('serial_number', ''),
            machine_type=request.form.get('machine_type', ''),
            manufacturer=request.form.get('manufacturer', ''),
            year_of_manufacture=int(request.form['year_of_manufacture']) if request.form.get('year_of_manufacture') else None,
            installation_location=request.form.get('installation_location', ''),
            contractor_id=int(request.form['contractor_id']) if request.form.get('contractor_id') else None,
            responsible_user_id=int(request.form['responsible_user_id']) if request.form.get('responsible_user_id') else None,
            responsible_person_id=int(request.form['responsible_person_id']) if request.form.get('responsible_person_id') else None,
            section_id=int(request.form['section_id']) if request.form.get('section_id') else None,
            marker_size=int(request.form.get('marker_size', 45)),
            marker_shape=request.form.get('marker_shape', 'circle'),
            floor_x=float(request.form.get('floor_x', 50)),
            floor_y=float(request.form.get('floor_y', 50))
        )
        if 'photo' in request.files and request.files['photo'].filename:
            filename = secure_filename(f"machine_{request.files['photo'].filename}")
            request.files['photo'].save(os.path.join(app.config['UPLOAD_FOLDER'], filename))
            m.photo = filename
        db.session.add(m)
        db.session.flush()
        # Assign users
        for uid in request.form.getlist('assigned_users'):
            u = User.query.get(int(uid))
            if u:
                m.assigned_users.append(u)
        db.session.commit()
        flash(_('Machine created'), 'success')
        return redirect(url_for('machine_detail', machine_id=m.id))
    users = User.query.filter(User.is_active_user == True).all()
    sections = FactorySection.query.all()
    contractors = Contractor.query.filter_by(is_active=True).order_by(Contractor.company_name).all()
    verantwoordelijken = Verantwoordelijke.query.order_by(Verantwoordelijke.naam).all()
    return render_template('machine_form.html', machine=None, users=users, sections=sections, contractors=contractors, verantwoordelijken=verantwoordelijken)

@app.route('/machines/<int:machine_id>')
@login_required
def machine_detail(machine_id):
    m = Machine.query.get_or_404(machine_id)
    faults = FaultReport.query.filter_by(machine_id=m.id).order_by(FaultReport.created_at.desc()).all()
    maintenance = MaintenanceRecord.query.filter_by(machine_id=m.id).order_by(MaintenanceRecord.date_performed.desc()).all()
    return render_template('machine_detail.html', machine=m, faults=faults, maintenance=maintenance)

@app.route('/machines/<int:machine_id>/edit', methods=['GET', 'POST'])
@login_required
@role_required('admin')
def machine_edit(machine_id):
    m = Machine.query.get_or_404(machine_id)
    if request.method == 'POST':
        m.name = request.form['name']
        m.description = request.form.get('description', '')
        m.serial_number = request.form.get('serial_number', '')
        m.machine_type = request.form.get('machine_type', '')
        m.manufacturer = request.form.get('manufacturer', '')
        m.year_of_manufacture = int(request.form['year_of_manufacture']) if request.form.get('year_of_manufacture') else None
        m.installation_location = request.form.get('installation_location', '')
        m.contractor_id = int(request.form['contractor_id']) if request.form.get('contractor_id') else None
        m.responsible_user_id = int(request.form['responsible_user_id']) if request.form.get('responsible_user_id') else None
        m.responsible_person_id = int(request.form['responsible_person_id']) if request.form.get('responsible_person_id') else None
        m.section_id = int(request.form['section_id']) if request.form.get('section_id') else None
        m.marker_size = int(request.form.get('marker_size', m.marker_size or 45))
        m.marker_shape = request.form.get('marker_shape', m.marker_shape or 'circle')
        m.status = request.form.get('status', m.status)
        m.floor_x = float(request.form.get('floor_x', m.floor_x))
        m.floor_y = float(request.form.get('floor_y', m.floor_y))
        if 'photo' in request.files and request.files['photo'].filename:
            filename = secure_filename(f"machine_{m.id}_{request.files['photo'].filename}")
            request.files['photo'].save(os.path.join(app.config['UPLOAD_FOLDER'], filename))
            m.photo = filename
        # Update assigned users
        m.assigned_users = []
        for uid in request.form.getlist('assigned_users'):
            u = User.query.get(int(uid))
            if u:
                m.assigned_users.append(u)
        db.session.commit()
        flash(_('Machine updated'), 'success')
        return redirect(url_for('machine_detail', machine_id=m.id))
    users = User.query.filter(User.is_active_user == True).all()
    sections = FactorySection.query.all()
    contractors = Contractor.query.filter_by(is_active=True).order_by(Contractor.company_name).all()
    verantwoordelijken = Verantwoordelijke.query.order_by(Verantwoordelijke.naam).all()
    return render_template('machine_form.html', machine=m, users=users, sections=sections, contractors=contractors, verantwoordelijken=verantwoordelijken)

@app.route('/machines/<int:machine_id>/upload-document', methods=['POST'])
@login_required
@role_required('admin', 'technician')
def machine_upload_document(machine_id):
    m = Machine.query.get_or_404(machine_id)
    if 'document' not in request.files or not request.files['document'].filename:
        flash(_('No file selected'), 'error')
        return redirect(url_for('machine_detail', machine_id=m.id))
    file = request.files['document']
    filename = secure_filename(f"doc_{m.id}_{file.filename}")
    file.save(os.path.join(app.config['UPLOAD_FOLDER'], filename))
    doc = MachineDocument(
        machine_id=m.id,
        doc_type=request.form.get('doc_type', 'other'),
        title=request.form.get('title', file.filename),
        filename=filename,
        uploaded_by=current_user.id
    )
    db.session.add(doc)
    db.session.commit()
    flash(_('Document uploaded'), 'success')
    return redirect(url_for('machine_detail', machine_id=m.id))

@app.route('/machines/<int:machine_id>/add-maintenance', methods=['GET', 'POST'])
@login_required
@role_required('admin', 'technician')
def machine_add_maintenance(machine_id):
    m = Machine.query.get_or_404(machine_id)
    if request.method == 'POST':
        mr = MaintenanceRecord(
            machine_id=m.id,
            maintenance_type=request.form['maintenance_type'],
            description=request.form['description'],
            performed_by=current_user.id,
            date_performed=datetime.strptime(request.form['date_performed'], '%Y-%m-%d'),
            next_maintenance=datetime.strptime(request.form['next_maintenance'], '%Y-%m-%d') if request.form.get('next_maintenance') else None,
            cost=float(request.form.get('cost', 0)),
            parts_used=request.form.get('parts_used', '[]'),
            notes=request.form.get('notes', '')
        )
        db.session.add(mr)
        db.session.commit()
        if 'photos' in request.files:
            for photo in request.files.getlist('photos'):
                if photo.filename:
                    fn = secure_filename(f"maint_{mr.id}_{photo.filename}")
                    photo.save(os.path.join(app.config['UPLOAD_FOLDER'], fn))
                    mp = MaintenancePhoto(maintenance_id=mr.id, filename=fn)
                    db.session.add(mp)
            db.session.commit()
        flash(_('Maintenance record added'), 'success')
        return redirect(url_for('machine_detail', machine_id=m.id))
    return render_template('maintenance_form.html', machine=m, now=datetime.utcnow())

@app.route('/machines/report', methods=['GET', 'POST'])
@login_required
@role_required('admin', 'director')
def machines_report():
    machine_ids = request.args.getlist('machine_ids') or request.form.getlist('machine_ids')
    date_from = request.args.get('date_from') or request.form.get('date_from')
    date_to = request.args.get('date_to') or request.form.get('date_to')
    
    if not date_from:
        date_from = (datetime.utcnow() - timedelta(days=30)).strftime('%Y-%m-%d')
    if not date_to:
        date_to = datetime.utcnow().strftime('%Y-%m-%d')
    
    d_from = datetime.strptime(date_from, '%Y-%m-%d')
    d_to = datetime.strptime(date_to, '%Y-%m-%d') + timedelta(days=1)
    
    if machine_ids:
        machines = Machine.query.filter(Machine.id.in_([int(i) for i in machine_ids])).all()
    else:
        machines = Machine.query.all()

    machine_ids_list = [m.id for m in machines]
    # Bulk-fetch faults and maintenance for all machines in 2 queries
    all_faults = FaultReport.query.filter(
        FaultReport.machine_id.in_(machine_ids_list),
        FaultReport.created_at >= d_from,
        FaultReport.created_at < d_to
    ).all()
    all_maintenance = MaintenanceRecord.query.filter(
        MaintenanceRecord.machine_id.in_(machine_ids_list),
        MaintenanceRecord.date_performed >= d_from,
        MaintenanceRecord.date_performed < d_to
    ).all()
    faults_by_machine = {}
    for f in all_faults:
        faults_by_machine.setdefault(f.machine_id, []).append(f)
    maint_by_machine = {}
    for mr in all_maintenance:
        maint_by_machine.setdefault(mr.machine_id, []).append(mr)

    report_data = []
    for m in machines:
        faults = faults_by_machine.get(m.id, [])
        maintenance = maint_by_machine.get(m.id, [])
        report_data.append({
            'machine': m,
            'faults_total': len(faults),
            'faults_open': len([f for f in faults if f.status in ['open', 'accepted', 'in_progress']]),
            'faults_resolved': len([f for f in faults if f.status in ['resolved', 'closed']]),
            'maintenance_total': len(maintenance),
            'maintenance_cost': sum(rec.cost for rec in maintenance),
            'fault_cost': sum(f.work_report.time_spent_hours * 50 for f in faults if f.work_report)  # estimate
        })
    
    return render_template('machines_report.html', report_data=report_data, 
                         date_from=date_from, date_to=date_to, all_machines=Machine.query.all(),
                         selected_machines=machine_ids)

@app.route('/floor')
@login_required
def floor_plan():
    if current_user.has_role('admin', 'director'):
        machines = Machine.query.all()
        sections = FactorySection.query.all()
    elif hasattr(current_user, '_person') and current_user.role == 'responsible':
        # Responsible person — show only their sections and machines
        person = current_user._person
        sections = list(person.resp_sections)
        section_ids = [s.id for s in sections]
        if person.access_level == 'full':
            machines = Machine.query.all()
        else:
            machines = Machine.query.filter(
                (Machine.responsible_person_id == person.id) |
                (Machine.section_id.in_(section_ids) if section_ids else False)
            ).all()
    elif current_user.has_role('technician'):
        machines = Machine.query.all()
        sections = FactorySection.query.all()
    else:
        machines = current_user.assigned_machines
        section_ids = list(set(m.section_id for m in machines if m.section_id))
        sections = FactorySection.query.filter(FactorySection.id.in_(section_ids)).all() if section_ids else []
    is_filtered = hasattr(current_user, '_person') and current_user.role == 'responsible'
    return render_template('floor_plan.html', machines=machines, sections=sections, is_filtered=is_filtered)

# ============================================================
# ROUTES — FACTORY SECTIONS
# ============================================================

@app.route('/sections')
@login_required
def sections_list():
    sections = FactorySection.query.all()
    return render_template('sections.html', sections=sections)

@app.route('/sections/new', methods=['GET', 'POST'])
@login_required
@role_required('admin')
def section_new():
    if request.method == 'POST':
        s = FactorySection(
            name=request.form['name'],
            description=request.form.get('description', ''),
            section_type=request.form.get('section_type', 'workshop'),
            color=request.form.get('color', '#3498db'),
            floor_x=float(request.form.get('floor_x', 10)),
            floor_y=float(request.form.get('floor_y', 10)),
            width=float(request.form.get('width', 25)),
            height=float(request.form.get('height', 25)),
        )
        s.responsible_user_id = int(request.form['responsible_user_id']) if request.form.get('responsible_user_id') else None
        person_ids = request.form.getlist('responsible_person_ids')
        s.responsible_persons = [Verantwoordelijke.query.get(int(pid)) for pid in person_ids if pid]
        db.session.add(s)
        db.session.commit()
        flash(_('Section created'), 'success')
        return redirect(url_for('sections_list'))
    verantwoordelijken = Verantwoordelijke.query.order_by(Verantwoordelijke.naam).all()
    users = User.query.filter(User.is_active_user == True, User.role.in_(['admin', 'technician', 'director'])).all()
    return render_template('section_form.html', section=None, verantwoordelijken=verantwoordelijken, users=users)

@app.route('/sections/<int:section_id>/edit', methods=['GET', 'POST'])
@login_required
@role_required('admin')
def section_edit(section_id):
    s = FactorySection.query.get_or_404(section_id)
    if request.method == 'POST':
        s.name = request.form['name']
        s.description = request.form.get('description', '')
        s.section_type = request.form.get('section_type', s.section_type)
        s.color = request.form.get('color', s.color)
        s.floor_x = float(request.form.get('floor_x', s.floor_x))
        s.floor_y = float(request.form.get('floor_y', s.floor_y))
        s.width = float(request.form.get('width', s.width))
        s.height = float(request.form.get('height', s.height))
        s.responsible_user_id = int(request.form['responsible_user_id']) if request.form.get('responsible_user_id') else None
        person_ids = request.form.getlist('responsible_person_ids')
        s.responsible_persons = [Verantwoordelijke.query.get(int(pid)) for pid in person_ids if pid]
        db.session.commit()
        flash(_('Section updated'), 'success')
        return redirect(url_for('sections_list'))
    verantwoordelijken = Verantwoordelijke.query.order_by(Verantwoordelijke.naam).all()
    users = User.query.filter(User.is_active_user == True, User.role.in_(['admin', 'technician', 'director'])).all()
    return render_template('section_form.html', section=s, verantwoordelijken=verantwoordelijken, users=users)

@app.route('/sections/<int:section_id>/delete', methods=['POST'])
@login_required
@role_required('admin')
def section_delete(section_id):
    s = FactorySection.query.get_or_404(section_id)
    for m in s.machines:
        m.section_id = None
    db.session.delete(s)
    db.session.commit()
    flash(_('Section deleted'), 'success')
    return redirect(url_for('sections_list'))

@app.route('/api/sections/<int:section_id>')
@login_required
def api_section_info(section_id):
    s = FactorySection.query.get_or_404(section_id)
    machines_data = []
    for m in s.machines:
        machines_data.append({
            'id': m.id, 'name': m.name, 'status': m.status,
            'type': m.machine_type or '', 'serial': m.serial_number or ''
        })
    return jsonify({
        'id': s.id, 'name': s.name, 'description': s.description or '',
        'type': s.section_type, 'color': s.color,
        'responsible': ', '.join(p.naam for p in s.responsible_persons) if s.responsible_persons else None,
        'responsible_ids': [p.id for p in s.responsible_persons],
        'machines': machines_data,
        'machines_count': len(machines_data),
        'active_count': len([m for m in s.machines if m.status == 'active']),
        'broken_count': len([m for m in s.machines if m.status == 'broken'])
    })

@app.route('/map-editor')
@login_required
@role_required('admin')
def map_editor():
    sections = FactorySection.query.all()
    machines = Machine.query.all()
    verantwoordelijken = Verantwoordelijke.query.order_by(Verantwoordelijke.naam).all()
    return render_template('map_editor.html', sections=sections, machines=machines, verantwoordelijken=verantwoordelijken)

@app.route('/settings')
@login_required
@role_required('admin')
def settings():
    users = User.query.order_by(User.display_name).all()
    sections = FactorySection.query.order_by(FactorySection.name).all()
    groups = ResponsibleGroup.query.order_by(ResponsibleGroup.name).all()
    responsible = Verantwoordelijke.query.order_by(Verantwoordelijke.naam).all()
    machines = Machine.query.order_by(Machine.name).all()
    
    # Fault statistics for settings page
    faults_by_priority = {}
    for p in ['low', 'normal', 'high', 'critical']:
        faults_by_priority[p] = FaultReport.query.filter_by(priority=p).filter(
            FaultReport.status.in_(['open', 'accepted', 'in_progress', 'parts_ordered', 'waiting_parts', 'reopened'])
        ).count()
    
    faults_by_status = {}
    for s in ['open', 'accepted', 'in_progress', 'parts_ordered', 'waiting_parts', 'resolved', 'closed', 'reopened']:
        faults_by_status[s] = FaultReport.query.filter_by(status=s).count()
    
    # Top machines with faults
    top_machines_faults = []
    for m in machines:
        fault_count = FaultReport.query.filter_by(machine_id=m.id).count()
        if fault_count > 0:
            open_count = FaultReport.query.filter(FaultReport.machine_id == m.id, FaultReport.status.in_(['open', 'accepted', 'in_progress'])).count()
            critical_count = FaultReport.query.filter(FaultReport.machine_id == m.id, FaultReport.priority == 'critical').count()
            m.fault_count = fault_count
            m.open_count = open_count
            m.critical_count = critical_count
            top_machines_faults.append(m)
    top_machines_faults.sort(key=lambda x: x.fault_count, reverse=True)
    
    return render_template('settings.html', users=users, sections=sections, groups=groups,
        responsible=responsible, machines=machines,
        faults_by_priority=faults_by_priority, faults_by_status=faults_by_status,
        top_machines_faults=top_machines_faults)

@app.route('/settings/user/<int:user_id>/access', methods=['POST'])
@login_required
@role_required('admin')
def settings_user_access(user_id):
    u = User.query.get_or_404(user_id)
    data = request.get_json()
    # Update role
    if 'role' in data:
        u.role = data['role']
    # Update section access
    if 'sections' in data:
        u.allowed_sections = []
        for key in data['sections']:
            db.session.add(UserSectionAccess(user_id=u.id, section_key=key))
    # Update active status
    if 'is_active' in data:
        u.is_active_user = data['is_active']
    db.session.commit()
    log_audit('update', 'user_access', u.id, f'Updated access for {u.username}')
    return jsonify({'ok': True})

@app.route('/settings/section/<int:section_id>/update', methods=['POST'])
@login_required
@role_required('admin')
def settings_section_update(section_id):
    s = FactorySection.query.get_or_404(section_id)
    data = request.get_json()
    if 'name' in data:
        s.name = data['name']
    if 'color' in data:
        s.color = data['color']
    if 'floor_x' in data:
        s.floor_x = data['floor_x']
    if 'floor_y' in data:
        s.floor_y = data['floor_y']
    if 'width' in data:
        s.width = data['width']
    if 'height' in data:
        s.height = data['height']
    db.session.commit()
    return jsonify({'ok': True})

@app.route('/api/map/save-section', methods=['POST'])
@login_required
@role_required('admin')
def api_map_save_section():
    data = request.get_json()
    section_id = data.get('id')
    if section_id:
        s = FactorySection.query.get_or_404(section_id)
        s.name = data.get('name', s.name)
        s.description = data.get('description', s.description)
        s.section_type = data.get('section_type', s.section_type)
        s.color = data.get('color', s.color)
        s.floor_x = data.get('x', s.floor_x)
        s.floor_y = data.get('y', s.floor_y)
        s.width = data.get('w', s.width)
        s.height = data.get('h', s.height)
        resp_ids = data.get('responsible_ids', [])
        s.responsible_persons = [Verantwoordelijke.query.get(int(pid)) for pid in resp_ids if pid]
    else:
        s = FactorySection(
            name=data.get('name', 'New Section'),
            description=data.get('description', ''),
            section_type=data.get('section_type', 'workshop'),
            color=data.get('color', '#3498db'),
            floor_x=data.get('x', 10),
            floor_y=data.get('y', 10),
            width=data.get('w', 25),
            height=data.get('h', 25),
        )
        db.session.add(s)
        db.session.flush()
        resp_ids = data.get('responsible_ids', [])
        s.responsible_persons = [Verantwoordelijke.query.get(int(pid)) for pid in resp_ids if pid]
    db.session.commit()
    return jsonify({'ok': True, 'id': s.id})

@app.route('/api/map/delete-section', methods=['POST'])
@login_required
@role_required('admin')
def api_map_delete_section():
    data = request.get_json()
    s = FactorySection.query.get_or_404(data['id'])
    for m in s.machines:
        m.section_id = None
    db.session.delete(s)
    db.session.commit()
    return jsonify({'ok': True})

@app.route('/api/map/save-machine-pos', methods=['POST'])
@login_required
@role_required('admin')
def api_map_save_machine_pos():
    data = request.get_json()
    m = Machine.query.get_or_404(data['id'])
    m.floor_x = data['x']
    m.floor_y = data['y']
    db.session.commit()
    return jsonify({'ok': True})

@app.route('/api/map/assign-machine', methods=['POST'])
@login_required
@role_required('admin')
def api_map_assign_machine():
    data = request.get_json()
    m = Machine.query.get_or_404(data['machine_id'])
    m.section_id = data.get('section_id')
    db.session.commit()
    return jsonify({'ok': True})

@app.route('/api/map/all')
@login_required
def api_map_all():
    sections = []
    for s in FactorySection.query.all():
        sections.append({
            'id': s.id, 'name': s.name, 'description': s.description or '',
            'type': s.section_type, 'color': s.color,
            'x': s.floor_x, 'y': s.floor_y, 'w': s.width, 'h': s.height,
            'responsible_ids': [p.id for p in s.responsible_persons],
            'responsible': ', '.join(p.naam for p in s.responsible_persons) if s.responsible_persons else None
        })
    machines = []
    for m in Machine.query.all():
        machines.append({
            'id': m.id, 'name': m.name, 'status': m.status,
            'type': m.machine_type or '', 'serial': m.serial_number or '',
            'x': m.floor_x, 'y': m.floor_y, 'section_id': m.section_id
        })
    return jsonify({'sections': sections, 'machines': machines})

# ============================================================
# ROUTES — MACHINE PARTS MAINTENANCE
# ============================================================

@app.route('/machines/<int:machine_id>/parts')
@login_required
def machine_parts(machine_id):
    m = Machine.query.get_or_404(machine_id)
    parts = MachinePart.query.filter_by(machine_id=m.id).order_by(MachinePart.name).all()
    today = datetime.utcnow().date()

    # Auto-update statuses and send reminders
    for p in parts:
        old_status = p.status
        if p.next_replacement and p.next_replacement <= today:
            p.status = 'needs_replacement'
        elif p.next_maintenance and p.next_maintenance <= today:
            p.status = 'needs_maintenance'
        elif p.status in ('needs_replacement', 'needs_maintenance'):
            p.status = 'ok'
        # Send notification to responsible person if status changed
        if p.status != old_status and p.status in ('needs_replacement', 'needs_maintenance'):
            target_user_id = p.responsible_user_id or current_user.id
            notif_title = f"{'Replacement' if p.status == 'needs_replacement' else 'Maintenance'} needed: {p.name}"
            notif_msg = f"{m.name}: {p.name} - {p.status.replace('_', ' ')}"
            existing = Notification.query.filter_by(user_id=target_user_id, is_read=False, title=notif_title).first()
            if not existing:
                create_notification(target_user_id, notif_title, notif_msg, 'warning', url_for('machine_parts', machine_id=m.id))
    db.session.commit()

    users = User.query.filter(User.is_active_user == True, User.role.in_(['admin', 'technician'])).all()
    return render_template('machine_parts.html', machine=m, parts=parts, today=today, users=users)

@app.route('/machines/<int:machine_id>/parts/new', methods=['GET', 'POST'])
@login_required
@role_required('admin', 'technician')
def machine_part_new(machine_id):
    m = Machine.query.get_or_404(machine_id)
    if request.method == 'POST':
        installed = datetime.strptime(request.form['installed_date'], '%Y-%m-%d').date() if request.form.get('installed_date') else None
        interval = int(request.form['replacement_interval_days']) if request.form.get('replacement_interval_days') else None
        maint_interval = int(request.form['maintenance_interval_days']) if request.form.get('maintenance_interval_days') else None
        
        p = MachinePart(
            machine_id=m.id,
            name=request.form['name'],
            part_number=request.form.get('part_number', ''),
            description=request.form.get('description', ''),
            category=request.form.get('category', 'mechanical'),
            installed_date=installed,
            last_replacement=installed,
            replacement_interval_days=interval,
            next_replacement=date_plus_days(installed, interval) if installed and interval else None,
            maintenance_interval_days=maint_interval,
            next_maintenance=date_plus_days(installed, maint_interval) if installed and maint_interval else None,
            responsible_user_id=int(request.form['responsible_user_id']) if request.form.get('responsible_user_id') else None,
        )
        db.session.add(p)
        db.session.commit()
        flash(_('Part added'), 'success')
        return redirect(url_for('machine_parts', machine_id=m.id))
    users = User.query.filter(User.is_active_user == True, User.role.in_(['admin', 'technician'])).all()
    return render_template('part_form.html', machine=m, part=None, users=users)

@app.route('/machines/<int:machine_id>/parts/<int:part_id>/edit', methods=['GET', 'POST'])
@login_required
@role_required('admin', 'technician')
def machine_part_edit(machine_id, part_id):
    m = Machine.query.get_or_404(machine_id)
    p = MachinePart.query.get_or_404(part_id)
    if request.method == 'POST':
        p.name = request.form['name']
        p.part_number = request.form.get('part_number', '')
        p.description = request.form.get('description', '')
        p.category = request.form.get('category', p.category)
        p.status = request.form.get('status', p.status)
        p.notes = request.form.get('notes', '')
        
        installed = datetime.strptime(request.form['installed_date'], '%Y-%m-%d').date() if request.form.get('installed_date') else p.installed_date
        p.installed_date = installed
        p.replacement_interval_days = int(request.form['replacement_interval_days']) if request.form.get('replacement_interval_days') else p.replacement_interval_days
        p.maintenance_interval_days = int(request.form['maintenance_interval_days']) if request.form.get('maintenance_interval_days') else p.maintenance_interval_days
        p.responsible_user_id = int(request.form['responsible_user_id']) if request.form.get('responsible_user_id') else None

        if installed and p.replacement_interval_days:
            p.next_replacement = date_plus_days(installed, p.replacement_interval_days)
        if installed and p.maintenance_interval_days:
            p.next_maintenance = date_plus_days(installed, p.maintenance_interval_days)

        db.session.commit()
        flash(_('Part updated'), 'success')
        return redirect(url_for('machine_parts', machine_id=m.id))
    users = User.query.filter(User.is_active_user == True, User.role.in_(['admin', 'technician'])).all()
    return render_template('part_form.html', machine=m, part=p, users=users)

@app.route('/machines/<int:machine_id>/parts/<int:part_id>/log', methods=['POST'])
@login_required
@role_required('admin', 'technician')
def machine_part_log(machine_id, part_id):
    m = Machine.query.get_or_404(machine_id)
    p = MachinePart.query.get_or_404(part_id)
    action = request.form['action']
    today = datetime.utcnow()
    
    log = PartMaintenanceLog(
        part_id=p.id,
        action=action,
        description=request.form.get('description', ''),
        performed_by=int(request.form['performed_by']) if request.form.get('performed_by') else current_user.id,
        date=today,
        cost=float(request.form.get('cost', 0)),
        notes=request.form.get('notes', '')
    )
    db.session.add(log)
    
    if action == 'replacement':
        p.last_replacement = today.date()
        p.last_maintenance = today.date()
        p.status = 'ok'
        if p.replacement_interval_days:
            p.next_replacement = date_plus_days(today.date(), p.replacement_interval_days)
        if p.maintenance_interval_days:
            p.next_maintenance = date_plus_days(today.date(), p.maintenance_interval_days)
    elif action == 'maintenance':
        p.last_maintenance = today.date()
        p.status = 'ok'
        if p.maintenance_interval_days:
            p.next_maintenance = date_plus_days(today.date(), p.maintenance_interval_days)
    
    db.session.commit()
    flash(_('Maintenance log added'), 'success')
    return redirect(url_for('machine_parts', machine_id=m.id))

@app.route('/machines/<int:machine_id>/parts/<int:part_id>/delete', methods=['POST'])
@login_required
@role_required('admin')
def machine_part_delete(machine_id, part_id):
    m = Machine.query.get_or_404(machine_id)
    p = MachinePart.query.get_or_404(part_id)
    PartMaintenanceLog.query.filter_by(part_id=p.id).delete()
    db.session.delete(p)
    db.session.commit()
    flash(_('Part deleted'), 'success')
    return redirect(url_for('machine_parts', machine_id=m.id))

# ============================================================
# ROUTES — ELECTRICITY / ELECTRICAL CABINETS
# ============================================================

@app.route('/electricity')
@login_required
def electricity_list():
    cabinets = ElectricalCabinet.query.filter_by(is_active=True).order_by(ElectricalCabinet.name).all()
    cab_map = {c.id: c for c in cabinets}
    all_breakers = [b for cab in cabinets for b in cab.breakers]
    stats = {
        'total': len(all_breakers),
        'on': sum(1 for b in all_breakers if b.status == 'on'),
        'off': sum(1 for b in all_breakers if b.status == 'off'),
    }
    # 4 physical panels from photos
    panels = [
        {'name': 'Panel 1 — VRACHTWAGEN / VERPAKKING / WERKVOORBEREIDING',
         'cabinet_ids': [3, 4, 5, 6]},
        {'name': 'Panel 2 — CANALIS / KOUDE / WATER / TECHNIEKEN',
         'cabinet_ids': [7, 8, 9, 10, 11]},
        {'name': 'Panel 3 — LOGISTIEK',
         'cabinet_ids': [12, 13]},
        {'name': 'Panel 4 — DISTRIBUTIE / UGL / ALGEMEEN / WB',
         'cabinet_ids': [14, 15, 16, 17, 18]},
    ]
    for p in panels:
        p['cabinets'] = [cab_map[cid] for cid in p['cabinet_ids'] if cid in cab_map]
        p['total_breakers'] = sum(len(c.breakers) for c in p['cabinets'])
        p['on_count'] = sum(1 for c in p['cabinets'] for b in c.breakers if b.status == 'on')
    return render_template('electricity.html', cabinets=cabinets, stats=stats, panels=panels)

@app.route('/electricity/cabinet/new', methods=['GET', 'POST'])
@login_required
@role_required('admin', 'director', 'technician')
def cabinet_new():
    if request.method == 'POST':
        c = ElectricalCabinet(
            name=request.form['name'],
            location=request.form.get('location', ''),
            cabinet_type=request.form.get('cabinet_type', 'distribution'),
            description=request.form.get('description', ''),
            manufacturer=request.form.get('manufacturer', ''),
            serial_number=request.form.get('serial_number', ''),
            main_fuse_amps=int(request.form['main_fuse_amps']) if request.form.get('main_fuse_amps') else None,
            voltage=request.form.get('voltage', '400V/230V'),
            schematic_x=int(request.form.get('schematic_x', 0)),
            schematic_y=int(request.form.get('schematic_y', 0))
        )
        if 'photo' in request.files and request.files['photo'].filename:
            filename = secure_filename(f"cabinet_{request.files['photo'].filename}")
            request.files['photo'].save(os.path.join(app.config['UPLOAD_FOLDER'], filename))
            c.photo = filename
        db.session.add(c)
        db.session.commit()
        flash(_('Cabinet created'), 'success')
        return redirect(url_for('cabinet_detail', cabinet_id=c.id))
    return render_template('cabinet_form.html', cabinet=None)

@app.route('/electricity/cabinet/<int:cabinet_id>')
@login_required
def cabinet_detail(cabinet_id):
    c = ElectricalCabinet.query.get_or_404(cabinet_id)
    return render_template('cabinet_detail.html', cabinet=c)

@app.route('/electricity/cabinet/<int:cabinet_id>/edit', methods=['GET', 'POST'])
@login_required
@role_required('admin', 'director', 'technician')
def cabinet_edit(cabinet_id):
    c = ElectricalCabinet.query.get_or_404(cabinet_id)
    if request.method == 'POST':
        c.name = request.form['name']
        c.location = request.form.get('location', '')
        c.cabinet_type = request.form.get('cabinet_type', 'distribution')
        c.description = request.form.get('description', '')
        c.manufacturer = request.form.get('manufacturer', '')
        c.serial_number = request.form.get('serial_number', '')
        c.main_fuse_amps = int(request.form['main_fuse_amps']) if request.form.get('main_fuse_amps') else None
        c.voltage = request.form.get('voltage', '400V/230V')
        c.schematic_x = int(request.form.get('schematic_x', 0))
        c.schematic_y = int(request.form.get('schematic_y', 0))
        if 'photo' in request.files and request.files['photo'].filename:
            filename = secure_filename(f"cabinet_{request.files['photo'].filename}")
            request.files['photo'].save(os.path.join(app.config['UPLOAD_FOLDER'], filename))
            c.photo = filename
        db.session.commit()
        flash(_('Cabinet updated'), 'success')
        return redirect(url_for('cabinet_detail', cabinet_id=c.id))
    return render_template('cabinet_form.html', cabinet=c)

@app.route('/electricity/cabinet/<int:cabinet_id>/delete', methods=['POST'])
@login_required
@role_required('admin')
def cabinet_delete(cabinet_id):
    c = ElectricalCabinet.query.get_or_404(cabinet_id)
    c.is_active = False
    db.session.commit()
    flash(_('Cabinet deleted'), 'success')
    return redirect(url_for('electricity_list'))

@app.route('/electricity/cabinet/<int:cabinet_id>/breaker/new', methods=['GET', 'POST'])
@login_required
@role_required('admin', 'director', 'technician')
def breaker_new(cabinet_id):
    c = ElectricalCabinet.query.get_or_404(cabinet_id)
    if request.method == 'POST':
        b = CircuitBreaker(
            cabinet_id=c.id,
            label=request.form['label'],
            description=request.form.get('description', ''),
            breaker_type=request.form.get('breaker_type', 'MCB'),
            amperage=int(request.form['amperage']) if request.form.get('amperage') else None,
            poles=int(request.form.get('poles', 1)),
            curve_type=request.form.get('curve_type', 'C'),
            phase=request.form.get('phase', ''),
            status=request.form.get('status', 'on'),
            connected_to=request.form.get('connected_to', ''),
            notes=request.form.get('notes', ''),
            row=int(request.form.get('row', 1)),
            position=int(request.form.get('position', 1))
        )
        db.session.add(b)
        db.session.commit()
        flash(_('Breaker added'), 'success')
        return redirect(url_for('cabinet_detail', cabinet_id=c.id))
    return render_template('breaker_form.html', cabinet=c, breaker=None)

@app.route('/electricity/breaker/<int:breaker_id>/edit', methods=['GET', 'POST'])
@login_required
@role_required('admin', 'director', 'technician')
def breaker_edit(breaker_id):
    b = CircuitBreaker.query.get_or_404(breaker_id)
    if request.method == 'POST':
        b.label = request.form['label']
        b.description = request.form.get('description', '')
        b.breaker_type = request.form.get('breaker_type', 'MCB')
        b.amperage = int(request.form['amperage']) if request.form.get('amperage') else None
        b.poles = int(request.form.get('poles', 1))
        b.curve_type = request.form.get('curve_type', 'C')
        b.phase = request.form.get('phase', '')
        b.status = request.form.get('status', 'on')
        b.connected_to = request.form.get('connected_to', '')
        b.notes = request.form.get('notes', '')
        b.row = int(request.form.get('row', 1))
        b.position = int(request.form.get('position', 1))
        db.session.commit()
        flash(_('Breaker updated'), 'success')
        return redirect(url_for('cabinet_detail', cabinet_id=b.cabinet_id))
    return render_template('breaker_form.html', cabinet=b.cabinet, breaker=b)

@app.route('/electricity/breaker/<int:breaker_id>/toggle', methods=['POST'])
@login_required
@role_required('admin', 'director', 'technician')
def breaker_toggle(breaker_id):
    b = CircuitBreaker.query.get_or_404(breaker_id)
    b.status = 'off' if b.status == 'on' else 'on'
    db.session.commit()
    return jsonify({'status': b.status})

@app.route('/electricity/breaker/<int:breaker_id>/delete', methods=['POST'])
@login_required
@role_required('admin')
def breaker_delete(breaker_id):
    b = CircuitBreaker.query.get_or_404(breaker_id)
    cabinet_id = b.cabinet_id
    db.session.delete(b)
    db.session.commit()
    flash(_('Breaker deleted'), 'success')
    return redirect(url_for('cabinet_detail', cabinet_id=cabinet_id))

@app.route('/electricity/schematic')
@login_required
def electricity_schematic():
    cabinets = ElectricalCabinet.query.filter_by(is_active=True).order_by(ElectricalCabinet.name).all()
    return render_template('electricity_schematic.html', cabinets=cabinets)

# ============================================================
# ROUTES — MAINTENANCE CALENDAR
# ============================================================

@app.route('/maintenance-calendar')
@login_required
def maintenance_calendar():
    today = datetime.utcnow().date()
    month = request.args.get('month', today.strftime('%Y-%m'))
    year, mon = map(int, month.split('-'))
    month_start = datetime(year, mon, 1).date()
    if mon == 12:
        month_end = datetime(year + 1, 1, 1).date()
    else:
        month_end = datetime(year, mon + 1, 1).date()
    
    # Get all parts with upcoming maintenance/replacement (eagerly load machine)
    from sqlalchemy.orm import joinedload
    if current_user.has_role('admin', 'director', 'technician'):
        parts = MachinePart.query.options(joinedload(MachinePart.machine)).all()
    else:
        machine_ids = [m.id for m in current_user.assigned_machines]
        parts = MachinePart.query.options(joinedload(MachinePart.machine)).filter(MachinePart.machine_id.in_(machine_ids)).all()
    
    # Build calendar events
    events = []
    for p in parts:
        if p.next_replacement and month_start <= p.next_replacement < month_end:
            events.append({
                'date': p.next_replacement,
                'type': 'replacement',
                'part': p.name,
                'machine': p.machine.name,
                'machine_id': p.machine_id,
                'part_id': p.id,
                'category': p.category,
                'overdue': p.next_replacement < today
            })
        if p.next_maintenance and month_start <= p.next_maintenance < month_end:
            events.append({
                'date': p.next_maintenance,
                'type': 'maintenance',
                'part': p.name,
                'machine': p.machine.name,
                'machine_id': p.machine_id,
                'part_id': p.id,
                'category': p.category,
                'overdue': p.next_maintenance < today
            })
        # Also check maintenance records
        for mr in MaintenanceRecord.query.filter_by(machine_id=p.machine_id).all():
            if mr.next_maintenance and month_start <= mr.next_maintenance.date() < month_end:
                events.append({
                    'date': mr.next_maintenance.date(),
                    'type': 'machine_maintenance',
                    'part': mr.description[:40],
                    'machine': p.machine.name,
                    'machine_id': p.machine_id,
                    'part_id': None,
                    'category': mr.maintenance_type,
                    'overdue': mr.next_maintenance.date() < today
                })
    
    # Add maintenance plans
    if current_user.has_role('admin', 'director', 'technician'):
        plans = MaintenancePlan.query.all()
    else:
        plan_machine_ids = [m.id for m in current_user.assigned_machines]
        plans = MaintenancePlan.query.filter(MaintenancePlan.machine_id.in_(plan_machine_ids)).all()
    for pl in plans:
        if pl.planned_start and month_start <= pl.planned_start < month_end:
            events.append({
                'date': pl.planned_start,
                'type': 'plan',
                'part': pl.title[:40],
                'machine': pl.machine.name,
                'machine_id': pl.machine_id,
                'part_id': None,
                'category': pl.maintenance_type,
                'overdue': pl.planned_start < today and pl.status not in ('completed', 'cancelled'),
                'plan_id': pl.id,
                'status': pl.status
            })

    # Get overdue items (before this month)
    overdue = []
    for p in parts:
        if p.next_replacement and p.next_replacement < today:
            overdue.append({'date': p.next_replacement, 'type': 'replacement', 'part': p.name, 'machine': p.machine.name, 'machine_id': p.machine_id, 'part_id': p.id, 'category': p.category})
        if p.next_maintenance and p.next_maintenance < today:
            overdue.append({'date': p.next_maintenance, 'type': 'maintenance', 'part': p.name, 'machine': p.machine.name, 'machine_id': p.machine_id, 'part_id': p.id, 'category': p.category})
    
    # Navigation
    prev_month = (month_start - timedelta(days=1)).strftime('%Y-%m')
    next_month = month_end.strftime('%Y-%m')
    
    return render_template('maintenance_calendar.html',
        month=month, month_start=month_start, month_end=month_end,
        events=sorted(events, key=lambda e: e['date']),
        overdue=overdue, today=today,
        prev_month=prev_month, next_month=next_month,
        timedelta=timedelta)

# ============================================================
# ROUTES — MAINTENANCE PLANS
# ============================================================

@app.route('/maintenance-plans')
@login_required
def maintenance_plans_list():
    if current_user.has_role('admin', 'director', 'technician'):
        plans = MaintenancePlan.query.order_by(MaintenancePlan.planned_start.desc()).all()
    else:
        machine_ids = [m.id for m in current_user.assigned_machines]
        plans = MaintenancePlan.query.filter(MaintenancePlan.machine_id.in_(machine_ids)).order_by(MaintenancePlan.planned_start.desc()).all()
    machines = Machine.query.order_by(Machine.name).all()
    return render_template('maintenance_plans.html', plans=plans, machines=machines)

@app.route('/maintenance-plans/new', methods=['GET', 'POST'])
@login_required
@role_required('admin', 'technician')
def maintenance_plan_new():
    if request.method == 'POST':
        p = MaintenancePlan(
            machine_id=int(request.form['machine_id']),
            title=request.form['title'],
            description=request.form.get('description', ''),
            maintenance_type=request.form.get('maintenance_type', 'preventive'),
            status=request.form.get('status', 'planned'),
            planned_start=datetime.strptime(request.form['planned_start'], '%Y-%m-%d').date(),
            planned_end=datetime.strptime(request.form['planned_end'], '%Y-%m-%d').date() if request.form.get('planned_end') else None,
            is_external='is_external' in request.form,
            company_name=request.form.get('company_name', ''),
            company_contact=request.form.get('company_contact', ''),
            company_person=request.form.get('company_person', ''),
            worker_id=int(request.form['worker_id']) if request.form.get('worker_id') else None,
            parts_used=request.form.get('parts_used', '[]'),
            cost=float(request.form.get('cost', 0)),
            report=request.form.get('report', ''),
            next_maintenance=datetime.strptime(request.form['next_maintenance'], '%Y-%m-%d').date() if request.form.get('next_maintenance') else None,
            notes=request.form.get('notes', ''),
            created_by=current_user.id
        )
        if 'offer_file' in request.files and request.files['offer_file'].filename:
            fn = secure_filename(f"offer_{request.files['offer_file'].filename}")
            request.files['offer_file'].save(os.path.join(app.config['UPLOAD_FOLDER'], fn))
            p.offer_file = fn
        if 'work_act_file' in request.files and request.files['work_act_file'].filename:
            fn = secure_filename(f"act_{request.files['work_act_file'].filename}")
            request.files['work_act_file'].save(os.path.join(app.config['UPLOAD_FOLDER'], fn))
            p.work_act_file = fn
        db.session.add(p)
        db.session.commit()
        flash(_('Maintenance plan created'), 'success')
        return redirect(url_for('maintenance_plan_detail', plan_id=p.id))
    machines = Machine.query.order_by(Machine.name).all()
    workers = Monteur.query.filter_by(actief=True).all()
    return render_template('maintenance_plan_form.html', plan=None, machines=machines, workers=workers)

@app.route('/maintenance-plans/<int:plan_id>')
@login_required
def maintenance_plan_detail(plan_id):
    p = MaintenancePlan.query.get_or_404(plan_id)
    return render_template('maintenance_plan_detail.html', plan=p)

@app.route('/maintenance-plans/<int:plan_id>/edit', methods=['GET', 'POST'])
@login_required
@role_required('admin', 'technician')
def maintenance_plan_edit(plan_id):
    p = MaintenancePlan.query.get_or_404(plan_id)
    if request.method == 'POST':
        p.machine_id = int(request.form['machine_id'])
        p.title = request.form['title']
        p.description = request.form.get('description', '')
        p.maintenance_type = request.form.get('maintenance_type', p.maintenance_type)
        p.status = request.form.get('status', p.status)
        p.planned_start = datetime.strptime(request.form['planned_start'], '%Y-%m-%d').date()
        p.planned_end = datetime.strptime(request.form['planned_end'], '%Y-%m-%d').date() if request.form.get('planned_end') else None
        p.actual_start = datetime.strptime(request.form['actual_start'], '%Y-%m-%d').date() if request.form.get('actual_start') else None
        p.actual_end = datetime.strptime(request.form['actual_end'], '%Y-%m-%d').date() if request.form.get('actual_end') else None
        p.is_external = 'is_external' in request.form
        p.company_name = request.form.get('company_name', '')
        p.company_contact = request.form.get('company_contact', '')
        p.company_person = request.form.get('company_person', '')
        p.worker_id = int(request.form['worker_id']) if request.form.get('worker_id') else None
        p.parts_used = request.form.get('parts_used', '[]')
        p.cost = float(request.form.get('cost', 0))
        p.report = request.form.get('report', '')
        p.next_maintenance = datetime.strptime(request.form['next_maintenance'], '%Y-%m-%d').date() if request.form.get('next_maintenance') else None
        p.notes = request.form.get('notes', '')
        if 'offer_file' in request.files and request.files['offer_file'].filename:
            fn = secure_filename(f"offer_{request.files['offer_file'].filename}")
            request.files['offer_file'].save(os.path.join(app.config['UPLOAD_FOLDER'], fn))
            p.offer_file = fn
        if 'work_act_file' in request.files and request.files['work_act_file'].filename:
            fn = secure_filename(f"act_{request.files['work_act_file'].filename}")
            request.files['work_act_file'].save(os.path.join(app.config['UPLOAD_FOLDER'], fn))
            p.work_act_file = fn
        db.session.commit()
        flash(_('Maintenance plan updated'), 'success')
        return redirect(url_for('maintenance_plan_detail', plan_id=p.id))
    machines = Machine.query.order_by(Machine.name).all()
    workers = Monteur.query.filter_by(actief=True).all()
    return render_template('maintenance_plan_form.html', plan=p, machines=machines, workers=workers)

@app.route('/maintenance-plans/<int:plan_id>/delete', methods=['POST'])
@login_required
@role_required('admin')
def maintenance_plan_delete(plan_id):
    p = MaintenancePlan.query.get_or_404(plan_id)
    db.session.delete(p)
    db.session.commit()
    flash(_('Maintenance plan deleted'), 'success')
    return redirect(url_for('maintenance_plans_list'))

@app.route('/api/maintenance/plans')
@login_required
def api_maintenance_plans():
    today = datetime.utcnow().date()
    month = request.args.get('month', today.strftime('%Y-%m'))
    year, mon = map(int, month.split('-'))
    month_start = datetime(year, mon, 1).date()
    if mon == 12:
        month_end = datetime(year + 1, 1, 1).date()
    else:
        month_end = datetime(year, mon + 1, 1).date()

    if current_user.has_role('admin', 'director', 'technician'):
        plans = MaintenancePlan.query.all()
    else:
        machine_ids = [m.id for m in current_user.assigned_machines]
        plans = MaintenancePlan.query.filter(MaintenancePlan.machine_id.in_(machine_ids)).all()

    result = []
    for p in plans:
        if p.planned_start and month_start <= p.planned_start < month_end:
            result.append({
                'id': p.id, 'title': p.title, 'machine': p.machine.name,
                'machine_id': p.machine_id, 'type': p.maintenance_type,
                'status': p.status, 'start': p.planned_start.isoformat(),
                'end': p.planned_end.isoformat() if p.planned_end else None,
                'is_external': p.is_external,
                'company': p.company_name if p.is_external else None,
                'worker': p.worker.naam if p.worker else None,
            })
    return jsonify(result)

@app.route('/api/maintenance/reminders')
@login_required
def api_maintenance_reminders():
    today = datetime.utcnow().date()
    soon = today + timedelta(days=14)
    
    if current_user.has_role('admin', 'director', 'technician'):
        parts = MachinePart.query.all()
    else:
        machine_ids = [m.id for m in current_user.assigned_machines]
        parts = MachinePart.query.filter(MachinePart.machine_id.in_(machine_ids)).all()
    
    reminders = []
    for p in parts:
        if p.next_replacement and p.next_replacement <= soon:
            days_left = (p.next_replacement - today).days
            reminders.append({
                'type': 'replacement', 'part': p.name, 'machine': p.machine.name,
                'machine_id': p.machine_id, 'part_id': p.id,
                'date': p.next_replacement.isoformat(),
                'days_left': days_left, 'overdue': days_left < 0
            })
        if p.next_maintenance and p.next_maintenance <= soon:
            days_left = (p.next_maintenance - today).days
            reminders.append({
                'type': 'maintenance', 'part': p.name, 'machine': p.machine.name,
                'machine_id': p.machine_id, 'part_id': p.id,
                'date': p.next_maintenance.isoformat(),
                'days_left': days_left, 'overdue': days_left < 0
            })
    
    reminders.sort(key=lambda r: r['date'])
    return jsonify(reminders)

# ============================================================
# ROUTES — MULE MAINTENANCE
# ============================================================

def gen_mule_number():
    vandaag = datetime.utcnow()
    prefix = vandaag.strftime('%Y%m%d')
    laatste = MuleMaintenance.query.filter(MuleMaintenance.number.like(f'MUL-{prefix}-%')).order_by(MuleMaintenance.id.desc()).first()
    if laatste:
        num = int(laatste.number.split('-')[2]) + 1
    else:
        num = 1
    return f'MUL-{prefix}-{num:04d}'

def gen_mule_order_number():
    vandaag = datetime.utcnow()
    prefix = vandaag.strftime('%Y%m%d')
    laatste = MulePartOrder.query.filter(MulePartOrder.order_number.like(f'MPO-{prefix}-%')).order_by(MulePartOrder.id.desc()).first()
    if laatste:
        num = int(laatste.order_number.split('-')[2]) + 1
    else:
        num = 1
    return f'MPO-{prefix}-{num:04d}'

@app.route('/mule')
@login_required
@role_required('admin', 'director', 'technician')
def mule_list():
    serial_filter = request.args.get('serial', '').strip()
    q = MuleMaintenance.query
    if serial_filter:
        q = q.filter(MuleMaintenance.mule_serial.ilike(f'%{serial_filter}%'))
    maintenance = q.order_by(MuleMaintenance.date.desc()).all()
    orders = MulePartOrder.query.order_by(MulePartOrder.ordered_at.desc()).limit(20).all()
    machines = Machine.query.order_by(Machine.name).all()
    return render_template('mule.html', maintenance=maintenance, orders=orders, machines=machines, serial_filter=serial_filter)

@app.route('/mule/new', methods=['GET', 'POST'])
@login_required
@role_required('admin', 'technician')
def mule_new():
    if request.method == 'POST':
        m = MuleMaintenance(
            number=gen_mule_number(),
            mule_number=request.form['mule_number'],
            mule_serial=request.form.get('mule_serial', ''),
            machine_id=int(request.form['machine_id']) if request.form.get('machine_id') else None,
            date=datetime.strptime(request.form['date'], '%Y-%m-%d').date(),
            reason=request.form['reason'],
            next_date=datetime.strptime(request.form['next_date'], '%Y-%m-%d').date() if request.form.get('next_date') else None,
            periodicity=request.form.get('periodicity', ''),
            notes=request.form.get('notes', ''),
            created_by=current_user.id
        )
        db.session.add(m)
        db.session.flush()
        # Add parts
        part_names = request.form.getlist('part_name')
        part_numbers = request.form.getlist('part_number')
        part_qtys = request.form.getlist('part_qty')
        for i, pname in enumerate(part_names):
            if pname.strip():
                pnum = part_numbers[i] if i < len(part_numbers) else ''
                pqty = float(part_qtys[i]) if i < len(part_qtys) and part_qtys[i] else 1
                db.session.add(MuleMaintenancePart(
                    maintenance_id=m.id, part_name=pname.strip(),
                    part_number=pnum.strip(), quantity=pqty
                ))
        # Add components
        comp_types = request.form.getlist('comp_type')
        comp_models = request.form.getlist('comp_model')
        comp_qtys = request.form.getlist('comp_qty')
        comp_knife = request.form.getlist('comp_knife')
        comp_cable_type = request.form.getlist('comp_cable_type')
        comp_cable_len = request.form.getlist('comp_cable_len')
        comp_gasket_len = request.form.getlist('comp_gasket_len')
        comp_spring = request.form.getlist('comp_spring')
        comp_bolt = request.form.getlist('comp_bolt')
        comp_filter = request.form.getlist('comp_filter')
        comp_oil = request.form.getlist('comp_oil')
        comp_volume = request.form.getlist('comp_volume')
        comp_replace = request.form.getlist('comp_replace')
        for i, ctype in enumerate(comp_types):
            if ctype:
                db.session.add(MuleComponent(
                    maintenance_id=m.id,
                    component_type=ctype,
                    model=comp_models[i] if i < len(comp_models) else '',
                    quantity=float(comp_qtys[i]) if i < len(comp_qtys) and comp_qtys[i] else 1,
                    knife_number=comp_knife[i] if i < len(comp_knife) else '',
                    cable_type=comp_cable_type[i] if i < len(comp_cable_type) else '',
                    cable_length=comp_cable_len[i] if i < len(comp_cable_len) else '',
                    gasket_length=comp_gasket_len[i] if i < len(comp_gasket_len) else '',
                    spring_size=comp_spring[i] if i < len(comp_spring) else '',
                    bolt_type=comp_bolt[i] if i < len(comp_bolt) else '',
                    filter_type=comp_filter[i] if i < len(comp_filter) else '',
                    oil_type=comp_oil[i] if i < len(comp_oil) else '',
                    volume=comp_volume[i] if i < len(comp_volume) else '',
                    replacement_date=datetime.strptime(comp_replace[i], '%Y-%m-%d').date() if i < len(comp_replace) and comp_replace[i] else None,
                ))
        db.session.commit()
        log_audit('create', 'mule_maintenance', m.id, f'{m.number} — {m.mule_number}')
        flash(_('Mule maintenance recorded'), 'success')
        return redirect(url_for('mule_list'))
    machines = Machine.query.order_by(Machine.name).all()
    return render_template('mule_form.html', mule=None, machines=machines)

@app.route('/mule/<int:mule_id>/edit', methods=['GET', 'POST'])
@login_required
@role_required('admin', 'technician')
def mule_edit(mule_id):
    m = MuleMaintenance.query.get_or_404(mule_id)
    if request.method == 'POST':
        m.mule_number = request.form['mule_number']
        m.mule_serial = request.form.get('mule_serial', '')
        m.machine_id = int(request.form['machine_id']) if request.form.get('machine_id') else None
        m.date = datetime.strptime(request.form['date'], '%Y-%m-%d').date()
        m.reason = request.form['reason']
        m.next_date = datetime.strptime(request.form['next_date'], '%Y-%m-%d').date() if request.form.get('next_date') else None
        m.periodicity = request.form.get('periodicity', '')
        m.notes = request.form.get('notes', '')
        # Update parts
        MuleMaintenancePart.query.filter_by(maintenance_id=m.id).delete()
        part_names = request.form.getlist('part_name')
        part_numbers = request.form.getlist('part_number')
        part_qtys = request.form.getlist('part_qty')
        for i, pname in enumerate(part_names):
            if pname.strip():
                pnum = part_numbers[i] if i < len(part_numbers) else ''
                pqty = float(part_qtys[i]) if i < len(part_qtys) and part_qtys[i] else 1
                db.session.add(MuleMaintenancePart(
                    maintenance_id=m.id, part_name=pname.strip(),
                    part_number=pnum.strip(), quantity=pqty
                ))
        # Update components
        MuleComponent.query.filter_by(maintenance_id=m.id).delete()
        comp_types = request.form.getlist('comp_type')
        comp_models = request.form.getlist('comp_model')
        comp_qtys = request.form.getlist('comp_qty')
        comp_knife = request.form.getlist('comp_knife')
        comp_cable_type = request.form.getlist('comp_cable_type')
        comp_cable_len = request.form.getlist('comp_cable_len')
        comp_gasket_len = request.form.getlist('comp_gasket_len')
        comp_spring = request.form.getlist('comp_spring')
        comp_bolt = request.form.getlist('comp_bolt')
        comp_filter = request.form.getlist('comp_filter')
        comp_oil = request.form.getlist('comp_oil')
        comp_volume = request.form.getlist('comp_volume')
        comp_replace = request.form.getlist('comp_replace')
        for i, ctype in enumerate(comp_types):
            if ctype:
                db.session.add(MuleComponent(
                    maintenance_id=m.id,
                    component_type=ctype,
                    model=comp_models[i] if i < len(comp_models) else '',
                    quantity=float(comp_qtys[i]) if i < len(comp_qtys) and comp_qtys[i] else 1,
                    knife_number=comp_knife[i] if i < len(comp_knife) else '',
                    cable_type=comp_cable_type[i] if i < len(comp_cable_type) else '',
                    cable_length=comp_cable_len[i] if i < len(comp_cable_len) else '',
                    gasket_length=comp_gasket_len[i] if i < len(comp_gasket_len) else '',
                    spring_size=comp_spring[i] if i < len(comp_spring) else '',
                    bolt_type=comp_bolt[i] if i < len(comp_bolt) else '',
                    filter_type=comp_filter[i] if i < len(comp_filter) else '',
                    oil_type=comp_oil[i] if i < len(comp_oil) else '',
                    volume=comp_volume[i] if i < len(comp_volume) else '',
                    replacement_date=datetime.strptime(comp_replace[i], '%Y-%m-%d').date() if i < len(comp_replace) and comp_replace[i] else None,
                ))
        db.session.commit()
        flash(_('Mule maintenance updated'), 'success')
        return redirect(url_for('mule_list'))
    machines = Machine.query.order_by(Machine.name).all()
    return render_template('mule_form.html', mule=m, machines=machines)

@app.route('/mule/<int:mule_id>/delete', methods=['POST'])
@login_required
@role_required('admin')
def mule_delete(mule_id):
    m = MuleMaintenance.query.get_or_404(mule_id)
    db.session.delete(m)
    db.session.commit()
    flash(_('Mule maintenance deleted'), 'success')
    return redirect(url_for('mule_list'))

@app.route('/mule/order', methods=['POST'])
@login_required
@role_required('admin', 'technician')
def mule_order():
    o = MulePartOrder(
        order_number=gen_mule_order_number(),
        mule_number=request.form.get('mule_number', ''),
        part_name=request.form['part_name'],
        part_number=request.form.get('part_number', ''),
        quantity=float(request.form.get('quantity', 1)),
        unit=request.form.get('unit', 'st'),
        supplier=request.form.get('supplier', ''),
        urgency=request.form.get('urgency', 'normal'),
        notes=request.form.get('notes', ''),
        ordered_by=current_user.id
    )
    db.session.add(o)
    db.session.commit()
    flash(_('Part order created'), 'success')
    return redirect(url_for('mule_list'))

@app.route('/mule/order/<int:order_id>/status', methods=['POST'])
@login_required
@role_required('admin', 'technician')
def mule_order_status(order_id):
    o = MulePartOrder.query.get_or_404(order_id)
    new_status = request.form.get('status')
    if new_status in ('pending', 'ordered', 'delivered', 'cancelled'):
        o.status = new_status
        if new_status == 'delivered':
            o.delivered_at = datetime.utcnow()
    db.session.commit()
    return redirect(url_for('mule_list'))

# ============================================================
# ROUTES — GAS CYLINDERS
# ============================================================

def enforce_cylinder_limit(gas_type):
    """Ensure max 2 cylinders are in_use for a gas type. Mark oldest as empty."""
    in_use = GasCylinder.query.filter_by(
        gas_type=gas_type, status='in_use'
    ).order_by(GasCylinder.installed_at.desc().nullslast(), GasCylinder.id.desc()).all()
    while len(in_use) > 2:
        oldest = in_use.pop()
        oldest.status = 'empty'
        oldest.installed_at = None
        log = CylinderLog(
            cylinder_id=oldest.id, action='auto_empty',
            old_cylinder_number=f'{oldest.cylinder_number} (in_use)',
            new_cylinder_number=f'{oldest.cylinder_number} (empty)',
            performed_by=current_user.id,
            notes='Автоматически: превышен лимит 2 баллона в работе'
        )
        db.session.add(log)

@app.route('/cylinders')
@login_required
@role_required('admin', 'director', 'technician')
def cylinders_dashboard():
    cylinders = GasCylinder.query.order_by(GasCylinder.gas_type, GasCylinder.id).all()
    # Active = only in_use (installed in system). Full = spare in storage.
    n2_active = sorted(
        [c for c in cylinders if c.gas_type == 'nitrogen' and c.status == 'in_use'],
        key=lambda c: (c.installed_at or datetime.min), reverse=True
    )
    co2_active = sorted(
        [c for c in cylinders if c.gas_type == 'co2' and c.status == 'in_use'],
        key=lambda c: (c.installed_at or datetime.min), reverse=True
    )
    # Pad to 2 each
    while len(n2_active) < 2:
        n2_active.append(None)
    while len(co2_active) < 2:
        co2_active.append(None)
    # Stats
    stats = {
        'n2_full': len([c for c in cylinders if c.gas_type == 'nitrogen' and c.status == 'full']),
        'n2_in_use': len([c for c in cylinders if c.gas_type == 'nitrogen' and c.status == 'in_use']),
        'n2_empty': len([c for c in cylinders if c.gas_type == 'nitrogen' and c.status == 'empty']),
        'n2_leak': len([c for c in cylinders if c.gas_type == 'nitrogen' and c.status == 'leak']),
        'n2_faulty': len([c for c in cylinders if c.gas_type == 'nitrogen' and c.status == 'faulty']),
        'co2_full': len([c for c in cylinders if c.gas_type == 'co2' and c.status == 'full']),
        'co2_in_use': len([c for c in cylinders if c.gas_type == 'co2' and c.status == 'in_use']),
        'co2_empty': len([c for c in cylinders if c.gas_type == 'co2' and c.status == 'empty']),
        'co2_leak': len([c for c in cylinders if c.gas_type == 'co2' and c.status == 'leak']),
        'co2_faulty': len([c for c in cylinders if c.gas_type == 'co2' and c.status == 'faulty']),
    }
    orders = CylinderOrder.query.order_by(CylinderOrder.ordered_at.desc()).limit(10).all()
    recent_logs = CylinderLog.query.order_by(CylinderLog.date.desc()).limit(20).all()
    # Available (full) cylinders for replacement dropdown
    available_n2 = GasCylinder.query.filter_by(gas_type='nitrogen', status='full').order_by(GasCylinder.cylinder_number).all()
    available_co2 = GasCylinder.query.filter_by(gas_type='co2', status='full').order_by(GasCylinder.cylinder_number).all()
    # Counts by gas type — usable only (full + in_use)
    count_n2 = stats['n2_full'] + stats['n2_in_use']
    count_co2 = stats['co2_full'] + stats['co2_in_use']
    # Total including all statuses (for reference)
    count_n2_all = len([c for c in cylinders if c.gas_type == 'nitrogen'])
    count_co2_all = len([c for c in cylinders if c.gas_type == 'co2'])
    # Low stock warnings
    n2_total = count_n2
    co2_total = count_co2
    warnings = []
    if n2_total <= 3:
        warnings.append({'gas': 'N₂', 'remaining': n2_total, 'threshold': 3, 'type': 'nitrogen'})
    if co2_total <= 1:
        warnings.append({'gas': 'CO₂', 'remaining': co2_total, 'threshold': 1, 'type': 'co2'})
    # Send notifications if low stock (once per session check)
    for w in warnings:
        existing = Notification.query.filter_by(
            user_id=current_user.id, is_read=False,
            title=_("Low stock") + f": {w['gas']}"
        ).first()
        if not existing:
            create_notification(
                current_user.id,
                _("Low stock") + f": {w['gas']}",
                f"{w['gas']}: {w['remaining']} {_('cylinders remaining')} ({_('threshold')}: {w['threshold']}). {_('Order new cylinders')}!",
                'warning',
                url_for('cylinders_dashboard')
            )
    # Gas system components
    n2_components = GasSystemComponent.query.filter_by(gas_type='nitrogen').order_by(GasSystemComponent.component_type).all()
    co2_components = GasSystemComponent.query.filter_by(gas_type='co2').order_by(GasSystemComponent.component_type).all()
    # Auto-create default components if none exist
    if not n2_components:
        defaults = [('valve', 'Main Valve N₂'), ('heater', 'Heater 1 N₂'), ('heater', 'Heater 2 N₂'), ('manometer', 'Manometer 1 N₂'), ('manometer', 'Manometer 2 N₂'), ('shut_off', 'Switch N₂')]
        for ctype, cname in defaults:
            db.session.add(GasSystemComponent(gas_type='nitrogen', component_type=ctype, name=cname))
        db.session.commit()
        n2_components = GasSystemComponent.query.filter_by(gas_type='nitrogen').order_by(GasSystemComponent.id).all()
    if not co2_components:
        defaults = [('valve', 'Main Valve CO₂'), ('heater', 'Heater 1 CO₂'), ('heater', 'Heater 2 CO₂'), ('manometer', 'Manometer 1 CO₂'), ('manometer', 'Manometer 2 CO₂'), ('shut_off', 'Switch CO₂')]
        for ctype, cname in defaults:
            db.session.add(GasSystemComponent(gas_type='co2', component_type=ctype, name=cname))
        db.session.commit()
        co2_components = GasSystemComponent.query.filter_by(gas_type='co2').order_by(GasSystemComponent.id).all()

    # Build JSON for JavaScript
    def comp_to_dict(c):
        return {'id': c.id, 'type': c.component_type, 'name': c.name, 'status': c.status or 'ok'}
    n2_json = json.dumps([comp_to_dict(c) for c in n2_components])
    co2_json = json.dumps([comp_to_dict(c) for c in co2_components])

    return render_template('cylinders.html',
        n2_active=n2_active[:2], co2_active=co2_active[:2],
        all_cylinders=cylinders, stats=stats,
        orders=orders, recent_logs=recent_logs,
        available_n2=available_n2, available_co2=available_co2,
        count_n2=count_n2, count_co2=count_co2,
        count_n2_all=count_n2_all, count_co2_all=count_co2_all,
        warnings=warnings,
        n2_components=n2_components, co2_components=co2_components,
        n2_json=n2_json, co2_json=co2_json)

@app.route('/cylinders/status', methods=['POST'])
@login_required
@role_required('admin', 'director', 'technician')
def cylinder_status_change():
    data = request.get_json()
    cyl_id = data.get('cylinder_id')
    new_status = data.get('status')
    if new_status not in ('full', 'in_use', 'empty', 'leak', 'faulty'):
        return jsonify({'error': 'Invalid status'}), 400
    cyl = GasCylinder.query.get(cyl_id)
    if not cyl:
        return jsonify({'error': 'Cylinder not found'}), 404
    old_status = cyl.status
    cyl.status = new_status
    if new_status == 'in_use':
        cyl.installed_at = datetime.utcnow()
        enforce_cylinder_limit(cyl.gas_type)
    elif new_status == 'empty':
        cyl.installed_at = None
    log = CylinderLog(
        cylinder_id=cyl.id, action='status_change',
        old_cylinder_number=f'{cyl.cylinder_number} ({old_status})',
        new_cylinder_number=f'{cyl.cylinder_number} ({new_status})',
        performed_by=current_user.id,
        notes=data.get('notes', '')
    )
    db.session.add(log)
    db.session.commit()
    log_audit('status_change', 'cylinder', cyl.id, f'{cyl.cylinder_number}: {old_status} → {new_status}')
    gas_label = 'N₂' if cyl.gas_type == 'nitrogen' else 'CO₂'
    add_work_report(f'🔴 Баллон {gas_label} #{cyl.cylinder_number}: {old_status} → {new_status}')
    return jsonify({'ok': True})

@app.route('/cylinders/api/active')
@login_required
def cylinders_active_api():
    """Return currently installed cylinders for schematic display (left/right). Only in_use."""
    n2_cylinders = GasCylinder.query.filter_by(gas_type='nitrogen', status='in_use').order_by(GasCylinder.installed_at.desc().nullslast(), GasCylinder.id.desc()).limit(2).all()
    co2_cylinders = GasCylinder.query.filter_by(gas_type='co2', status='in_use').order_by(GasCylinder.installed_at.desc().nullslast(), GasCylinder.id.desc()).limit(2).all()
    
    def cyl_info(c):
        return {'id': c.id, 'number': c.cylinder_number, 'status': c.status, 'installed_at': c.installed_at.strftime('%d.%m.%Y') if c.installed_at else ''} if c else None
    
    return jsonify({
        'n2_left': cyl_info(n2_cylinders[0]) if len(n2_cylinders) > 0 else None,
        'n2_right': cyl_info(n2_cylinders[1]) if len(n2_cylinders) > 1 else None,
        'co2_left': cyl_info(co2_cylinders[0]) if len(co2_cylinders) > 0 else None,
        'co2_right': cyl_info(co2_cylinders[1]) if len(co2_cylinders) > 1 else None,
    })

@app.route('/cylinders/component/<int:comp_id>/update', methods=['POST'])
@login_required
@role_required('admin', 'director', 'technician')
def cylinder_component_update(comp_id):
    comp = GasSystemComponent.query.get_or_404(comp_id)
    data = request.get_json()
    old_status = comp.status
    new_status = data.get('status', comp.status)
    comp.status = new_status
    comp.notes = data.get('notes', comp.notes)
    comp.last_check = datetime.utcnow().date()
    # Log the change
    log_entry = CylinderLog(
        cylinder_id=None,
        action=f'component_{new_status}',
        old_cylinder_number=f'{comp.name} ({old_status})',
        new_cylinder_number=f'{comp.name} ({new_status})',
        performed_by=current_user.id,
        notes=data.get('notes', '')
    )
    db.session.add(log_entry)
    db.session.commit()
    return jsonify({'ok': True, 'status': comp.status})

@app.route('/cylinders/replace', methods=['POST'])
@login_required
@role_required('admin', 'director', 'technician')
def cylinder_replace():
    data = request.get_json()
    old_id = data.get('cylinder_id')
    new_cyl_id = data.get('new_cylinder_id')  # ID of existing full cylinder
    new_number = data.get('new_number', '').strip()
    gas_type = data.get('gas_type')
    notes = data.get('notes', '')

    # Determine the new cylinder
    new_cyl = None
    if new_cyl_id:
        new_cyl = GasCylinder.query.get(new_cyl_id)
        if not new_cyl:
            return jsonify({'error': 'Cylinder not found'}), 404
        new_number = new_cyl.cylinder_number
    elif new_number:
        # Manual entry — find or create
        new_cyl = GasCylinder.query.filter_by(cylinder_number=new_number, gas_type=gas_type).first()
        if not new_cyl:
            new_cyl = GasCylinder(
                gas_type=gas_type, cylinder_number=new_number,
                status='full', received_at=datetime.utcnow()
            )
            db.session.add(new_cyl)
            db.session.flush()
    else:
        return jsonify({'error': 'No cylinder selected'}), 400

    # Mark old cylinder as empty
    old_number = None
    if old_id:
        old = GasCylinder.query.get(old_id)
        if old:
            old_number = old.cylinder_number
            old.status = 'empty'
            old.installed_at = None
            log = CylinderLog(
                cylinder_id=old.id, action='replaced',
                old_cylinder_number=old.cylinder_number,
                new_cylinder_number=new_number,
                performed_by=current_user.id, notes=notes
            )
            db.session.add(log)

    # Install new cylinder
    new_cyl.status = 'in_use'
    new_cyl.installed_at = datetime.utcnow()
    if notes:
        new_cyl.notes = notes
    log = CylinderLog(
        cylinder_id=new_cyl.id, action='installed',
        old_cylinder_number=old_number,
        new_cylinder_number=new_number,
        performed_by=current_user.id, notes=notes
    )
    db.session.add(log)
    enforce_cylinder_limit(gas_type)
    db.session.commit()
    gas_label = 'N₂' if gas_type == 'nitrogen' else 'CO₂'
    log_audit('replace', 'cylinder', new_cyl.id, f'{gas_label}: {old_number or "—"} → {new_number}')
    add_work_report(f'🔀 Замена баллона {gas_label}: {old_number or "—"} → {new_number}')
    return jsonify({'ok': True})

@app.route('/cylinders/order', methods=['POST'])
@login_required
@role_required('admin', 'director', 'technician')
def cylinder_order():
    gas_type = request.form.get('gas_type')
    quantity = int(request.form.get('quantity', 1))
    supplier = request.form.get('supplier', '')
    reason = request.form.get('reason', '')
    notes = request.form.get('notes', '')

    order = CylinderOrder(
        gas_type=gas_type, quantity=quantity,
        supplier=supplier, reason=reason,
        ordered_by=current_user.id, notes=notes
    )
    db.session.add(order)
    db.session.commit()
    flash(_('Cylinder order created'), 'success')
    return redirect(url_for('cylinders_dashboard'))

@app.route('/cylinders/order/<int:order_id>/status', methods=['POST'])
@login_required
@role_required('admin', 'technician')
def cylinder_order_status(order_id):
    o = CylinderOrder.query.get_or_404(order_id)
    new_status = request.form.get('status')
    if new_status in ('pending', 'ordered', 'delivered', 'cancelled'):
        o.status = new_status
        if new_status == 'delivered':
            o.delivered_at = datetime.utcnow()
            # Create cylinder records from entered numbers
            numbers_raw = request.form.get('cylinder_numbers', '').strip()
            if numbers_raw:
                numbers = list(set(n.strip() for n in numbers_raw.replace(',', '\n').split('\n') if n.strip()))
            else:
                # Auto-generate numbers if not provided
                numbers = [f"{o.gas_type.upper()}-{datetime.utcnow().strftime('%Y%m%d')}-{i+1:03d}" for i in range(o.quantity)]
            now = datetime.utcnow()
            for num in numbers:
                existing = GasCylinder.query.filter_by(cylinder_number=num, gas_type=o.gas_type).first()
                if existing:
                    existing.status = 'full'
                    existing.received_at = now
                    existing.notes = f"Re-delivered via order #{o.id}"
                else:
                    cyl = GasCylinder(
                        gas_type=o.gas_type, cylinder_number=num,
                        status='full', received_at=now,
                        notes=f"Order #{o.id}"
                    )
                    db.session.add(cyl)
    db.session.commit()
    flash(_('Order status updated'), 'success')
    return redirect(url_for('cylinders_dashboard'))

@app.route('/api/cylinders/status')
@login_required
def api_cylinders_status():
    cylinders = GasCylinder.query.filter(GasCylinder.status.in_(['full', 'in_use'])).all()
    n2 = [{'id': c.id, 'number': c.cylinder_number, 'status': c.status,
            'installed': c.installed_at.strftime('%d-%m-%Y %H:%M') if c.installed_at else None}
           for c in cylinders if c.gas_type == 'nitrogen']
    co2 = [{'id': c.id, 'number': c.cylinder_number, 'status': c.status,
             'installed': c.installed_at.strftime('%d-%m-%Y %H:%M') if c.installed_at else None}
            for c in cylinders if c.gas_type == 'co2']
    return jsonify({'nitrogen': n2, 'co2': co2})

# ============================================================
# ROUTES — EQUIPMENT REPAIRS
# ============================================================

@app.route('/repairs')
@login_required
@role_required('admin', 'director', 'technician')
def repairs_list():
    gas_type = request.args.get('gas', '')
    status = request.args.get('status', '')
    q = EquipmentRepair.query
    if status:
        q = q.filter_by(status=status)
    if gas_type:
        q = q.join(GasSystemComponent).filter(GasSystemComponent.gas_type == gas_type)
    repairs = q.order_by(EquipmentRepair.date_broken.desc()).all()
    components = GasSystemComponent.query.order_by(GasSystemComponent.gas_type, GasSystemComponent.component_type).all()
    return render_template('repairs.html', repairs=repairs, components=components, gas_filter=gas_type, status_filter=status)

@app.route('/repairs/new', methods=['GET', 'POST'])
@login_required
@role_required('admin', 'technician')
def repair_new():
    if request.method == 'POST':
        comp_id = request.form.get('component_id')
        if not comp_id:
            flash(_('Select component'), 'error')
            return redirect(url_for('repair_new'))
        comp_id = int(comp_id)
        comp = GasSystemComponent.query.get(comp_id)
        cost_str = request.form.get('repair_cost', '').strip()
        r = EquipmentRepair(
            component_id=comp_id,
            fault_description=request.form['fault_description'],
            date_broken=datetime.strptime(request.form['date_broken'], '%Y-%m-%dT%H:%M'),
            repair_company=request.form.get('repair_company', ''),
            repair_description=request.form.get('repair_description', ''),
            repair_cost=float(cost_str) if cost_str else 0,
            date_sent=datetime.strptime(request.form['date_sent'], '%Y-%m-%dT%H:%M') if request.form.get('date_sent') else None,
            date_repaired=datetime.strptime(request.form['date_repaired'], '%Y-%m-%dT%H:%M') if request.form.get('date_repaired') else None,
            date_installed=datetime.strptime(request.form['date_installed'], '%Y-%m-%dT%H:%M') if request.form.get('date_installed') else None,
            status=request.form.get('status', 'broken'),
            notes=request.form.get('notes', ''),
            created_by=current_user.id
        )
        # Update component status
        if comp:
            if r.status == 'broken':
                comp.status = 'faulty'
            elif r.status in ('in_repair', 'repaired'):
                comp.status = 'replaced'
            elif r.status == 'installed':
                comp.status = 'ok'
                comp.installed_at = r.date_installed
        db.session.add(r)
        db.session.commit()
        flash(_('Repair record created'), 'success')
        return redirect(url_for('repairs_list'))
    components = GasSystemComponent.query.order_by(GasSystemComponent.gas_type, GasSystemComponent.component_type).all()
    return render_template('repair_form.html', repair=None, components=components)

@app.route('/repairs/<int:repair_id>/edit', methods=['GET', 'POST'])
@login_required
@role_required('admin', 'technician')
def repair_edit(repair_id):
    r = EquipmentRepair.query.get_or_404(repair_id)
    if request.method == 'POST':
        r.fault_description = request.form['fault_description']
        r.date_broken = datetime.strptime(request.form['date_broken'], '%Y-%m-%dT%H:%M')
        r.repair_company = request.form.get('repair_company', '')
        r.repair_description = request.form.get('repair_description', '')
        cost_str = request.form.get('repair_cost', '').strip()
        r.repair_cost = float(cost_str) if cost_str else 0
        r.date_sent = datetime.strptime(request.form['date_sent'], '%Y-%m-%dT%H:%M') if request.form.get('date_sent') else None
        r.date_repaired = datetime.strptime(request.form['date_repaired'], '%Y-%m-%dT%H:%M') if request.form.get('date_repaired') else None
        r.date_installed = datetime.strptime(request.form['date_installed'], '%Y-%m-%dT%H:%M') if request.form.get('date_installed') else None
        r.status = request.form.get('status', r.status)
        r.notes = request.form.get('notes', '')
        # Update component
        comp = r.component
        if comp:
            if r.status == 'installed':
                comp.status = 'ok'
                comp.installed_at = r.date_installed
            elif r.status == 'broken':
                comp.status = 'faulty'
        db.session.commit()
        flash(_('Repair record updated'), 'success')
        return redirect(url_for('repairs_list'))
    components = GasSystemComponent.query.order_by(GasSystemComponent.gas_type, GasSystemComponent.component_type).all()
    return render_template('repair_form.html', repair=r, components=components)

@app.route('/repairs/<int:repair_id>/status', methods=['POST'])
@login_required
@role_required('admin', 'technician')
def repair_status(repair_id):
    r = EquipmentRepair.query.get_or_404(repair_id)
    new_status = request.form.get('status')
    if new_status in ('broken', 'in_repair', 'repaired', 'installed'):
        r.status = new_status
        comp = r.component
        if new_status == 'in_repair':
            r.date_sent = datetime.utcnow()
            if comp: comp.status = 'replaced'
        elif new_status == 'repaired':
            r.date_repaired = datetime.utcnow()
            if comp: comp.status = 'replaced'
        elif new_status == 'installed':
            r.date_installed = datetime.utcnow()
            if comp:
                comp.status = 'ok'
                comp.installed_at = r.date_installed
        db.session.commit()
    flash(_('Status updated'), 'success')
    return redirect(url_for('repairs_list'))

@app.route('/repairs/stats')
@login_required
@role_required('admin', 'director', 'technician')
def repairs_stats():
    all_repairs = EquipmentRepair.query.order_by(EquipmentRepair.date_broken.desc()).all()
    # Stats
    total = len(all_repairs)
    by_component = {}
    by_company = {}
    total_cost = 0
    total_days = 0
    days_count = 0
    for r in all_repairs:
        # By component type
        ctype = r.component.component_type if r.component else 'unknown'
        if ctype not in by_component:
            by_component[ctype] = {'count': 0, 'cost': 0, 'days': []}
        by_component[ctype]['count'] += 1
        by_component[ctype]['cost'] += r.repair_cost or 0
        # Days in repair
        if r.date_broken and r.date_installed:
            days = (r.date_installed - r.date_broken).days
            by_component[ctype]['days'].append(days)
            total_days += days
            days_count += 1
        # By company
        if r.repair_company:
            if r.repair_company not in by_company:
                by_company[r.repair_company] = {'count': 0, 'cost': 0}
            by_company[r.repair_company]['count'] += 1
            by_company[r.repair_company]['cost'] += r.repair_cost or 0
        total_cost += r.repair_cost or 0
    # Average days
    avg_days = round(total_days / days_count, 1) if days_count else 0
    # Add averages to by_component
    for ctype in by_component:
        days_list = by_component[ctype]['days']
        by_component[ctype]['avg_days'] = round(sum(days_list) / len(days_list), 1) if days_list else 0
    return render_template('repairs_stats.html',
        total=total, total_cost=total_cost, avg_days=avg_days,
        by_component=by_component, by_company=by_company,
        recent=all_repairs[:20])

# ============================================================
# ROUTES — TWO (Technical Work Orders)
# ============================================================

def gen_two_number():
    vandaag = datetime.utcnow()
    prefix = vandaag.strftime('%Y%m%d')
    laatste = TechnicalWorkOrder.query.filter(TechnicalWorkOrder.number.like(f'TWO-{prefix}-%')).order_by(TechnicalWorkOrder.id.desc()).first()
    if laatste:
        num = int(laatste.number.split('-')[2]) + 1
    else:
        num = 1
    return f'TWO-{prefix}-{num:04d}'

@app.route('/two')
@login_required
@role_required('admin', 'director', 'technician')
def two_list():
    if current_user.has_role('admin', 'director'):
        orders = TechnicalWorkOrder.query.order_by(TechnicalWorkOrder.created_at.desc()).all()
    elif current_user.has_role('technician'):
        orders = TechnicalWorkOrder.query.order_by(TechnicalWorkOrder.created_at.desc()).all()
    else:
        orders = TechnicalWorkOrder.query.filter_by(created_by=current_user.id).order_by(TechnicalWorkOrder.created_at.desc()).all()
    return render_template('two_list.html', orders=orders)

@app.route('/two/new', methods=['GET', 'POST'])
@login_required
@role_required('admin', 'technician')
def two_new():
    if request.method == 'POST':
        fault_id = int(request.form['fault_id']) if request.form.get('fault_id') else None
        two = TechnicalWorkOrder(
            number=gen_two_number(),
            fault_id=fault_id,
            machine_id=int(request.form['machine_id']) if request.form.get('machine_id') else None,
            section_id=int(request.form['section_id']) if request.form.get('section_id') else None,
            description=request.form['description'],
            additional_work=request.form.get('additional_work', ''),
            planned_date=datetime.strptime(request.form['planned_date'], '%Y-%m-%d').date() if request.form.get('planned_date') else None,
            status=request.form.get('status', 'draft'),
            notes=request.form.get('notes', ''),
            created_by=current_user.id
        )
        # Assign workers
        worker_ids = request.form.getlist('worker_ids')
        for wid in worker_ids:
            w = Monteur.query.get(int(wid))
            if w:
                two.workers.append(w)
        db.session.add(two)
        db.session.flush()

        # Process section assignments
        section_ids = request.form.getlist('section_ids')
        section_descs = request.form.getlist('section_descriptions')
        for i, sid in enumerate(section_ids):
            if not sid:
                continue
            assignment = TWOAssignment(
                two_id=two.id,
                section_id=int(sid),
                description=section_descs[i] if i < len(section_descs) else '',
                sort_order=i
            )
            db.session.add(assignment)
            db.session.flush()
            # Add checklist items for this section
            work_items = request.form.getlist(f'section_work_{i}')
            for j, text in enumerate(work_items):
                if text.strip():
                    db.session.add(TWOChecklistItem(
                        two_id=two.id, assignment_id=assignment.id,
                        text=text.strip(), sort_order=j
                    ))

        # Process machine assignments
        machine_ids = request.form.getlist('machine_ids')
        machine_descs = request.form.getlist('machine_descriptions')
        for i, mid in enumerate(machine_ids):
            if not mid:
                continue
            assignment = TWOAssignment(
                two_id=two.id,
                machine_id=int(mid),
                description=machine_descs[i] if i < len(machine_descs) else '',
                sort_order=i
            )
            db.session.add(assignment)
            db.session.flush()
            # Add checklist items for this machine
            work_items = request.form.getlist(f'machine_work_{i}')
            for j, text in enumerate(work_items):
                if text.strip():
                    db.session.add(TWOChecklistItem(
                        two_id=two.id, assignment_id=assignment.id,
                        text=text.strip(), sort_order=j
                    ))

        db.session.commit()
        # Handle photos
        if 'photos' in request.files:
            for photo in request.files.getlist('photos'):
                if photo.filename:
                    fn = secure_filename(f"two_{two.id}_{photo.filename}")
                    photo.save(os.path.join(app.config['UPLOAD_FOLDER'], fn))
                    db.session.add(TWOPhoto(two_id=two.id, filename=fn))
            db.session.commit()
        log_audit('create', 'two', two.id, two.number)
        flash(_('TWO created') + f': {two.number}', 'success')
        return redirect(url_for('two_detail', two_id=two.id))
    faults = FaultReport.query.filter(FaultReport.status.in_(['open', 'accepted', 'in_progress'])).order_by(FaultReport.created_at.desc()).all()
    workers = Monteur.query.filter_by(actief=True).order_by(Monteur.naam).all()
    machines = Machine.query.order_by(Machine.name).all()
    sections = FactorySection.query.order_by(FactorySection.name).all()
    return render_template('two_form.html', two=None, faults=faults, workers=workers, machines=machines, sections=sections)

@app.route('/api/check-worker-availability')
@login_required
def check_worker_availability():
    """Check if a worker is available on a given date."""
    worker_id = request.args.get('worker_id', type=int)
    date_str = request.args.get('date', '')
    if not worker_id or not date_str:
        return jsonify({'available': True})
    date = datetime.strptime(date_str, '%Y-%m-%d').date()
    
    # Check weekend shifts (off, sick)
    shift = WeekendShift.query.filter_by(user_id=worker_id, date=date).first()
    if shift and shift.shift_type in ('off', 'sick'):
        worker = User.query.get(worker_id)
        # Find alternative workers
        alternatives = suggest_available_workers(date, [worker_id])
        next_dates = suggest_available_dates(worker_id, date)
        return jsonify({
            'available': False,
            'reason': 'off' if shift.shift_type == 'off' else 'sick',
            'worker': worker.display_name if worker else '',
            'date': date.strftime('%d-%m-%Y'),
            'alternatives': alternatives,
            'next_dates': next_dates
        })
    
    # Check Belgian holidays
    holidays = get_belgian_holidays(date.year)
    if date in holidays:
        alternatives = suggest_available_workers(date, [worker_id])
        next_dates = suggest_available_dates(worker_id, date)
        return jsonify({
            'available': False,
            'reason': 'holiday',
            'holiday': holidays[date],
            'date': date.strftime('%d-%m-%Y'),
            'alternatives': alternatives,
            'next_dates': next_dates
        })
    
    # Check work schedule - is this a working day for the worker?
    weekday = date.isoweekday()  # 1=Monday, 7=Sunday
    schedule = WorkSchedule.query.filter_by(user_id=worker_id, is_active=True).first()
    if schedule:
        work_days = [int(d.strip()) for d in schedule.work_days.split(',')]
        if weekday not in work_days:
            worker = User.query.get(worker_id)
            alternatives = suggest_available_workers(date, [worker_id])
            next_dates = suggest_available_dates(worker_id, date)
            day_names = {1:'Mon',2:'Tue',3:'Wed',4:'Thu',5:'Fri',6:'Sat',7:'Sun'}
            work_day_names = ', '.join(day_names.get(d, str(d)) for d in work_days)
            return jsonify({
                'available': False,
                'reason': 'not_work_day',
                'worker': worker.display_name if worker else '',
                'date': date.strftime('%d-%m-%Y'),
                'schedule': f"{schedule.shift_start}-{schedule.shift_end}",
                'work_days': work_day_names,
                'alternatives': alternatives,
                'next_dates': next_dates
            })
    else:
        # No schedule = assume Mon-Fri working week
        if weekday >= 6:  # Saturday or Sunday
            worker = User.query.get(worker_id)
            alternatives = suggest_available_workers(date, [worker_id])
            next_dates = suggest_available_dates(worker_id, date)
            return jsonify({
                'available': False,
                'reason': 'weekend',
                'worker': worker.display_name if worker else '',
                'date': date.strftime('%d-%m-%Y'),
                'alternatives': alternatives,
                'next_dates': next_dates
            })
    
    return jsonify({'available': True})

def suggest_available_workers(date, exclude_ids=None):
    """Find workers available on a given date."""
    if exclude_ids is None:
        exclude_ids = []
    
    # Check holidays
    holidays = get_belgian_holidays(date.year)
    if date in holidays:
        return []
    
    weekday = date.isoweekday()
    available = []
    
    workers = Monteur.query.filter_by(actief=True).all()
    for w in workers:
        if w.user_id in exclude_ids:
            continue
        
        # Check if off/sick
        shift = WeekendShift.query.filter_by(user_id=w.user_id, date=date).first()
        if shift and shift.shift_type in ('off', 'sick'):
            continue
        
        # Check work schedule
        schedule = WorkSchedule.query.filter_by(user_id=w.user_id, is_active=True).first()
        if schedule:
            work_days = [int(d.strip()) for d in schedule.work_days.split(',')]
            if weekday not in work_days:
                continue
        else:
            # No schedule = assume Mon-Fri, skip weekends
            if weekday >= 6:
                continue
        
        available.append({'id': w.user_id, 'name': w.naam, 'specialty': w.specialisatie or ''})
    
    return available[:5]  # Return max 5 suggestions

def suggest_available_dates(worker_id, from_date, count=5):
    """Find next available dates for a worker."""
    dates = []
    current = from_date + timedelta(days=1)
    holidays = get_belgian_holidays(from_date.year)
    
    for _ in range(30):  # Check next 30 days
        if len(dates) >= count:
            break
        
        # Check holiday
        if current in holidays:
            current += timedelta(days=1)
            continue
        
        # Check off/sick
        shift = WeekendShift.query.filter_by(user_id=worker_id, date=current).first()
        if shift and shift.shift_type in ('off', 'sick'):
            current += timedelta(days=1)
            continue
        
        # Check work schedule
        weekday = current.isoweekday()
        schedule = WorkSchedule.query.filter_by(user_id=worker_id, is_active=True).first()
        if schedule:
            work_days = [int(d.strip()) for d in schedule.work_days.split(',')]
            if weekday not in work_days:
                current += timedelta(days=1)
                continue
        else:
            # No schedule = assume Mon-Fri, skip weekends
            if weekday >= 6:
                current += timedelta(days=1)
                continue
        
        dates.append(current.strftime('%Y-%m-%d'))
        current += timedelta(days=1)
    
    return dates

@app.route('/two/<int:two_id>')
@login_required
@role_required('admin', 'director', 'technician')
def two_detail(two_id):
    two = TechnicalWorkOrder.query.get_or_404(two_id)
    return render_template('two_detail.html', two=two)

@app.route('/two/<int:two_id>/checklist/add', methods=['POST'])
@login_required
@role_required('admin', 'technician')
def two_checklist_add(two_id):
    two = TechnicalWorkOrder.query.get_or_404(two_id)
    text = request.form.get('text', '').strip()
    if text:
        max_order = max([i.sort_order for i in two.checklist_items], default=0)
        item = TWOChecklistItem(two_id=two_id, text=text, sort_order=max_order + 1)
        db.session.add(item)
        db.session.commit()
    return redirect(url_for('two_detail', two_id=two_id))

@app.route('/two/checklist/<int:item_id>/toggle', methods=['POST'])
@login_required
@role_required('admin', 'technician')
def two_checklist_toggle(item_id):
    item = TWOChecklistItem.query.get_or_404(item_id)
    item.is_done = not item.is_done
    item.done_at = datetime.utcnow() if item.is_done else None
    item.done_by = current_user.id if item.is_done else None
    db.session.commit()
    return jsonify({'ok': True, 'is_done': item.is_done, 'done_at': item.done_at.strftime('%d.%m.%Y %H:%M') if item.done_at else None})

@app.route('/two/checklist/<int:item_id>/delete', methods=['POST'])
@login_required
@role_required('admin', 'technician')
def two_checklist_delete(item_id):
    item = TWOChecklistItem.query.get_or_404(item_id)
    two_id = item.two_id
    db.session.delete(item)
    db.session.commit()
    return redirect(url_for('two_detail', two_id=two_id))

@app.route('/two/<int:two_id>/signature', methods=['POST'])
@login_required
@role_required('admin', 'technician')
def two_add_signature(two_id):
    two = TechnicalWorkOrder.query.get_or_404(two_id)
    signer_name = request.form.get('signer_name', '').strip()
    signature_data = request.form.get('signature_data', '')
    if signer_name and signature_data:
        sig = TWOSignature(two_id=two_id, signer_name=signer_name, signature_data=signature_data)
        db.session.add(sig)
        two.status = 'completed'
        two.completed_at = datetime.utcnow()
        db.session.commit()
        flash(_('Signature saved'), 'success')
    return redirect(url_for('two_detail', two_id=two_id))

@app.route('/two/<int:two_id>/edit', methods=['GET', 'POST'])
@login_required
@role_required('admin', 'technician')
def two_edit(two_id):
    two = TechnicalWorkOrder.query.get_or_404(two_id)
    if request.method == 'POST':
        two.fault_id = int(request.form['fault_id']) if request.form.get('fault_id') else None
        two.machine_id = int(request.form['machine_id']) if request.form.get('machine_id') else None
        two.section_id = int(request.form['section_id']) if request.form.get('section_id') else None
        two.description = request.form['description']
        two.additional_work = request.form.get('additional_work', '')
        two.planned_date = datetime.strptime(request.form['planned_date'], '%Y-%m-%d').date() if request.form.get('planned_date') else None
        two.status = request.form.get('status', two.status)
        two.result = request.form.get('result', '')
        two.parts_used = request.form.get('parts_used', '')
        two.time_spent_hours = float(request.form.get('time_spent_hours', 0))
        two.notes = request.form.get('notes', '')
        if request.form.get('started_at'):
            two.started_at = datetime.strptime(request.form['started_at'], '%Y-%m-%dT%H:%M')
        if request.form.get('completed_at'):
            two.completed_at = datetime.strptime(request.form['completed_at'], '%Y-%m-%dT%H:%M')
        # Update workers
        two.workers = []
        for wid in request.form.getlist('worker_ids'):
            w = Monteur.query.get(int(wid))
            if w:
                two.workers.append(w)
        # Update assignments — delete old, create new
        for a in two.assignments:
            db.session.delete(a)
        db.session.flush()
        # Process section assignments
        section_ids = request.form.getlist('section_ids')
        section_descs = request.form.getlist('section_descriptions')
        for i, sid in enumerate(section_ids):
            if not sid:
                continue
            assignment = TWOAssignment(
                two_id=two.id,
                section_id=int(sid),
                description=section_descs[i] if i < len(section_descs) else '',
                sort_order=i
            )
            db.session.add(assignment)
            db.session.flush()
            work_items = request.form.getlist(f'section_work_{i}')
            for j, text in enumerate(work_items):
                if text.strip():
                    db.session.add(TWOChecklistItem(
                        two_id=two.id, assignment_id=assignment.id,
                        text=text.strip(), sort_order=j
                    ))
        # Process machine assignments
        machine_ids = request.form.getlist('machine_ids')
        machine_descs = request.form.getlist('machine_descriptions')
        for i, mid in enumerate(machine_ids):
            if not mid:
                continue
            assignment = TWOAssignment(
                two_id=two.id,
                machine_id=int(mid),
                description=machine_descs[i] if i < len(machine_descs) else '',
                sort_order=i
            )
            db.session.add(assignment)
            db.session.flush()
            work_items = request.form.getlist(f'machine_work_{i}')
            for j, text in enumerate(work_items):
                if text.strip():
                    db.session.add(TWOChecklistItem(
                        two_id=two.id, assignment_id=assignment.id,
                        text=text.strip(), sort_order=j
                    ))
        # Handle photos
        if 'photos' in request.files:
            for photo in request.files.getlist('photos'):
                if photo.filename:
                    fn = secure_filename(f"two_{two.id}_{photo.filename}")
                    photo.save(os.path.join(app.config['UPLOAD_FOLDER'], fn))
                    db.session.add(TWOPhoto(two_id=two.id, filename=fn))
        db.session.commit()
        log_audit('update', 'two', two.id, two.number)
        flash(_('TWO updated'), 'success')
        return redirect(url_for('two_detail', two_id=two.id))
    faults = FaultReport.query.filter(FaultReport.status.in_(['open', 'accepted', 'in_progress'])).order_by(FaultReport.created_at.desc()).all()
    workers = Monteur.query.filter_by(actief=True).order_by(Monteur.naam).all()
    machines = Machine.query.order_by(Machine.name).all()
    sections = FactorySection.query.order_by(FactorySection.name).all()
    return render_template('two_form.html', two=two, faults=faults, workers=workers, machines=machines, sections=sections)

@app.route('/two/<int:two_id>/complete', methods=['POST'])
@login_required
@role_required('admin', 'technician')
def two_complete(two_id):
    two = TechnicalWorkOrder.query.get_or_404(two_id)
    two.status = 'completed'
    two.completed_at = datetime.utcnow()
    two.result = request.form.get('result', two.result)
    # Do NOT auto-resolve linked fault — close manually
    db.session.commit()
    log_audit('complete', 'two', two.id, two.number)
    flash(_('TWO completed'), 'success')
    return redirect(url_for('two_detail', two_id=two.id))

@app.route('/two/<int:two_id>/delete', methods=['POST'])
@login_required
@role_required('admin')
def two_delete(two_id):
    two = TechnicalWorkOrder.query.get_or_404(two_id)
    db.session.delete(two)
    db.session.commit()
    flash(_('TWO deleted'), 'success')
    return redirect(url_for('two_list'))

@app.route('/two/<int:two_id>/print')
@login_required
@role_required('admin', 'director', 'technician')
def two_print(two_id):
    two = TechnicalWorkOrder.query.get_or_404(two_id)
    return render_template('two_print.html', two=two)

@app.route('/api/two/from-fault/<int:fault_id>')
@login_required
@role_required('admin', 'director', 'technician')
def two_from_fault(fault_id):
    """Get fault data for pre-filling TWO form"""
    f = FaultReport.query.get_or_404(fault_id)
    return jsonify({
        'id': f.id,
        'title': f.title,
        'description': f.description,
        'machine_id': f.machine_id,
        'machine_name': f.machine.name,
        'section_id': f.machine.section_id,
        'section_name': f.machine.section.name if f.machine.section else '',
        'priority': f.priority,
        'reporter': f.reporter.display_name if f.reporter else ''
    })

# ============================================================
# ROUTES — FAULT REPORTS
# ============================================================

@app.route('/faults')
@login_required
def faults_list():
    if current_user.has_role('admin', 'director'):
        faults = FaultReport.query.order_by(FaultReport.created_at.desc()).all()
    elif current_user.has_role('technician'):
        faults = FaultReport.query.filter(
            (FaultReport.technician_id == current_user.id) | 
            (FaultReport.status == 'open')
        ).order_by(FaultReport.created_at.desc()).all()
    else:
        faults = FaultReport.query.filter_by(reporter_id=current_user.id).order_by(FaultReport.created_at.desc()).all()
    return render_template('faults.html', faults=faults)

@app.route('/faults/new', methods=['GET', 'POST'])
@login_required
def fault_new():
    if request.method == 'POST':
        f = FaultReport(
            title=request.form['title'],
            description=request.form['description'],
            priority=request.form.get('priority', 'normal'),
            machine_id=int(request.form['machine_id']),
            reporter_id=current_user.id
        )
        db.session.add(f)
        db.session.flush()

        # Assign selected technicians
        tech_ids = request.form.getlist('technician_ids')
        for tid in tech_ids:
            tech = User.query.get(int(tid))
            if tech:
                f.assigned_technicians.append(tech)
        # Set primary technician (first selected)
        if tech_ids:
            f.technician_id = int(tech_ids[0])
            f.status = 'accepted'
            f.accepted_at = datetime.utcnow()

        db.session.commit()

        # Handle photo uploads
        if 'photos' in request.files:
            for photo in request.files.getlist('photos'):
                if photo.filename:
                    filename = secure_filename(f"fault_{f.id}_{photo.filename}")
                    photo.save(os.path.join(app.config['UPLOAD_FOLDER'], filename))
                    fp = FaultPhoto(fault_id=f.id, filename=filename)
                    db.session.add(fp)

        # Handle video uploads
        if 'videos' in request.files:
            for video in request.files.getlist('videos'):
                if video.filename:
                    filename = secure_filename(f"fault_{f.id}_video_{video.filename}")
                    video.save(os.path.join(app.config['UPLOAD_FOLDER'], filename))
                    fv = FaultVideo(fault_id=f.id, filename=filename)
                    db.session.add(fv)
            db.session.commit()

        # Notify assigned technicians
        for tech in f.assigned_technicians:
            create_notification(
                tech.id,
                _('Fault assigned to you'),
                f"{_('Machine')}: {f.machine.name} - {f.title} ({_('Priority')}: {f.priority})",
                'fault',
                url_for('fault_detail', fault_id=f.id)
            )
        # If no technicians selected, notify all
        if not f.assigned_technicians:
            for tech in User.query.filter_by(role='technician', is_active_user=True).all():
                create_notification(
                    tech.id,
                    _('New fault report'),
                    f"{_('Machine')}: {f.machine.name} - {f.title}",
                    'fault',
                    url_for('fault_detail', fault_id=f.id)
                )

        log_audit('create', 'fault', f.id, f'{f.title} — {f.machine.name} (приоритет: {f.priority})')
        add_work_report(f'⚠️ Новая поломка: {f.title} — {f.machine.name} (приоритет: {f.priority})')

        flash(_('Fault report created'), 'success')
        return redirect(url_for('faults_list'))

    technicians = User.query.filter_by(role='technician', is_active_user=True).order_by(User.display_name).all()
    machines = Machine.query.order_by(Machine.name).all()
    return render_template('fault_form.html', fault=None, machines=machines, technicians=technicians)

@app.route('/faults/<int:fault_id>')
@login_required
def fault_detail(fault_id):
    f = FaultReport.query.get_or_404(fault_id)
    technicians = User.query.filter_by(role='technician', is_active_user=True).order_by(User.display_name).all()
    contractors = Contractor.query.filter_by(is_active=True).order_by(Contractor.company_name).all()
    return render_template('fault_detail.html', fault=f, technicians=technicians, contractors=contractors)

@app.route('/faults/<int:fault_id>/accept', methods=['POST'])
@login_required
@role_required('technician', 'admin')
def fault_accept(fault_id):
    f = FaultReport.query.get_or_404(fault_id)
    f.status = 'accepted'
    f.technician_id = current_user.id
    f.accepted_at = datetime.utcnow()
    db.session.commit()
    
    log_audit('accept', 'fault', f.id, f'{f.title} — {f.machine.name}')
    
    create_notification(
        f.reporter_id,
        _('Fault accepted'),
        f"{_('Technician')} {current_user.display_name} {_('accepted your fault report')}: {f.title}",
        'info',
        url_for('fault_detail', fault_id=f.id)
    )
    
    flash(_('Fault report accepted'), 'success')
    return redirect(url_for('fault_detail', fault_id=f.id))

@app.route('/faults/<int:fault_id>/assign', methods=['POST'])
@login_required
@role_required('admin', 'director')
def fault_assign(fault_id):
    f = FaultReport.query.get_or_404(fault_id)
    tech_ids = request.form.getlist('technician_ids')
    contractor_id = request.form.get('contractor_id', '')
    if not tech_ids and not contractor_id:
        flash(_('Select at least one technician or contractor'), 'error')
        return redirect(url_for('fault_detail', fault_id=f.id))
    # Clear and reassign technicians
    f.assigned_technicians = []
    names = []
    if tech_ids:
        for tid in tech_ids:
            tech = User.query.get(int(tid))
            if tech and tech.role == 'technician':
                f.assigned_technicians.append(tech)
                names.append(tech.display_name or tech.username)
        f.technician_id = int(tech_ids[0])
    # Assign contractor
    if contractor_id:
        f.contractor_id = int(contractor_id)
        c = Contractor.query.get(int(contractor_id))
        if c:
            names.append(f"🏢 {c.company_name}")
    f.status = 'accepted'
    f.accepted_at = datetime.utcnow()
    db.session.commit()

    for tech in f.assigned_technicians:
        create_notification(
            tech.id,
            _('Fault assigned to you'),
            f"{_('Admin assigned fault to you')}: {f.title} ({_('Machine')}: {f.machine.name})",
            'fault',
            url_for('fault_detail', fault_id=f.id)
        )
    create_notification(
        f.reporter_id,
        _('Fault assigned'),
        f"{_('Your fault assigned to')} {', '.join(names)}: {f.title}",
        'info',
        url_for('fault_detail', fault_id=f.id)
    )
    
    log_audit('assign', 'fault', f.id, f'{f.title} → {", ".join(names)}')
    
    flash(_('Fault assigned to') + ' ' + (tech.display_name or tech.username), 'success')
    return redirect(url_for('fault_detail', fault_id=f.id))

@app.route('/faults/<int:fault_id>/resolve', methods=['POST'])
@login_required
@role_required('technician', 'admin')
def fault_resolve(fault_id):
    f = FaultReport.query.get_or_404(fault_id)
    f.status = 'resolved'
    f.resolved_at = datetime.utcnow()
    db.session.commit()
    
    log_audit('resolve', 'fault', f.id, f'{f.title} — {f.machine.name}')

    create_notification(
        f.reporter_id,
        _('Fault resolved'),
        f"{_('Your fault report has been resolved')}: {f.title}",
        'info',
        url_for('fault_detail', fault_id=f.id)
    )

    flash(_('Fault report resolved'), 'success')
    return redirect(url_for('fault_detail', fault_id=f.id))

@app.route('/faults/<int:fault_id>/status', methods=['POST'])
@login_required
@role_required('technician', 'admin')
def fault_status_change(fault_id):
    f = FaultReport.query.get_or_404(fault_id)
    data = request.get_json()
    new_status = data.get('status')
    reason = data.get('reason', '')
    allowed = ['open', 'accepted', 'in_progress', 'parts_ordered', 'waiting_parts', 'resolved', 'reopened']
    if new_status not in allowed:
        return jsonify({'error': 'Invalid status'}), 400
    old_status = f.status
    f.status = new_status
    if new_status == 'reopened':
        f.resolved_at = None
    # Save history
    history = FaultStatusHistory(
        fault_id=f.id, old_status=old_status, new_status=new_status,
        reason=reason, changed_by=current_user.id
    )
    db.session.add(history)
    db.session.commit()
    log_audit('status_change', 'fault', f.id, f'{old_status} → {new_status}')
    add_work_report(f'🔄 Поломка #{f.id} "{f.title}": статус {old_status} → {new_status}')
    return jsonify({'ok': True, 'old': old_status, 'new': new_status})

@app.route('/faults/<int:fault_id>/close', methods=['POST'])
@login_required
@role_required('technician', 'admin')
def fault_close(fault_id):
    f = FaultReport.query.get_or_404(fault_id)
    data = request.get_json() if request.is_json else request.form
    has_report = data.get('has_report', '')
    has_parts = data.get('has_parts', '')
    close_notes = data.get('close_notes', '')
    # Validate work report exists
    if has_report == 'no':
        return jsonify({'error': _('Work report is required to close this fault')}), 400
    f.status = 'closed'
    f.resolved_at = datetime.utcnow()
    db.session.commit()
    create_notification(
        f.reporter_id,
        _('Fault closed'),
        f"{_('Your fault report has been closed')}: {f.title}. {close_notes}",
        'success',
        url_for('fault_detail', fault_id=f.id)
    )
    log_audit('close', 'fault', f.id, close_notes)
    add_work_report(f'🔒 Поломка #{f.id} "{f.title}" закрыта. {close_notes}')
    flash(_('Fault report closed'), 'success')
    if request.is_json:
        return jsonify({'ok': True})
    return redirect(url_for('fault_detail', fault_id=f.id))

@app.route('/faults/<int:fault_id>/reopen', methods=['POST'])
@login_required
@role_required('technician', 'admin')
def fault_reopen(fault_id):
    f = FaultReport.query.get_or_404(fault_id)
    data = request.get_json() if request.is_json else request.form
    reason = data.get('reason', '')
    reopen_date = data.get('reopen_date', datetime.utcnow().strftime('%Y-%m-%d'))
    old_status = f.status
    f.status = 'reopened'
    f.resolved_at = None
    # Save history
    history = FaultStatusHistory(
        fault_id=f.id, old_status=old_status, new_status='reopened',
        reason=f'{reopen_date}: {reason}', changed_by=current_user.id
    )
    db.session.add(history)
    db.session.commit()
    create_notification(
        f.reporter_id,
        _('Fault reopened'),
        f"{_('Fault reopened')}: {f.title}. {reason}",
        'warning',
        url_for('fault_detail', fault_id=f.id)
    )
    log_audit('reopen', 'fault', f.id, f'{old_status} → reopened: {reason} ({reopen_date})')
    add_work_report(f'🔓 Поломка #{f.id} "{f.title}" переоткрыта. Причина: {reason}')
    if request.is_json:
        return jsonify({'ok': True})
    return redirect(url_for('fault_detail', fault_id=f.id))

# ============================================================
# ROUTES — WORK REPORTS
# ============================================================

@app.route('/faults/<int:fault_id>/work-report', methods=['GET', 'POST'])
@login_required
@role_required('technician', 'admin')
def work_report_new(fault_id):
    f = FaultReport.query.get_or_404(fault_id)
    if request.method == 'POST':
        wr = WorkReport(
            fault_id=f.id,
            technician_id=current_user.id,
            work_description=request.form['work_description'],
            parts_used=request.form.get('parts_used', '[]'),
            time_spent_hours=float(request.form.get('time_spent_hours', 0))
        )
        db.session.add(wr)
        db.session.commit()
        
        # Handle photo uploads
        if 'photos' in request.files:
            for photo in request.files.getlist('photos'):
                if photo.filename:
                    filename = secure_filename(f"work_{wr.id}_{photo.filename}")
                    photo.save(os.path.join(app.config['UPLOAD_FOLDER'], filename))
                    wp = WorkReportPhoto(report_id=wr.id, filename=filename, description=request.form.get('photo_desc', ''))
                    db.session.add(wp)
            db.session.commit()
        
        # Deduct parts from warehouse
        try:
            parts = json.loads(wr.parts_used)
            for part in parts:
                item = VoorraadItem.query.get(part['part_id'])
                if item:
                    item.hoeveelheid -= part['quantity']
                    mutatie = VoorraadMutatie(
                        item_id=item.id,
                        type='uitgaand',
                        hoeveelheid=part['quantity'],
                        opmerking=f"Work report #{wr.id} for fault #{f.id}"
                    )
                    db.session.add(mutatie)
            db.session.commit()
        except (ValueError, KeyError, TypeError) as e:
            db.session.rollback()
        
        f.status = 'resolved'
        f.resolved_at = datetime.utcnow()
        db.session.commit()
        
        log_audit('create', 'work_report', wr.id, f'Отчёт по поломке #{f.id}: {f.title} ({wr.time_spent_hours}ч)')
        add_work_report(f'📝 Отчёт о работе по поломке #{f.id}: {f.title} ({wr.time_spent_hours}ч)')
        
        flash(_('Work report created'), 'success')
        return redirect(url_for('fault_detail', fault_id=f.id))
    
    return render_template('work_report_form.html', fault=f, warehouse_items=VoorraadItem.query.all(), report=None)

@app.route('/faults/<int:fault_id>/work-report/<int:report_id>/edit', methods=['GET', 'POST'])
@login_required
@role_required('technician', 'admin')
def work_report_edit(fault_id, report_id):
    f = FaultReport.query.get_or_404(fault_id)
    wr = WorkReport.query.get_or_404(report_id)
    if request.method == 'POST':
        # Reverse old parts deduction
        try:
            old_parts = json.loads(wr.parts_used) if wr.parts_used else []
            for part in old_parts:
                item = VoorraadItem.query.get(part['part_id'])
                if item:
                    item.hoeveelheid += part['quantity']
                    mutatie = VoorraadMutatie(
                        item_id=item.id,
                        type='inkomend',
                        hoeveelheid=part['quantity'],
                        opmerking=f"Reversed: work report #{wr.id} edit"
                    )
                    db.session.add(mutatie)
            db.session.commit()
        except (ValueError, KeyError, TypeError):
            db.session.rollback()
        
        # Update report
        wr.work_description = request.form['work_description']
        wr.parts_used = request.form.get('parts_used', '[]')
        wr.time_spent_hours = float(request.form.get('time_spent_hours', 0))
        
        # Handle new photos
        if 'photos' in request.files:
            for photo in request.files.getlist('photos'):
                if photo.filename:
                    filename = secure_filename(f"work_{wr.id}_{photo.filename}")
                    photo.save(os.path.join(app.config['UPLOAD_FOLDER'], filename))
                    wp = WorkReportPhoto(report_id=wr.id, filename=filename, description=request.form.get('photo_desc', ''))
                    db.session.add(wp)
        
        # Deduct new parts from warehouse
        try:
            new_parts = json.loads(wr.parts_used)
            for part in new_parts:
                item = VoorraadItem.query.get(part['part_id'])
                if item:
                    item.hoeveelheid -= part['quantity']
                    mutatie = VoorraadMutatie(
                        item_id=item.id,
                        type='uitgaand',
                        hoeveelheid=part['quantity'],
                        opmerking=f"Work report #{wr.id} (edited) for fault #{f.id}"
                    )
                    db.session.add(mutatie)
            db.session.commit()
        except (ValueError, KeyError, TypeError):
            db.session.rollback()
        
        log_audit('update', 'work_report', wr.id, f'Отчёт по поломке #{f.id}: {f.title}')
        
        flash(_('Work report updated'), 'success')
        return redirect(url_for('fault_detail', fault_id=f.id))
    
    return render_template('work_report_form.html', fault=f, warehouse_items=VoorraadItem.query.all(), report=wr)

# ============================================================
# ROUTES — MESSAGES
# ============================================================

@app.route('/messages')
@login_required
def messages_list():
    received = Message.query.filter_by(receiver_id=current_user.id).order_by(Message.created_at.desc()).all()
    sent = Message.query.filter_by(sender_id=current_user.id).order_by(Message.created_at.desc()).all()
    all_messages = []
    if current_user.has_role('admin'):
        all_messages = Message.query.order_by(Message.created_at.desc()).limit(200).all()
    return render_template('messages.html', received=received, sent=sent, all_messages=all_messages)

@app.route('/messages/new', methods=['GET', 'POST'])
@login_required
def message_new():
    if request.method == 'POST':
        receiver_ids = request.form.getlist('receiver_ids')
        verant_ids = request.form.getlist('verant_ids')
        subject = request.form.get('subject', '')
        body = request.form['body']
        fault_id = request.form.get('fault_id') or None
        # Resolve verantwoordelijke to linked users
        skipped = []
        for vid in verant_ids:
            v = Verantwoordelijke.query.get(int(vid))
            if v:
                linked_user = User.query.filter_by(person_id=v.id, is_active_user=True).first()
                if linked_user and str(linked_user.id) not in receiver_ids:
                    receiver_ids.append(str(linked_user.id))
                elif not linked_user:
                    skipped.append(v.naam)
        # If nobody selected — send to ALL active users
        if not receiver_ids:
            all_users = User.query.filter(User.id != current_user.id, User.is_active_user == True).all()
            receiver_ids = [str(u.id) for u in all_users]
        sent = 0
        for rid in receiver_ids:
            m = Message(
                sender_id=current_user.id,
                receiver_id=int(rid),
                subject=subject,
                body=body,
                fault_id=fault_id
            )
            db.session.add(m)
            create_notification(
                int(rid),
                _('New message'),
                f"{_('From')}: {current_user.display_name} - {subject}",
                'message',
                url_for('messages_list')
            )
            sent += 1
        db.session.commit()
        msg = _('Message sent to') + f' {sent} ' + _('users')
        if skipped:
            msg += f'. {_("No user account for")}: {", ".join(skipped)}'
        flash(msg, 'success')
        return redirect(url_for('messages_list'))
    
    users = User.query.filter(User.id != current_user.id, User.is_active_user == True).all()
    verantwoordelijken = Verantwoordelijke.query.order_by(Verantwoordelijke.naam).all()
    fault_id = request.args.get('fault_id')
    return render_template('message_form.html', users=users, fault_id=fault_id, verantwoordelijken=verantwoordelijken)

@app.route('/messages/<int:message_id>')
@login_required
def message_detail(message_id):
    m = Message.query.get_or_404(message_id)
    if m.sender_id != current_user.id and m.receiver_id != current_user.id:
        from flask import abort
        abort(403)
    if m.receiver_id == current_user.id:
        m.is_read = True
        db.session.commit()
    return render_template('message_detail.html', message=m)

# ============================================================
# ROUTES — NOTIFICATIONS
# ============================================================

@app.route('/notifications')
@login_required
def notifications_list():
    notifications = Notification.query.filter_by(user_id=current_user.id).order_by(Notification.created_at.desc()).limit(50).all()
    return render_template('notifications.html', notifications=notifications)

@app.route('/notifications/unread')
@login_required
def notifications_unread():
    count = Notification.query.filter_by(user_id=current_user.id, is_read=False).count()
    return jsonify({'count': count})

@app.route('/notifications/<int:notif_id>/read', methods=['POST'])
@login_required
def notification_read(notif_id):
    n = Notification.query.get_or_404(notif_id)
    if n.user_id == current_user.id:
        n.is_read = True
        db.session.commit()
    return jsonify({'success': True})

@app.route('/notifications/read-all', methods=['POST'])
@login_required
def notifications_read_all():
    Notification.query.filter_by(user_id=current_user.id, is_read=False).update({'is_read': True})
    db.session.commit()
    return jsonify({'success': True})

# ============================================================
# ROUTES — DIRECTOR STATISTICS
# ============================================================

@app.route('/stats/faults')
@login_required
@role_required('admin', 'director', 'technician')
def stats_faults():
    total = FaultReport.query.count()
    open_count = FaultReport.query.filter_by(status='open').count()
    in_progress = FaultReport.query.filter(FaultReport.status.in_(['accepted', 'in_progress'])).count()
    resolved = FaultReport.query.filter_by(status='resolved').count()
    
    by_machine = db.session.query(
        Machine.name, db.func.count(FaultReport.id)
    ).join(FaultReport).group_by(Machine.name).all()
    
    by_priority = db.session.query(
        FaultReport.priority, db.func.count(FaultReport.id)
    ).group_by(FaultReport.priority).all()
    
    return render_template('stats_faults.html',
        total=total, open_count=open_count, in_progress=in_progress, resolved=resolved,
        by_machine=by_machine, by_priority=by_priority)

@app.route('/stats/full')
@login_required
@role_required('admin', 'director')
def stats_full():
    """Full statistics page with all metrics"""
    date_from = request.args.get('date_from', '')
    date_to = request.args.get('date_to', '')
    
    if date_from and date_to:
        d_from = datetime.strptime(date_from, '%Y-%m-%d')
        d_to = datetime.strptime(date_to, '%Y-%m-%d') + timedelta(days=1)
        period = (d_to - d_from).days
    else:
        period = int(request.args.get('period', '30'))
        d_from = datetime.utcnow() - timedelta(days=period)
        d_to = datetime.utcnow() + timedelta(days=1)
        date_from = d_from.strftime('%Y-%m-%d')
        date_to = (d_to - timedelta(days=1)).strftime('%Y-%m-%d')
    
    # Faults stats
    faults_total = FaultReport.query.count()
    faults_period = FaultReport.query.filter(FaultReport.created_at >= d_from).count()
    faults_open = FaultReport.query.filter(FaultReport.status.in_(['open', 'accepted', 'in_progress'])).count()
    faults_resolved = FaultReport.query.filter(FaultReport.status == 'resolved').count()
    faults_critical = FaultReport.query.filter(FaultReport.priority == 'critical').count()
    
    faults_by_priority = db.session.query(
        FaultReport.priority, db.func.count(FaultReport.id)
    ).filter(FaultReport.created_at >= d_from).group_by(FaultReport.priority).all()
    
    faults_by_status = db.session.query(
        FaultReport.status, db.func.count(FaultReport.id)
    ).group_by(FaultReport.status).all()
    
    # Top machines with faults
    top_machines = db.session.query(
        Machine.name, db.func.count(FaultReport.id).label('cnt')
    ).join(FaultReport).filter(FaultReport.created_at >= d_from).group_by(Machine.name).order_by(db.desc('cnt')).limit(10).all()
    
    # Warehouse stats
    warehouse_total = VoorraadItem.query.count()
    warehouse_low = VoorraadItem.query.filter(VoorraadItem.hoeveelheid <= VoorraadItem.minimum).count()
    warehouse_incoming = VoorraadMutatie.query.filter(VoorraadMutatie.type == 'inkomend', VoorraadMutatie.aangemaakt >= d_from).count()
    warehouse_outgoing = VoorraadMutatie.query.filter(VoorraadMutatie.type == 'uitgaand', VoorraadMutatie.aangemaakt >= d_from).count()
    
    # Users stats
    users_total = User.query.filter(User.is_active_user == True).count()
    users_online = UserActivityLog.query.filter(UserActivityLog.action == 'login', UserActivityLog.created_at >= datetime.utcnow() - timedelta(hours=1)).count()
    
    # Cylinder stats
    cyl_n2_full = GasCylinder.query.filter_by(gas_type='nitrogen', status='full').count()
    cyl_n2_in_use = GasCylinder.query.filter_by(gas_type='nitrogen', status='in_use').count()
    cyl_co2_full = GasCylinder.query.filter_by(gas_type='co2', status='full').count()
    cyl_co2_in_use = GasCylinder.query.filter_by(gas_type='co2', status='in_use').count()
    
    # TWO stats
    two_total = TechnicalWorkOrder.query.count()
    two_active = TechnicalWorkOrder.query.filter(TechnicalWorkOrder.status.in_(['draft', 'assigned', 'in_progress'])).count()
    two_completed = TechnicalWorkOrder.query.filter_by(status='completed').count()
    
    # Machine report with filtering
    section_filter = request.args.get('section', '')
    group_filter = request.args.get('group', '')
    
    machines_q = Machine.query.order_by(Machine.name)
    if section_filter:
        machines_q = machines_q.filter_by(section_id=int(section_filter))
    
    all_machines = machines_q.all()
    sections = FactorySection.query.order_by(FactorySection.name).all()
    groups = ResponsibleGroup.query.order_by(ResponsibleGroup.name).all()
    
    machine_report = []
    for m in all_machines:
        total_faults = FaultReport.query.filter(FaultReport.machine_id == m.id).count()
        period_faults = FaultReport.query.filter(FaultReport.machine_id == m.id, FaultReport.created_at >= d_from).count()
        open_faults = FaultReport.query.filter(FaultReport.machine_id == m.id, FaultReport.status.in_(['open', 'accepted', 'in_progress'])).count()
        critical_faults = FaultReport.query.filter(FaultReport.machine_id == m.id, FaultReport.priority == 'critical').count()
        last_fault = FaultReport.query.filter(FaultReport.machine_id == m.id).order_by(FaultReport.created_at.desc()).first()
        machine_report.append({
            'machine': m,
            'total_faults': total_faults,
            'period_faults': period_faults,
            'open_faults': open_faults,
            'critical_faults': critical_faults,
            'last_fault': last_fault,
        })
    
    # Sort by period faults descending
    machine_report.sort(key=lambda x: x['period_faults'], reverse=True)
    
    return render_template('stats_full.html',
        period=period, d_from=d_from, date_from=date_from, date_to=date_to,
        faults_total=faults_total, faults_period=faults_period, faults_open=faults_open,
        faults_resolved=faults_resolved, faults_critical=faults_critical,
        faults_by_priority=faults_by_priority, faults_by_status=faults_by_status,
        top_machines=top_machines,
        warehouse_total=warehouse_total, warehouse_low=warehouse_low,
        warehouse_incoming=warehouse_incoming, warehouse_outgoing=warehouse_outgoing,
        users_total=users_total, users_online=users_online,
        cyl_n2_full=cyl_n2_full, cyl_n2_in_use=cyl_n2_in_use,
        cyl_co2_full=cyl_co2_full, cyl_co2_in_use=cyl_co2_in_use,
        two_total=two_total, two_active=two_active, two_completed=two_completed,
        machine_report=machine_report, sections=sections, groups=groups,
        section_filter=section_filter, group_filter=group_filter)

@app.route('/stats/full/export')
@login_required
@role_required('admin', 'director')
def stats_export():
    """Export full statistics in various formats"""
    format_type = request.args.get('format', 'pdf')
    date_from = request.args.get('date_from', '')
    date_to = request.args.get('date_to', '')
    
    if date_from and date_to:
        d_from = datetime.strptime(date_from, '%Y-%m-%d')
        d_to = datetime.strptime(date_to, '%Y-%m-%d') + timedelta(days=1)
        period = (d_to - d_from).days
    else:
        period = int(request.args.get('period', '30'))
        d_from = datetime.utcnow() - timedelta(days=period)
        d_to = datetime.utcnow() + timedelta(days=1)
        date_from = d_from.strftime('%Y-%m-%d')
        date_to = (d_to - timedelta(days=1)).strftime('%Y-%m-%d')
    
    # Collect all stats
    stats = {
        'period': f'{date_from} — {date_to}',
        'date': datetime.utcnow().strftime('%d.%m.%Y %H:%M'),
        'faults': {
            'total': FaultReport.query.count(),
            'period': FaultReport.query.filter(FaultReport.created_at >= d_from).count(),
            'open': FaultReport.query.filter(FaultReport.status.in_(['open', 'accepted', 'in_progress'])).count(),
            'resolved': FaultReport.query.filter(FaultReport.status == 'resolved').count(),
            'critical': FaultReport.query.filter(FaultReport.priority == 'critical').count(),
        },
        'warehouse': {
            'total': VoorraadItem.query.count(),
            'low_stock': VoorraadItem.query.filter(VoorraadItem.hoeveelheid <= VoorraadItem.minimum).count(),
            'incoming': VoorraadMutatie.query.filter(VoorraadMutatie.type == 'inkomend', VoorraadMutatie.aangemaakt >= d_from).count(),
            'outgoing': VoorraadMutatie.query.filter(VoorraadMutatie.type == 'uitgaand', VoorraadMutatie.aangemaakt >= d_from).count(),
        },
        'users': {
            'total': User.query.filter(User.is_active_user == True).count(),
        },
        'cylinders': {
            'n2_full': GasCylinder.query.filter_by(gas_type='nitrogen', status='full').count(),
            'n2_in_use': GasCylinder.query.filter_by(gas_type='nitrogen', status='in_use').count(),
            'co2_full': GasCylinder.query.filter_by(gas_type='co2', status='full').count(),
            'co2_in_use': GasCylinder.query.filter_by(gas_type='co2', status='in_use').count(),
        },
        'two': {
            'total': TechnicalWorkOrder.query.count(),
            'active': TechnicalWorkOrder.query.filter(TechnicalWorkOrder.status.in_(['draft', 'assigned', 'in_progress'])).count(),
            'completed': TechnicalWorkOrder.query.filter_by(status='completed').count(),
        }
    }
    
    # Machine report
    section_filter = request.args.get('section', '')
    machines_q = Machine.query.order_by(Machine.name)
    if section_filter:
        machines_q = machines_q.filter_by(section_id=int(section_filter))
    
    machine_data = []
    for m in machines_q.all():
        total_f = FaultReport.query.filter(FaultReport.machine_id == m.id).count()
        period_f = FaultReport.query.filter(FaultReport.machine_id == m.id, FaultReport.created_at >= d_from).count()
        open_f = FaultReport.query.filter(FaultReport.machine_id == m.id, FaultReport.status.in_(['open', 'accepted', 'in_progress'])).count()
        critical_f = FaultReport.query.filter(FaultReport.machine_id == m.id, FaultReport.priority == 'critical').count()
        machine_data.append({
            'name': m.name, 'type': m.machine_type or '', 'serial': m.serial_number or '',
            'section': m.section.name if m.section else '',
            'total': total_f, 'period': period_f, 'open': open_f, 'critical': critical_f
        })
    machine_data.sort(key=lambda x: x['period'], reverse=True)
    
    # Top machines
    top_machines = db.session.query(
        Machine.name, db.func.count(FaultReport.id).label('cnt')
    ).join(FaultReport).filter(FaultReport.created_at >= d_from).group_by(Machine.name).order_by(db.desc('cnt')).limit(10).all()
    
    # By priority
    by_priority = db.session.query(
        FaultReport.priority, db.func.count(FaultReport.id)
    ).filter(FaultReport.created_at >= d_from).group_by(FaultReport.priority).all()
    
    filename = f'statistics_{period}days_{datetime.utcnow().strftime("%Y%m%d")}'
    
    # === TXT ===
    if format_type == 'txt':
        lines = []
        lines.append('=' * 60)
        lines.append('СТАТИСТИКА — CRM МАСТЕРСКАЯ')
        lines.append(f'Период: {stats["period"]}')
        lines.append(f'Дата: {stats["date"]}')
        lines.append('=' * 60)
        lines.append('')
        lines.append('--- ЗАЯВКИ О НЕИСПРАВНОСТИ ---')
        lines.append(f'Всего: {stats["faults"]["total"]}')
        lines.append(f'За период: {stats["faults"]["period"]}')
        lines.append(f'Открытых: {stats["faults"]["open"]}')
        lines.append(f'Решённых: {stats["faults"]["resolved"]}')
        lines.append(f'Критичных: {stats["faults"]["critical"]}')
        lines.append('')
        lines.append('По приоритету:')
        for p, c in by_priority:
            lines.append(f'  {p}: {c}')
        lines.append('')
        lines.append('Топ станков по заявкам:')
        for name, cnt in top_machines:
            lines.append(f'  {name}: {cnt}')
        lines.append('')
        lines.append('--- СКЛАД ---')
        lines.append(f'Всего позиций: {stats["warehouse"]["total"]}')
        lines.append(f'Низкий запас: {stats["warehouse"]["low_stock"]}')
        lines.append(f'Приход за период: {stats["warehouse"]["incoming"]}')
        lines.append(f'Расход за период: {stats["warehouse"]["outgoing"]}')
        lines.append('')
        lines.append('--- БАЛЛОНЫ ---')
        lines.append(f'N₂ полных: {stats["cylinders"]["n2_full"]}')
        lines.append(f'N₂ в работе: {stats["cylinders"]["n2_in_use"]}')
        lines.append(f'CO₂ полных: {stats["cylinders"]["co2_full"]}')
        lines.append(f'CO₂ в работе: {stats["cylinders"]["co2_in_use"]}')
        lines.append('')
        lines.append('--- TWO (НАРЯДЫ) ---')
        lines.append(f'Всего: {stats["two"]["total"]}')
        lines.append(f'Активных: {stats["two"]["active"]}')
        lines.append(f'Завершённых: {stats["two"]["completed"]}')
        lines.append('')
        lines.append('--- ПОЛЬЗОВАТЕЛИ ---')
        lines.append(f'Активных: {stats["users"]["total"]}')
        lines.append('')
        lines.append('--- ОТЧЁТ ПО СТАНКАМ ---')
        lines.append(f'Всего станков: {len(machine_data)}')
        lines.append('')
        lines.append(f'{"Станок":<25} {"Тип":<15} {"Отдел":<15} {"Всего":>6} {"Период":>7} {"Откр":>5} {"Крит":>5}')
        lines.append('-' * 80)
        for m in machine_data:
            lines.append(f'{m["name"]:<25} {m["type"]:<15} {m["section"]:<15} {m["total"]:>6} {m["period"]:>7} {m["open"]:>5} {m["critical"]:>5}')
        lines.append('')
        lines.append('=' * 60)
        lines.append('CRM Мастерская — Статистика')
        
        from flask import Response
        return Response('\n'.join(lines), mimetype='text/plain',
            headers={'Content-Disposition': f'attachment;filename={filename}.txt'})
    
    # === CSV ===
    elif format_type == 'csv':
        import csv, io
        output = io.StringIO()
        writer = csv.writer(output)
        writer.writerow(['CRM Мастерская — Статистика', stats['period'], stats['date']])
        writer.writerow([])
        writer.writerow(['ЗАЯВКИ', 'Количество'])
        writer.writerow(['Всего', stats['faults']['total']])
        writer.writerow(['За период', stats['faults']['period']])
        writer.writerow(['Открытых', stats['faults']['open']])
        writer.writerow(['Решённых', stats['faults']['resolved']])
        writer.writerow(['Критичных', stats['faults']['critical']])
        writer.writerow([])
        writer.writerow(['Приоритет', 'Количество'])
        for p, c in by_priority:
            writer.writerow([p, c])
        writer.writerow([])
        writer.writerow(['Станок', 'Заявок'])
        for name, cnt in top_machines:
            writer.writerow([name, cnt])
        writer.writerow([])
        writer.writerow(['СКЛАД', 'Количество'])
        writer.writerow(['Всего позиций', stats['warehouse']['total']])
        writer.writerow(['Низкий запас', stats['warehouse']['low_stock']])
        writer.writerow(['Приход', stats['warehouse']['incoming']])
        writer.writerow(['Расход', stats['warehouse']['outgoing']])
        writer.writerow([])
        writer.writerow(['ОТЧЁТ ПО СТАНКАМ'])
        writer.writerow(['Станок', 'Тип', 'Серийный номер', 'Отдел', 'Всего заявок', 'За период', 'Открытых', 'Критичных'])
        for m in machine_data:
            writer.writerow([m['name'], m['type'], m['serial'], m['section'], m['total'], m['period'], m['open'], m['critical']])
        output.seek(0)
        from flask import Response
        return Response(output.getvalue(), mimetype='text/csv',
            headers={'Content-Disposition': f'attachment;filename={filename}.csv'})
    
    # === EXCEL ===
    elif format_type == 'excel':
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        wb = Workbook()
        ws = wb.active
        ws.title = 'Статистика'
        
        header_fill = PatternFill(start_color='2C3E50', end_color='2C3E50', fill_type='solid')
        header_font = Font(bold=True, color='FFFFFF', size=11)
        title_font = Font(bold=True, size=14)
        section_font = Font(bold=True, size=12, color='2C3E50')
        thin_border = Border(left=Side(style='thin'), right=Side(style='thin'), top=Side(style='thin'), bottom=Side(style='thin'))
        
        ws['A1'] = 'Статистика — CRM Мастерская'
        ws['A1'].font = title_font
        ws['A2'] = f'Период: {stats["period"]}'
        ws['A3'] = f'Дата: {stats["date"]}'
        
        row = 5
        ws.cell(row=row, column=1, value='ЗАЯВКИ О НЕИСПРАВНОСТИ').font = section_font
        row += 1
        for label, val in [('Всего', stats['faults']['total']), ('За период', stats['faults']['period']),
                           ('Открытых', stats['faults']['open']), ('Решённых', stats['faults']['resolved']),
                           ('Критичных', stats['faults']['critical'])]:
            ws.cell(row=row, column=1, value=label).border = thin_border
            ws.cell(row=row, column=2, value=val).border = thin_border
            row += 1
        
        row += 1
        ws.cell(row=row, column=1, value='По приоритету').font = section_font
        row += 1
        for p, c in by_priority:
            ws.cell(row=row, column=1, value=p).border = thin_border
            ws.cell(row=row, column=2, value=c).border = thin_border
            row += 1
        
        row += 1
        ws.cell(row=row, column=1, value='Топ станков').font = section_font
        row += 1
        for name, cnt in top_machines:
            ws.cell(row=row, column=1, value=name).border = thin_border
            ws.cell(row=row, column=2, value=cnt).border = thin_border
            row += 1
        
        row += 1
        ws.cell(row=row, column=1, value='СКЛАД').font = section_font
        row += 1
        for label, val in [('Всего позиций', stats['warehouse']['total']), ('Низкий запас', stats['warehouse']['low_stock']),
                           ('Приход', stats['warehouse']['incoming']), ('Расход', stats['warehouse']['outgoing'])]:
            ws.cell(row=row, column=1, value=label).border = thin_border
            ws.cell(row=row, column=2, value=val).border = thin_border
            row += 1
        
        row += 1
        ws.cell(row=row, column=1, value='БАЛЛОНЫ').font = section_font
        row += 1
        for label, val in [('N₂ полных', stats['cylinders']['n2_full']), ('N₂ в работе', stats['cylinders']['n2_in_use']),
                           ('CO₂ полных', stats['cylinders']['co2_full']), ('CO₂ в работе', stats['cylinders']['co2_in_use'])]:
            ws.cell(row=row, column=1, value=label).border = thin_border
            ws.cell(row=row, column=2, value=val).border = thin_border
            row += 1
        
        row += 1
        ws.cell(row=row, column=1, value='ОТЧЁТ ПО СТАНКАМ').font = section_font
        row += 1
        machine_headers = ['Станок', 'Тип', 'Серийный номер', 'Отдел', 'Всего заявок', 'За период', 'Открытых', 'Критичных']
        for col, h in enumerate(machine_headers, 1):
            cell = ws.cell(row=row, column=col, value=h)
            cell.fill = header_fill
            cell.font = header_font
            cell.border = thin_border
        row += 1
        for m in machine_data:
            ws.cell(row=row, column=1, value=m['name']).border = thin_border
            ws.cell(row=row, column=2, value=m['type']).border = thin_border
            ws.cell(row=row, column=3, value=m['serial']).border = thin_border
            ws.cell(row=row, column=4, value=m['section']).border = thin_border
            ws.cell(row=row, column=5, value=m['total']).border = thin_border
            ws.cell(row=row, column=6, value=m['period']).border = thin_border
            ws.cell(row=row, column=7, value=m['open']).border = thin_border
            ws.cell(row=row, column=8, value=m['critical']).border = thin_border
            row += 1
        
        ws.column_dimensions['A'].width = 25
        ws.column_dimensions['B'].width = 15
        ws.column_dimensions['C'].width = 18
        ws.column_dimensions['D'].width = 15
        
        buf = io.BytesIO()
        wb.save(buf)
        buf.seek(0)
        return send_file(buf, mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            download_name=f'{filename}.xlsx', as_attachment=True)
    
    # === WORD ===
    elif format_type == 'word':
        from docx import Document
        from docx.shared import Pt, RGBColor
        doc = Document()
        doc.add_heading('Статистика — CRM Мастерская', level=1)
        doc.add_paragraph(f'Период: {stats["period"]}')
        doc.add_paragraph(f'Дата: {stats["date"]}')
        
        doc.add_heading('Заявки о неисправности', level=2)
        table = doc.add_table(rows=6, cols=2)
        table.style = 'Light Grid Accent 1'
        table.rows[0].cells[0].text = 'Показатель'
        table.rows[0].cells[1].text = 'Количество'
        for i, (label, val) in enumerate([('Всего', stats['faults']['total']), ('За период', stats['faults']['period']),
                                           ('Открытых', stats['faults']['open']), ('Решённых', stats['faults']['resolved']),
                                           ('Критичных', stats['faults']['critical'])], 1):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[1].text = str(val)
        
        doc.add_heading('Топ станков по заявкам', level=2)
        table2 = doc.add_table(rows=len(top_machines)+1, cols=2)
        table2.style = 'Light Grid Accent 1'
        table2.rows[0].cells[0].text = 'Станок'
        table2.rows[0].cells[1].text = 'Заявок'
        for i, (name, cnt) in enumerate(top_machines, 1):
            table2.rows[i].cells[0].text = name
            table2.rows[i].cells[1].text = str(cnt)
        
        doc.add_heading('Склад', level=2)
        table3 = doc.add_table(rows=5, cols=2)
        table3.style = 'Light Grid Accent 1'
        table3.rows[0].cells[0].text = 'Показатель'
        table3.rows[0].cells[1].text = 'Количество'
        for i, (label, val) in enumerate([('Всего позиций', stats['warehouse']['total']), ('Низкий запас', stats['warehouse']['low_stock']),
                                           ('Приход', stats['warehouse']['incoming']), ('Расход', stats['warehouse']['outgoing'])], 1):
            table3.rows[i].cells[0].text = label
            table3.rows[i].cells[1].text = str(val)
        
        doc.add_heading('Баллоны', level=2)
        table4 = doc.add_table(rows=5, cols=2)
        table4.style = 'Light Grid Accent 1'
        table4.rows[0].cells[0].text = 'Показатель'
        table4.rows[0].cells[1].text = 'Количество'
        for i, (label, val) in enumerate([('N₂ полных', stats['cylinders']['n2_full']), ('N₂ в работе', stats['cylinders']['n2_in_use']),
                                           ('CO₂ полных', stats['cylinders']['co2_full']), ('CO₂ в работе', stats['cylinders']['co2_in_use'])], 1):
            table4.rows[i].cells[0].text = label
            table4.rows[i].cells[1].text = str(val)
        
        doc.add_heading('Отчёт по станкам', level=2)
        table5 = doc.add_table(rows=len(machine_data)+1, cols=8)
        table5.style = 'Light Grid Accent 1'
        for i, h in enumerate(['Станок', 'Тип', 'Серийный номер', 'Отдел', 'Всего', 'Период', 'Открытых', 'Критичных']):
            table5.rows[0].cells[i].text = h
        for i, m in enumerate(machine_data, 1):
            table5.rows[i].cells[0].text = m['name']
            table5.rows[i].cells[1].text = m['type']
            table5.rows[i].cells[2].text = m['serial']
            table5.rows[i].cells[3].text = m['section']
            table5.rows[i].cells[4].text = str(m['total'])
            table5.rows[i].cells[5].text = str(m['period'])
            table5.rows[i].cells[6].text = str(m['open'])
            table5.rows[i].cells[7].text = str(m['critical'])
        
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return send_file(buf, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            download_name=f'{filename}.docx', as_attachment=True)
    
    # === PDF ===
    elif format_type == 'pdf':
        html = f'''<!DOCTYPE html>
<html><head><meta charset="utf-8">
<style>
body {{ font-family: Arial; padding: 30px; font-size: 12px; }}
h1 {{ font-size: 18px; border-bottom: 2px solid #2c3e50; padding-bottom: 8px; }}
h2 {{ font-size: 14px; color: #2c3e50; margin-top: 20px; }}
table {{ width: 100%; border-collapse: collapse; margin: 10px 0; }}
th {{ background: #2c3e50; color: white; padding: 8px; text-align: left; }}
td {{ padding: 6px 8px; border-bottom: 1px solid #eee; }}
tr:nth-child(even) {{ background: #f9f9f9; }}
.stat {{ display: inline-block; padding: 10px 20px; margin: 5px; background: #f0f7ff; border-radius: 8px; text-align: center; }}
.stat .num {{ font-size: 24px; font-weight: bold; color: #2c3e50; }}
.stat .lbl {{ font-size: 10px; color: #888; }}
.footer {{ margin-top: 30px; font-size: 9px; color: #999; text-align: center; border-top: 1px solid #eee; padding-top: 10px; }}
</style></head><body>
<h1>Статистика — CRM Мастерская</h1>
<p>Период: {stats["period"]} | Дата: {stats["date"]}</p>

<h2>Заявки о неисправности</h2>
<div>
<div class="stat"><div class="num">{stats["faults"]["total"]}</div><div class="lbl">Всего</div></div>
<div class="stat"><div class="num">{stats["faults"]["period"]}</div><div class="lbl">За период</div></div>
<div class="stat"><div class="num">{stats["faults"]["open"]}</div><div class="lbl">Открытых</div></div>
<div class="stat"><div class="num">{stats["faults"]["resolved"]}</div><div class="lbl">Решённых</div></div>
<div class="stat"><div class="num">{stats["faults"]["critical"]}</div><div class="lbl">Критичных</div></div>
</div>

<h2>По приоритету</h2>
<table><tr><th>Приоритет</th><th>Количество</th></tr>'''
        for p, c in by_priority:
            html += f'<tr><td>{p}</td><td>{c}</td></tr>'
        html += '</table>'
        
        html += '<h2>Топ станков по заявкам</h2>'
        html += '<table><tr><th>Станок</th><th>Заявок</th></tr>'
        for name, cnt in top_machines:
            html += f'<tr><td>{name}</td><td>{cnt}</td></tr>'
        html += '</table>'
        
        html += f'''
<h2>Склад</h2>
<div>
<div class="stat"><div class="num">{stats["warehouse"]["total"]}</div><div class="lbl">Всего позиций</div></div>
<div class="stat"><div class="num">{stats["warehouse"]["low_stock"]}</div><div class="lbl">Низкий запас</div></div>
<div class="stat"><div class="num">{stats["warehouse"]["incoming"]}</div><div class="lbl">Приход</div></div>
<div class="stat"><div class="num">{stats["warehouse"]["outgoing"]}</div><div class="lbl">Расход</div></div>
</div>

<h2>Баллоны</h2>
<div>
<div class="stat"><div class="num">{stats["cylinders"]["n2_full"]}</div><div class="lbl">N₂ полных</div></div>
<div class="stat"><div class="num">{stats["cylinders"]["n2_in_use"]}</div><div class="lbl">N₂ в работе</div></div>
<div class="stat"><div class="num">{stats["cylinders"]["co2_full"]}</div><div class="lbl">CO₂ полных</div></div>
<div class="stat"><div class="num">{stats["cylinders"]["co2_in_use"]}</div><div class="lbl">CO₂ в работе</div></div>
</div>

<h2>TWO (Наряды)</h2>
<div>
<div class="stat"><div class="num">{stats["two"]["total"]}</div><div class="lbl">Всего</div></div>
<div class="stat"><div class="num">{stats["two"]["active"]}</div><div class="lbl">Активных</div></div>
<div class="stat"><div class="num">{stats["two"]["completed"]}</div><div class="lbl">Завершённых</div></div>
</div>

<h2>Отчёт по станкам</h2>
<table><tr><th>Станок</th><th>Тип</th><th>Отдел</th><th>Всего</th><th>Период</th><th>Открытых</th><th>Критичных</th></tr>'''
        for m in machine_data:
            html += f'<tr><td>{m["name"]}</td><td>{m["type"]}</td><td>{m["section"]}</td><td>{m["total"]}</td><td>{m["period"]}</td><td>{m["open"]}</td><td>{m["critical"]}</td></tr>'
        html += '</table>'
        
        html += f'''
<div class="footer">CRM Мастерская — Статистика — {stats["date"]}</div>
</body></html>'''
        
        from flask import Response
        return Response(html, mimetype='text/html',
            headers={'Content-Disposition': f'attachment;filename={filename}.html'})
    
    return jsonify({'error': 'Unknown format'}), 400

@app.route('/stats/parts')
@login_required
@role_required('admin', 'director')
def stats_parts():
    low_stock = VoorraadItem.query.filter(VoorraadItem.hoeveelheid <= VoorraadItem.minimum).all()
    recent_usage = VoorraadMutatie.query.filter_by(type='uitgaand').order_by(VoorraadMutatie.aangemaakt.desc()).limit(20).all()
    return render_template('stats_parts.html', low_stock=low_stock, recent_usage=recent_usage)

# ============================================================
# ROUTES — DASHBOARD
# ============================================================

@app.route('/')
@login_required
def index():
    # Responsible persons go directly to floor plan
    if hasattr(current_user, '_person') and current_user.role == 'responsible':
        return redirect(url_for('floor_plan'))
    stats = {
        'opdrachten_totaal': Opdracht.query.count(),
        'opdrachten_actief': Opdracht.query.filter(Opdracht.status.notin_(['afgeleverd', 'geannuleerd'])).count(),
        'opdrachten_vandaag': Opdracht.query.filter(Opdracht.aangemaakt >= datetime.utcnow().date()).count(),
        'verantwoordelijken': Verantwoordelijke.query.count(),
        'monteurs': Monteur.query.filter_by(actief=True).count(),
        'voorraad_laag': VoorraadItem.query.filter(VoorraadItem.hoeveelheid <= VoorraadItem.minimum).count(),
        'faults_open': FaultReport.query.filter_by(status='open').count(),
        'faults_active': FaultReport.query.filter(FaultReport.status.in_(['open', 'accepted', 'in_progress'])).count(),
    }
    recent = Opdracht.query.order_by(Opdracht.aangemaakt.desc()).limit(10).all()
    laag = VoorraadItem.query.filter(VoorraadItem.hoeveelheid <= VoorraadItem.minimum).all()
    
    recent_faults = []
    if current_user.has_role('admin', 'director', 'technician'):
        recent_faults = FaultReport.query.order_by(FaultReport.created_at.desc()).limit(5).all()
    else:
        recent_faults = FaultReport.query.filter_by(reporter_id=current_user.id).order_by(FaultReport.created_at.desc()).limit(5).all()
    
    users = User.query.order_by(User.id).all() if current_user.has_role('admin') else []
    
    # Dashboard statistics for admin/director
    dashboard_stats = {}
    if current_user.has_role('admin', 'director'):
        dashboard_stats['low'] = FaultReport.query.filter_by(priority='low').filter(FaultReport.status.in_(['open', 'accepted', 'in_progress', 'parts_ordered', 'reopened'])).count()
        dashboard_stats['normal'] = FaultReport.query.filter_by(priority='normal').filter(FaultReport.status.in_(['open', 'accepted', 'in_progress', 'parts_ordered', 'reopened'])).count()
        dashboard_stats['high'] = FaultReport.query.filter_by(priority='high').filter(FaultReport.status.in_(['open', 'accepted', 'in_progress', 'parts_ordered', 'reopened'])).count()
        dashboard_stats['critical'] = FaultReport.query.filter_by(priority='critical').filter(FaultReport.status.in_(['open', 'accepted', 'in_progress', 'parts_ordered', 'reopened'])).count()
        
        # Top machines with faults
        all_machines = Machine.query.all()
        top_machines = []
        for m in all_machines:
            fc = FaultReport.query.filter_by(machine_id=m.id).count()
            if fc > 0:
                m.fault_count = fc
                m.open_count = FaultReport.query.filter(FaultReport.machine_id == m.id, FaultReport.status.in_(['open', 'accepted', 'in_progress'])).count()
                m.critical_count = FaultReport.query.filter(FaultReport.machine_id == m.id, FaultReport.priority == 'critical').count()
                top_machines.append(m)
        top_machines.sort(key=lambda x: x.fault_count, reverse=True)
        dashboard_stats['top_machines'] = top_machines
    
    return render_template('index.html', stats=stats, recent_orders=recent, low_stock=laag, recent_faults=recent_faults, users=users, now=datetime.utcnow(), dashboard_stats=dashboard_stats)

# ============================================================
# ROUTES — OPDRACHTEN (existing)
# ============================================================

@app.route('/orders')
@login_required
@role_required('admin', 'director', 'technician')
def orders_list():
    sf = request.args.get('status', '')
    wf = request.args.get('worker', '')
    q = Opdracht.query
    if sf: q = q.filter_by(status=sf)
    if wf: q = q.filter_by(monteur_id=wf)
    orders = q.order_by(Opdracht.aangemaakt.desc()).all()
    workers = Monteur.query.filter_by(actief=True).all()
    return render_template('orders.html', orders=orders, workers=workers, status_filter=sf, worker_filter=wf)

@app.route('/orders/new', methods=['GET', 'POST'])
@login_required
@role_required('admin', 'director', 'technician')
def order_new():
    if request.method == 'POST':
        o = Opdracht(
            nummer=genereer_nummer(),
            responsible_id=request.form['klant_id'],
            monteur_id=request.form.get('monteur_id') or None,
            apparaat=request.form['apparaat'],
            model=request.form.get('model', ''),
            serienummer=request.form.get('serienummer', ''),
            probleem=request.form['probleem'],
            arbeidskosten=float(request.form.get('arbeidskosten', 0)),
            status='aangenomen'
        )
        db.session.add(o)
        db.session.commit()
        flash(_('Work Order created') + f' {o.nummer}', 'success')
        return redirect(url_for('order_detail', order_id=o.id))
    verantwoordelijken = Verantwoordelijke.query.order_by(Verantwoordelijke.naam).all()
    monteurs = Monteur.query.filter_by(actief=True).all()
    return render_template('order_form.html', verantwoordelijken=verantwoordelijken, workers=monteurs, order=None)

@app.route('/orders/<int:order_id>')
@login_required
@role_required('admin', 'director', 'technician')
def order_detail(order_id):
    order = Opdracht.query.get_or_404(order_id)
    return render_template('order_detail.html', order=order)

@app.route('/orders/<int:order_id>/edit', methods=['GET', 'POST'])
@login_required
@role_required('admin', 'director', 'technician')
def order_edit(order_id):
    order = Opdracht.query.get_or_404(order_id)
    if request.method == 'POST':
        order.monteur_id = request.form.get('monteur_id') or None
        order.apparaat = request.form['apparaat']
        order.model = request.form.get('model', '')
        order.serienummer = request.form.get('serienummer', '')
        order.probleem = request.form['probleem']
        order.diagnose = request.form.get('diagnose', '')
        order.uitgevoerd = request.form.get('uitgevoerd', '')
        order.arbeidskosten = float(request.form.get('arbeidskosten', 0))
        order.onderdelenkosten = float(request.form.get('onderdelenkosten', 0))
        order.totaal = order.arbeidskosten + order.onderdelenkosten
        ns = request.form.get('status', order.status)
        if ns != order.status:
            if ns == 'in behandeling' and not order.gestart: order.gestart = datetime.utcnow()
            elif ns == 'gereed' and not order.gereed: order.gereed = datetime.utcnow()
            elif ns == 'afgeleverd' and not order.afgeleverd: order.afgeleverd = datetime.utcnow()
            order.status = ns
        db.session.commit()
        flash(_('Work Order updated'), 'success')
        return redirect(url_for('order_detail', order_id=order.id))
    verantwoordelijken = Verantwoordelijke.query.order_by(Verantwoordelijke.naam).all()
    monteurs = Monteur.query.filter_by(actief=True).all()
    return render_template('order_form.html', verantwoordelijken=verantwoordelijken, workers=monteurs, order=order)

# ============================================================
# ROUTES — KLANTEN (existing)
# ============================================================

@app.route('/responsible')
@login_required
@role_required('admin', 'director')
def responsible_list():
    group_id = request.args.get('group', '')
    q = Verantwoordelijke.query
    if group_id:
        q = q.filter_by(group_id=int(group_id))
    verantwoordelijken = q.order_by(Verantwoordelijke.naam).all()
    groups = ResponsibleGroup.query.order_by(ResponsibleGroup.name).all()
    all_sections = FactorySection.query.order_by(FactorySection.name).all()
    all_machines = Machine.query.order_by(Machine.name).all()
    return render_template('responsible.html', verantwoordelijken=verantwoordelijken,
        groups=groups, group_filter=int(group_id) if group_id else None,
        all_sections=all_sections, all_machines=all_machines)

@app.route('/responsible/phonebook')
@login_required
def phone_directory():
    persons = Verantwoordelijke.query.filter(
        db.or_(Verantwoordelijke.telefoon != '', Verantwoordelijke.internal_phone != '', Verantwoordelijke.email != '')
    ).filter(Verantwoordelijke.is_active == True).order_by(Verantwoordelijke.naam).all()
    workers = Monteur.query.filter(Monteur.actief == True).order_by(Monteur.naam).all()
    users = User.query.filter(User.is_active_user == True).order_by(User.display_name).all()
    groups = ResponsibleGroup.query.order_by(ResponsibleGroup.name).all()
    return render_template('phone_directory.html', persons=persons, workers=workers, users=users, groups=groups)

@app.route('/responsible/groups')
@login_required
@role_required('admin', 'director')
def responsible_groups():
    groups = ResponsibleGroup.query.order_by(ResponsibleGroup.name).all()
    return render_template('responsible_groups.html', groups=groups)

@app.route('/responsible/groups/new', methods=['GET', 'POST'])
@login_required
@role_required('admin')
def responsible_group_new():
    if request.method == 'POST':
        g = ResponsibleGroup(name=request.form['name'], description=request.form.get('description', ''))
        db.session.add(g); db.session.commit()
        flash(_('Group created'), 'success')
        return redirect(url_for('responsible_groups'))
    return render_template('responsible_group_form.html', group=None)

@app.route('/responsible/groups/<int:group_id>')
@login_required
@role_required('admin', 'director')
def responsible_group_detail(group_id):
    g = ResponsibleGroup.query.get_or_404(group_id)
    return render_template('responsible_group_detail.html', group=g)

@app.route('/responsible/groups/<int:group_id>/edit', methods=['GET', 'POST'])
@login_required
@role_required('admin')
def responsible_group_edit(group_id):
    g = ResponsibleGroup.query.get_or_404(group_id)
    if request.method == 'POST':
        g.name = request.form['name']
        g.description = request.form.get('description', '')
        db.session.commit()
        flash(_('Group updated'), 'success')
        return redirect(url_for('responsible_groups'))
    return render_template('responsible_group_form.html', group=g)

@app.route('/responsible/groups/<int:group_id>/delete', methods=['POST'])
@login_required
@role_required('admin')
def responsible_group_delete(group_id):
    g = ResponsibleGroup.query.get_or_404(group_id)
    for m in g.members:
        m.group_id = None
    db.session.delete(g); db.session.commit()
    flash(_('Group deleted'), 'success')
    return redirect(url_for('responsible_groups'))

@app.route('/responsible/groups/<int:group_id>/permissions', methods=['GET', 'POST'])
@login_required
@role_required('admin')
def group_permissions(group_id):
    group = ResponsibleGroup.query.get_or_404(group_id)
    
    if request.method == 'POST':
        # Clear existing permissions
        GroupPermission.query.filter_by(group_id=group.id).delete()
        
        # Add new permissions from form
        for key, name, icon in SECTIONS_LIST:
            can_view = f'{key}_view' in request.form
            can_create = f'{key}_create' in request.form
            can_edit = f'{key}_edit' in request.form
            can_delete = f'{key}_delete' in request.form
            
            if can_view or can_create or can_edit or can_delete:
                perm = GroupPermission(
                    group_id=group.id,
                    section_key=key,
                    can_view=can_view,
                    can_create=can_create,
                    can_edit=can_edit,
                    can_delete=can_delete
                )
                db.session.add(perm)
        
        db.session.commit()
        flash(_('Permissions updated'), 'success')
        return redirect(url_for('responsible_groups'))
    
    # Get current permissions
    perms = {p.section_key: p for p in group.permissions}
    
    return render_template('group_permissions.html', group=group, sections=SECTIONS_LIST, perms=perms)

# Sections list for permissions UI
SECTIONS_LIST = [
    ('machines', 'Machines', '⚙️'),
    ('floor', 'Floor Plan', '🏭'),
    ('sections', 'Sections', '🏗️'),
    ('faults', 'Faults', '⚠️'),
    ('two', 'TWO', '🔧'),
    ('maintenance', 'Maintenance', '📅'),
    ('warehouse', 'Warehouse', '📦'),
    ('orders', 'Work Orders', '📋'),
    ('clients', 'Clients', '👤'),
    ('workers', 'Workers', '🔧'),
    ('contractors', 'Contractors', '🏢'),
    ('invoices', 'Invoices', '📄'),
    ('messages', 'Messages', '💬'),
    ('notifications', 'Notifications', '🔔'),
    ('purchase_requests', 'Purchase Requests', '🛒'),
    ('reports', 'Reports', '📊'),
    ('schedule', 'Schedule', '📅'),
    ('time_tracking', 'Time Tracking', '⏱'),
    ('vacations', 'Vacations', '🏖'),
    ('cylinders', 'Gas Cylinders', '🔴'),
    ('quality', 'Quality Control', '✅'),
    ('users', 'Users', '👥'),
    ('audit_log', 'Audit Log', '📋'),
]

@app.route('/responsible/new', methods=['GET', 'POST'])
@login_required
@role_required('admin', 'director')
def responsible_new():
    if request.method == 'POST':
        first_name = request.form.get('first_name', '').strip()
        last_name = request.form.get('last_name', '').strip()
        full_name = f"{first_name} {last_name}".strip()

        c = Verantwoordelijke(
            naam=full_name,
            telefoon=request.form.get('telefoon', ''),
            internal_phone=request.form.get('internal_phone', ''),
            email=request.form.get('email', ''),
            group_id=int(request.form['group_id']) if request.form.get('group_id') else None,
            access_level=request.form.get('access_level', 'floor'),
            notities=request.form.get('notities', '')
        )
        # Set password if provided
        password = request.form.get('password', '').strip()
        if password:
            c.set_password(password)
        db.session.add(c)
        db.session.flush()
        # Assign sections
        section_ids = request.form.getlist('sections')
        c.resp_sections = []
        for sid in section_ids:
            s = FactorySection.query.get(int(sid))
            if s:
                c.resp_sections.append(s)
        db.session.commit()
        flash(_('Responsible person added') + f': {c.naam}', 'success')
        return redirect(url_for('responsible_list'))
    groups = ResponsibleGroup.query.order_by(ResponsibleGroup.name).all()
    sections = FactorySection.query.order_by(FactorySection.name).all()
    return render_template('responsible_form.html', verantwoordelijke=None, groups=groups, sections=sections)

@app.route('/responsible/<int:resp_id>')
@login_required
@role_required('admin', 'director')
def responsible_detail(resp_id):
    c = Verantwoordelijke.query.get_or_404(resp_id)
    # Machines are now linked via Contractor, not directly to Verantwoordelijke
    machines = []
    # Find linked users
    users = User.query.filter_by(person_id=c.id).all()
    # Find linked sections
    sections = FactorySection.query.filter(
        FactorySection.responsible_persons.any(Verantwoordelijke.id == c.id)
    ).all()
    # Find linked orders
    orders = Opdracht.query.filter_by(responsible_id=c.id).order_by(Opdracht.aangemaakt.desc()).limit(20).all()
    # Find linked workers (via group)
    workers = Monteur.query.filter_by(group_id=c.group_id).all() if c.group_id else []
    today = datetime.utcnow().date()
    return render_template('responsible_detail.html',
        person=c, machines=machines, users=users, sections=sections, orders=orders, workers=workers, today=today)

@app.route('/responsible/<int:resp_id>/assign-group', methods=['POST'])
@login_required
@role_required('admin', 'director')
def responsible_assign_group(resp_id):
    c = Verantwoordelijke.query.get_or_404(resp_id)
    data = request.get_json()
    c.group_id = data.get('group_id')
    db.session.commit()
    return jsonify({'ok': True})

@app.route('/responsible/<int:resp_id>/assign-sections', methods=['POST'])
@login_required
@role_required('admin', 'director')
def responsible_assign_sections(resp_id):
    c = Verantwoordelijke.query.get_or_404(resp_id)
    data = request.get_json()
    section_ids = data.get('section_ids', [])
    c.resp_sections = []
    for sid in section_ids:
        s = FactorySection.query.get(int(sid))
        if s:
            c.resp_sections.append(s)
    db.session.commit()
    return jsonify({'ok': True})

@app.route('/responsible/<int:resp_id>/quick-edit', methods=['POST'])
@login_required
@role_required('admin', 'director')
def responsible_quick_edit(resp_id):
    c = Verantwoordelijke.query.get_or_404(resp_id)
    c.naam = request.form.get('naam', c.naam).strip()
    c.position = request.form.get('position', '').strip()
    c.telefoon = request.form.get('telefoon', '').strip()
    c.internal_phone = request.form.get('internal_phone', '').strip()
    c.email = request.form.get('email', '').strip()
    db.session.commit()
    flash(_('Responsible person updated'), 'success')
    return redirect(url_for('responsible_list'))

@app.route('/responsible/<int:resp_id>/delete', methods=['POST'])
@login_required
@role_required('admin')
def responsible_delete(resp_id):
    c = Verantwoordelijke.query.get_or_404(resp_id)
    c.resp_sections = []
    db.session.delete(c)
    db.session.commit()
    flash(_('Responsible person deleted'), 'success')
    return redirect(url_for('responsible_list'))

@app.route('/responsible/quick-add', methods=['POST'])
@login_required
@role_required('admin', 'director')
def responsible_quick_add():
    naam = request.form.get('naam', '').strip()
    if naam:
        c = Verantwoordelijke(
            naam=naam,
            position=request.form.get('position', '').strip(),
            telefoon=request.form.get('telefoon', '').strip(),
            email=request.form.get('email', '').strip()
        )
        db.session.add(c)
        db.session.commit()
        flash(_('Responsible person added'), 'success')
    return redirect(url_for('responsible_list'))

@app.route('/sections/<int:section_id>/assign-machine', methods=['POST'])
@login_required
@role_required('admin', 'director')
def section_assign_machine(section_id):
    section = FactorySection.query.get_or_404(section_id)
    data = request.get_json()
    machine_ids = data.get('machine_ids', [])
    if not machine_ids:
        single = data.get('machine_id')
        if single:
            machine_ids = [int(single)]
    if not machine_ids:
        return jsonify({'error': 'No machines selected'}), 400
    assigned = 0
    for mid in machine_ids:
        machine = Machine.query.get(int(mid))
        if machine:
            machine.section_id = section_id
            assigned += 1
    db.session.commit()
    return jsonify({'ok': True, 'assigned': assigned})

@app.route('/sections/<int:section_id>/remove-machine/<int:machine_id>', methods=['POST'])
@login_required
@role_required('admin', 'director')
def section_remove_machine(section_id, machine_id):
    machine = Machine.query.get_or_404(machine_id)
    if machine.section_id == section_id:
        machine.section_id = None
        db.session.commit()
    return jsonify({'ok': True})

@app.route('/responsible/<int:resp_id>/edit', methods=['GET', 'POST'])
@login_required
@role_required('admin', 'director')
def responsible_edit(resp_id):
    c = Verantwoordelijke.query.get_or_404(resp_id)
    if request.method == 'POST':
        first_name = request.form.get('first_name', '').strip()
        last_name = request.form.get('last_name', '').strip()
        c.naam = f"{first_name} {last_name}".strip()
        c.telefoon = request.form.get('telefoon', '')
        c.internal_phone = request.form.get('internal_phone', '')
        c.email = request.form.get('email', '')
        c.group_id = int(request.form['group_id']) if request.form.get('group_id') else None
        c.access_level = request.form.get('access_level', c.access_level or 'floor')
        c.is_active = 'is_active' in request.form
        c.notities = request.form.get('notities', '')
        # Update password if provided
        password = request.form.get('password', '').strip()
        if password:
            c.set_password(password)
        # Update sections
        section_ids = request.form.getlist('sections')
        c.resp_sections = []
        for sid in section_ids:
            s = FactorySection.query.get(int(sid))
            if s:
                c.resp_sections.append(s)
        db.session.commit()
        flash(_('Responsible person updated'), 'success')
        return redirect(url_for('responsible_list'))
    groups = ResponsibleGroup.query.order_by(ResponsibleGroup.name).all()
    sections = FactorySection.query.order_by(FactorySection.name).all()
    return render_template('responsible_form.html', verantwoordelijke=c, groups=groups, sections=sections)

# ============================================================
# ROUTES — MONTEURS (existing)
# ============================================================

@app.route('/workers')
@login_required
@role_required('admin', 'director')
def workers_list():
    return render_template('workers.html', workers=Monteur.query.order_by(Monteur.naam).all())

@app.route('/workers/new', methods=['GET', 'POST'])
@login_required
@role_required('admin', 'director')
def worker_new():
    if request.method == 'POST':
        w = Monteur(naam=request.form['naam'], telefoon=request.form.get('telefoon',''),
                    specialisatie=request.form.get('specialisatie',''),
                    tarief_per_uur=float(request.form.get('tarief_per_uur',0)),
                    user_id=int(request.form['user_id']) if request.form.get('user_id') else None,
                    group_id=int(request.form['group_id']) if request.form.get('group_id') else None)
        db.session.add(w); db.session.commit()
        flash(_('Worker added') + f': {w.naam}', 'success')
        return redirect(url_for('workers_list'))
    users = User.query.filter(User.is_active_user == True, User.role.in_(['technician', 'user'])).order_by(User.display_name).all()
    groups = ResponsibleGroup.query.order_by(ResponsibleGroup.name).all()
    return render_template('worker_form.html', worker=None, users=users, groups=groups)

@app.route('/workers/<int:worker_id>/edit', methods=['GET', 'POST'])
@login_required
@role_required('admin', 'director')
def worker_edit(worker_id):
    w = Monteur.query.get_or_404(worker_id)
    if request.method == 'POST':
        w.naam = request.form['naam']; w.telefoon = request.form.get('telefoon','')
        w.specialisatie = request.form.get('specialisatie','')
        w.tarief_per_uur = float(request.form.get('tarief_per_uur',0))
        w.actief = 'actief' in request.form
        w.user_id = int(request.form['user_id']) if request.form.get('user_id') else None
        w.group_id = int(request.form['group_id']) if request.form.get('group_id') else None
        db.session.commit()
        flash(_('Worker updated'), 'success')
        return redirect(url_for('workers_list'))
    users = User.query.filter(User.is_active_user == True, User.role.in_(['technician', 'user'])).order_by(User.display_name).all()
    groups = ResponsibleGroup.query.order_by(ResponsibleGroup.name).all()
    return render_template('worker_form.html', worker=w, users=users, groups=groups)

# ============================================================
# ROUTES — WAREHOUSE GROUPS
# ============================================================

@app.route('/warehouse/groups')
@login_required
@role_required('admin', 'director', 'technician')
def warehouse_groups():
    groups = WarehouseGroup.query.order_by(WarehouseGroup.name).all()
    # Auto-create groups from machine manufacturers
    manufacturers = [m[0] for m in db.session.query(Machine.manufacturer).distinct().all() if m[0]]
    return render_template('warehouse_groups.html', groups=groups, manufacturers=manufacturers)

@app.route('/warehouse/groups/new', methods=['GET', 'POST'])
@login_required
@role_required('admin', 'technician')
def warehouse_group_new():
    if request.method == 'POST':
        g = WarehouseGroup(
            name=request.form['name'],
            manufacturer=request.form.get('manufacturer', ''),
            description=request.form.get('description', '')
        )
        db.session.add(g); db.session.commit()
        flash(_('Group created'), 'success')
        return redirect(url_for('warehouse_groups'))
    manufacturers = [m[0] for m in db.session.query(Machine.manufacturer).distinct().all() if m[0]]
    return render_template('warehouse_group_form.html', group=None, manufacturers=manufacturers)

@app.route('/warehouse/groups/<int:group_id>/edit', methods=['GET', 'POST'])
@login_required
@role_required('admin', 'technician')
def warehouse_group_edit(group_id):
    g = WarehouseGroup.query.get_or_404(group_id)
    if request.method == 'POST':
        g.name = request.form['name']
        g.manufacturer = request.form.get('manufacturer', '')
        g.description = request.form.get('description', '')
        db.session.commit()
        flash(_('Group updated'), 'success')
        return redirect(url_for('warehouse_groups'))
    manufacturers = [m[0] for m in db.session.query(Machine.manufacturer).distinct().all() if m[0]]
    return render_template('warehouse_group_form.html', group=g, manufacturers=manufacturers)

@app.route('/warehouse/groups/<int:group_id>/delete', methods=['POST'])
@login_required
@role_required('admin')
def warehouse_group_delete(group_id):
    g = WarehouseGroup.query.get_or_404(group_id)
    for item in g.items:
        item.group_id = None
    db.session.delete(g); db.session.commit()
    flash(_('Group deleted'), 'success')
    return redirect(url_for('warehouse_groups'))

@app.route('/warehouse/groups/auto-create', methods=['POST'])
@login_required
@role_required('admin')
def warehouse_groups_auto():
    manufacturers = [m[0] for m in db.session.query(Machine.manufacturer).distinct().all() if m[0]]
    created = 0
    for mfg in manufacturers:
        existing = WarehouseGroup.query.filter_by(manufacturer=mfg).first()
        if not existing:
            g = WarehouseGroup(name=mfg, manufacturer=mfg, description=f'Auto-created from manufacturer: {mfg}')
            db.session.add(g)
            created += 1
    db.session.commit()
    flash(_('{} groups created').format(created), 'success')
    return redirect(url_for('warehouse_groups'))

# ============================================================
# ROUTES — MAGAZIJN (existing)
# ============================================================

@app.route('/warehouse')
@login_required
@role_required('admin', 'director', 'technician')
def warehouse_list():
    cat = request.args.get('categorie', '')
    group_id = request.args.get('group', '')
    q = VoorraadItem.query
    if cat: q = q.filter_by(categorie=cat)
    if group_id: q = q.filter_by(group_id=int(group_id))
    items = q.order_by(VoorraadItem.naam).all()
    cats = [c[0] for c in db.session.query(VoorraadItem.categorie).distinct().all() if c[0]]
    groups = WarehouseGroup.query.order_by(WarehouseGroup.name).all()
    laag = [i for i in items if i.hoeveelheid <= i.minimum]
    return render_template('warehouse.html', items=items, categories=cats, category_filter=cat,
        groups=groups, group_filter=int(group_id) if group_id else None, low_stock=laag)

@app.route('/warehouse/new', methods=['GET', 'POST'])
@login_required
@role_required('admin', 'director', 'technician')
def warehouse_new():
    if request.method == 'POST':
        i = VoorraadItem(
            naam=request.form['naam'],
            description=request.form.get('description', ''),
            categorie=request.form.get('categorie',''),
            group_id=int(request.form['group_id']) if request.form.get('group_id') else None,
            supplier_part_number=request.form.get('supplier_part_number','').strip() or None,
            eenheid=request.form.get('eenheid','st'),
            hoeveelheid=float(request.form.get('hoeveelheid',0)),
            minimum=float(request.form.get('minimum',0)),
            prijs=float(request.form.get('prijs',0)),
            locatie=request.form.get('locatie',''),
            consumable_type=request.form.get('consumable_type',''),
            consumable_subtype=request.form.get('consumable_subtype',''),
            volume=request.form.get('volume',''),
            compatible_machines=request.form.get('compatible_machines',''),
            replacement_interval=request.form.get('replacement_interval',''),
            last_replacement=datetime.strptime(request.form['last_replacement'], '%Y-%m-%d').date() if request.form.get('last_replacement') else None,
            next_replacement=datetime.strptime(request.form['next_replacement'], '%Y-%m-%d').date() if request.form.get('next_replacement') else None,
        )
        db.session.add(i); db.session.commit()
        flash(_('Item added') + f': {i.naam}', 'success')
        return redirect(url_for('warehouse_list'))
    groups = WarehouseGroup.query.order_by(WarehouseGroup.name).all()
    return render_template('warehouse_form.html', item=None, groups=groups)

@app.route('/warehouse/<int:item_id>/edit', methods=['GET', 'POST'])
@login_required
@role_required('admin', 'director', 'technician')
def warehouse_edit(item_id):
    item = VoorraadItem.query.get_or_404(item_id)
    if request.method == 'POST':
        item.naam = request.form['naam']
        item.description = request.form.get('description', '')
        item.categorie = request.form.get('categorie','')
        item.group_id = int(request.form['group_id']) if request.form.get('group_id') else None
        item.supplier_part_number = request.form.get('supplier_part_number','').strip() or None
        item.eenheid = request.form.get('eenheid','st')
        item.hoeveelheid = float(request.form.get('hoeveelheid',0))
        item.minimum = float(request.form.get('minimum',0))
        item.prijs = float(request.form.get('prijs',0))
        item.locatie = request.form.get('locatie','')
        item.consumable_type = request.form.get('consumable_type','')
        item.consumable_subtype = request.form.get('consumable_subtype','')
        item.volume = request.form.get('volume','')
        item.compatible_machines = request.form.get('compatible_machines','')
        item.replacement_interval = request.form.get('replacement_interval','')
        item.last_replacement = datetime.strptime(request.form['last_replacement'], '%Y-%m-%d').date() if request.form.get('last_replacement') else None
        item.next_replacement = datetime.strptime(request.form['next_replacement'], '%Y-%m-%d').date() if request.form.get('next_replacement') else None
        db.session.commit()
        flash(_('Item updated'), 'success')
        return redirect(url_for('warehouse_list'))
    groups = WarehouseGroup.query.order_by(WarehouseGroup.name).all()
    return render_template('warehouse_form.html', item=item, groups=groups)

@app.route('/warehouse/duplicates')
@login_required
@role_required('admin')
def warehouse_duplicates():
    from sqlalchemy import func
    # Find items with same name (case-insensitive)
    dupes_name = db.session.query(
        func.lower(VoorraadItem.naam).label('name'),
        func.count().label('cnt')
    ).group_by(func.lower(VoorraadItem.naam)).having(func.count() > 1).all()
    dupe_groups = []
    for name, cnt in dupes_name:
        items = VoorraadItem.query.filter(func.lower(VoorraadItem.naam) == name).order_by(VoorraadItem.id).all()
        dupe_groups.append({'name': name, 'count': cnt, 'item_list': items})
    # Find items with same supplier_part_number
    dupes_spn = db.session.query(
        VoorraadItem.supplier_part_number.label('spn'),
        func.count().label('cnt')
    ).filter(VoorraadItem.supplier_part_number.isnot(None), VoorraadItem.supplier_part_number != '').group_by(VoorraadItem.supplier_part_number).having(func.count() > 1).all()
    spn_groups = []
    for spn, cnt in dupes_spn:
        items = VoorraadItem.query.filter_by(supplier_part_number=spn).order_by(VoorraadItem.id).all()
        spn_groups.append({'spn': spn, 'count': cnt, 'item_list': items})
    return render_template('warehouse_duplicates.html', dupe_groups=dupe_groups, spn_groups=spn_groups)

@app.route('/warehouse/<int:item_id>/delete', methods=['POST'])
@login_required
@role_required('admin')
def warehouse_delete(item_id):
    item = VoorraadItem.query.get_or_404(item_id)
    name = item.naam
    # Check for movements
    if item.mutaties:
        flash(_('Cannot delete item with movement history. Deactivate instead.'), 'error')
        return redirect(url_for('warehouse_edit', item_id=item.id))
    db.session.delete(item)
    db.session.commit()
    log_audit('delete', 'warehouse_item', item_id, name)
    flash(_('Item deleted') + f': {name}', 'success')
    return redirect(url_for('warehouse_duplicates'))

@app.route('/warehouse/<int:item_id>/move', methods=['POST'])
@login_required
@role_required('admin', 'director', 'technician')
def warehouse_move(item_id):
    item = VoorraadItem.query.get_or_404(item_id)
    mt = request.form['type']
    qty = float(request.form['hoeveelheid'])
    if mt == 'uitgaand' and qty > item.hoeveelheid:
        flash(_('Insufficient stock!'), 'error')
        return redirect(url_for('warehouse_list'))
    m = VoorraadMutatie(item_id=item_id, type=mt, hoeveelheid=qty,
                        opdracht_id=request.form.get('opdracht_id') or None,
                        opmerking=request.form.get('opmerking',''))
    if mt == 'inkomend': item.hoeveelheid += qty
    else: item.hoeveelheid -= qty
    db.session.add(m); db.session.commit()
    flash(_('{} {} {} — {}').format(mt.capitalize(), qty, item.eenheid, item.naam), 'success')
    return redirect(url_for('warehouse_list'))

# ============================================================
# ROUTES — INVOICES
# ============================================================

@app.route('/invoices')
@login_required
@role_required('admin', 'director')
def invoices_list():
    if current_user.has_role('admin', 'director'):
        invoices = Invoice.query.order_by(Invoice.created_at.desc()).all()
    else:
        invoices = Invoice.query.filter_by(created_by=current_user.id).order_by(Invoice.created_at.desc()).all()
    return render_template('invoices.html', invoices=invoices)

@app.route('/invoices/new', methods=['GET', 'POST'])
@login_required
@role_required('admin', 'director')
def invoice_new():
    if request.method == 'POST':
        inv = Invoice(
            invoice_number=request.form['invoice_number'],
            supplier=request.form['supplier'],
            invoice_date=datetime.strptime(request.form['invoice_date'], '%Y-%m-%d').date(),
            due_date=datetime.strptime(request.form['due_date'], '%Y-%m-%d').date() if request.form.get('due_date') else None,
            total=float(request.form.get('total', 0)),
            status='pending',
            notes=request.form.get('notes', ''),
            created_by=current_user.id
        )
        db.session.add(inv)
        db.session.commit()
        # Add items
        descriptions = request.form.getlist('item_desc[]')
        quantities = request.form.getlist('item_qty[]')
        prices = request.form.getlist('item_price[]')
        for i in range(len(descriptions)):
            if descriptions[i].strip():
                qty = float(quantities[i]) if i < len(quantities) and quantities[i] else 1
                price = float(prices[i]) if i < len(prices) and prices[i] else 0
                item = InvoiceItem(
                    invoice_id=inv.id, description=descriptions[i],
                    quantity=qty, unit_price=price, total_price=qty * price
                )
                db.session.add(item)
        db.session.commit()
        flash(_('Invoice created'), 'success')
        return redirect(url_for('invoice_detail', invoice_id=inv.id))
    items = VoorraadItem.query.order_by(VoorraadItem.naam).all()
    return render_template('invoice_form.html', invoice=None, warehouse_items=items)

@app.route('/invoices/<int:invoice_id>')
@login_required
@role_required('admin', 'director')
def invoice_detail(invoice_id):
    inv = Invoice.query.get_or_404(invoice_id)
    return render_template('invoice_detail.html', invoice=inv)

@app.route('/invoices/<int:invoice_id>/edit', methods=['GET', 'POST'])
@login_required
@role_required('admin', 'director')
def invoice_edit(invoice_id):
    inv = Invoice.query.get_or_404(invoice_id)
    if request.method == 'POST':
        inv.invoice_number = request.form['invoice_number']
        inv.supplier = request.form['supplier']
        inv.invoice_date = datetime.strptime(request.form['invoice_date'], '%Y-%m-%d').date()
        inv.due_date = datetime.strptime(request.form['due_date'], '%Y-%m-%d').date() if request.form.get('due_date') else None
        inv.total = float(request.form.get('total', 0))
        inv.notes = request.form.get('notes', '')
        # Update items
        inv.items = []
        descriptions = request.form.getlist('item_desc[]')
        quantities = request.form.getlist('item_qty[]')
        prices = request.form.getlist('item_price[]')
        for i in range(len(descriptions)):
            if descriptions[i].strip():
                qty = float(quantities[i]) if i < len(quantities) and quantities[i] else 1
                price = float(prices[i]) if i < len(prices) and prices[i] else 0
                item = InvoiceItem(
                    invoice_id=inv.id, description=descriptions[i],
                    quantity=qty, unit_price=price, total_price=qty * price
                )
                db.session.add(item)
        db.session.commit()
        flash(_('Invoice updated'), 'success')
        return redirect(url_for('invoice_detail', invoice_id=inv.id))
    items = VoorraadItem.query.order_by(VoorraadItem.naam).all()
    return render_template('invoice_form.html', invoice=inv, warehouse_items=items)

@app.route('/invoices/<int:invoice_id>/approve', methods=['POST'])
@login_required
@role_required('admin', 'director')
def invoice_approve(invoice_id):
    inv = Invoice.query.get_or_404(invoice_id)
    inv.status = 'approved'
    inv.signed_by = current_user.id
    inv.signed_at = datetime.utcnow()
    db.session.commit()
    flash(_('Invoice approved for payment'), 'success')
    return redirect(url_for('invoice_detail', invoice_id=inv.id))

@app.route('/invoices/<int:invoice_id>/reject', methods=['POST'])
@login_required
@role_required('admin', 'director')
def invoice_reject(invoice_id):
    inv = Invoice.query.get_or_404(invoice_id)
    inv.status = 'rejected'
    inv.rejection_reason = request.form.get('rejection_reason', '')
    inv.signed_by = current_user.id
    inv.signed_at = datetime.utcnow()
    db.session.commit()
    flash(_('Invoice rejected'), 'error')
    return redirect(url_for('invoice_detail', invoice_id=inv.id))

@app.route('/invoices/<int:invoice_id>/pay', methods=['POST'])
@login_required
@role_required('admin', 'director')
def invoice_pay(invoice_id):
    inv = Invoice.query.get_or_404(invoice_id)
    inv.status = 'paid'
    db.session.commit()
    flash(_('Invoice marked as paid'), 'success')
    return redirect(url_for('invoice_detail', invoice_id=inv.id))

@app.route('/invoices/<int:invoice_id>/delete', methods=['POST'])
@login_required
@role_required('admin')
def invoice_delete(invoice_id):
    inv = Invoice.query.get_or_404(invoice_id)
    db.session.delete(inv)
    db.session.commit()
    flash(_('Invoice deleted'), 'success')
    return redirect(url_for('invoices_list'))

# ============================================================
# ROUTES — RAPPORTEN (existing)
# ============================================================

@app.route('/reports')
@login_required
@role_required('admin', 'director', 'technician')
def reports():
    return render_template('reports.html')

@app.route('/reports/advanced')
@login_required
@role_required('admin', 'director')
def reports_advanced():
    report_type = request.args.get('type', 'activity')
    date_from = request.args.get('date_from', (datetime.utcnow() - timedelta(days=30)).strftime('%Y-%m-%d'))
    date_to = request.args.get('date_to', datetime.utcnow().strftime('%Y-%m-%d'))
    user_id = request.args.get('user_id', '')
    section_id = request.args.get('section_id', '')
    
    d_from = datetime.strptime(date_from, '%Y-%m-%d')
    d_to = datetime.strptime(date_to, '%Y-%m-%d') + timedelta(days=1)
    
    users = User.query.filter(User.is_active_user == True).order_by(User.display_name).all()
    sections = FactorySection.query.order_by(FactorySection.name).all()
    
    data = []
    stats = {}
    
    if report_type == 'activity':
        q = UserActivityLog.query.filter(UserActivityLog.created_at >= d_from, UserActivityLog.created_at < d_to)
        if user_id:
            q = q.filter_by(user_id=int(user_id))
        data = q.order_by(UserActivityLog.created_at.desc()).limit(500).all()
        stats['total_actions'] = q.count()
        stats['unique_users'] = db.session.query(db.func.count(db.distinct(UserActivityLog.user_id))).filter(UserActivityLog.created_at >= d_from, UserActivityLog.created_at < d_to).scalar()
        
    elif report_type == 'faults':
        q = FaultReport.query.filter(FaultReport.created_at >= d_from, FaultReport.created_at < d_to)
        if user_id:
            q = q.filter_by(reporter_id=int(user_id))
        if section_id:
            q = q.filter(FaultReport.machine.has(Machine.section_id == int(section_id)))
        data = q.order_by(FaultReport.created_at.desc()).all()
        stats['total'] = q.count()
        stats['open'] = q.filter(FaultReport.status.in_(['open', 'accepted', 'in_progress'])).count()
        stats['resolved'] = q.filter_by(status='resolved').count()
        stats['critical'] = q.filter_by(priority='critical').count()
        
    elif report_type == 'warehouse':
        q = VoorraadMutatie.query.filter(VoorraadMutatie.aangemaakt >= d_from, VoorraadMutatie.aangemaakt < d_to)
        data = q.order_by(VoorraadMutatie.aangemaakt.desc()).limit(500).all()
        stats['total_movements'] = q.count()
        stats['incoming'] = q.filter_by(type='inkomend').count()
        stats['outgoing'] = q.filter_by(type='uitgaand').count()
        
    elif report_type == 'errors':
        q = SystemLog.query.filter(SystemLog.created_at >= d_from, SystemLog.created_at < d_to)
        if user_id:
            q = q.filter_by(user_id=int(user_id))
        data = q.order_by(SystemLog.created_at.desc()).limit(500).all()
        stats['total'] = q.count()
        stats['errors'] = q.filter_by(level='ERROR').count()
        stats['warnings'] = q.filter_by(level='WARNING').count()
        
    elif report_type == 'users':
        q = AuditLog.query.filter(AuditLog.created_at >= d_from, AuditLog.created_at < d_to)
        if user_id:
            q = q.filter_by(user_id=int(user_id))
        data = q.order_by(AuditLog.created_at.desc()).limit(500).all()
        stats['total'] = q.count()
        
    elif report_type == 'responsible':
        # Report on responsible persons - who submitted what and when
        persons = Verantwoordelijke.query.filter(Verantwoordelijke.is_active == True).order_by(Verantwoordelijke.naam).all()
        responsible_data = []
        for p in persons:
            faults = FaultReport.query.filter(
                FaultReport.reporter_id.in_(
                    db.session.query(User.id).filter(User.person_id == p.id)
                ),
                FaultReport.created_at >= d_from,
                FaultReport.created_at < d_to
            ).order_by(FaultReport.created_at.desc()).all()
            if faults or not user_id:
                responsible_data.append({
                    'person': p,
                    'faults': faults,
                    'total': len(faults),
                    'open': len([f for f in faults if f.status in ['open', 'accepted', 'in_progress']]),
                    'resolved': len([f for f in faults if f.status == 'resolved']),
                })
        data = responsible_data
        stats['total_persons'] = len(responsible_data)
        stats['total_faults'] = sum(r['total'] for r in responsible_data)
        
    elif report_type == 'machines':
        # Report by machines - faults, maintenance, status
        machines_q = Machine.query.order_by(Machine.name)
        if section_id:
            machines_q = machines_q.filter_by(section_id=int(section_id))
        machines_list = machines_q.all()
        machine_data = []
        for m in machines_list:
            faults = FaultReport.query.filter(
                FaultReport.machine_id == m.id,
                FaultReport.created_at >= d_from,
                FaultReport.created_at < d_to
            ).order_by(FaultReport.created_at.desc()).all()
            total_faults = FaultReport.query.filter(FaultReport.machine_id == m.id).count()
            machine_data.append({
                'machine': m,
                'faults': faults,
                'period_count': len(faults),
                'total_count': total_faults,
                'open': len([f for f in faults if f.status in ['open', 'accepted', 'in_progress']]),
                'critical': len([f for f in faults if f.priority == 'critical']),
            })
        data = machine_data
        stats['total_machines'] = len(machine_data)
        stats['total_faults'] = sum(r['period_count'] for r in machine_data)
        stats['machines_with_faults'] = len([r for r in machine_data if r['period_count'] > 0])
        
    return render_template('reports_advanced.html',
        report_type=report_type, data=data, stats=stats,
        users=users, sections=sections,
        date_from=date_from, date_to=date_to,
        user_id=user_id, section_id=section_id)

@app.route('/reports/advanced/export')
@login_required
@role_required('admin', 'director')
def reports_export():
    report_type = request.args.get('type', 'activity')
    format_type = request.args.get('format', 'csv')
    date_from = request.args.get('date_from', (datetime.utcnow() - timedelta(days=30)).strftime('%Y-%m-%d'))
    date_to = request.args.get('date_to', datetime.utcnow().strftime('%Y-%m-%d'))
    user_id = request.args.get('user_id', '')
    
    d_from = datetime.strptime(date_from, '%Y-%m-%d')
    d_to = datetime.strptime(date_to, '%Y-%m-%d') + timedelta(days=1)
    
    # Collect data
    headers = []
    rows = []
    title = ''
    
    if report_type == 'activity':
        title = 'Активность пользователей'
        headers = ['Дата', 'Пользователь', 'Действие', 'Страница', 'Детали', 'IP']
        q = UserActivityLog.query.filter(UserActivityLog.created_at >= d_from, UserActivityLog.created_at < d_to)
        if user_id: q = q.filter_by(user_id=int(user_id))
        for r in q.order_by(UserActivityLog.created_at.desc()).limit(1000).all():
            rows.append([r.created_at.strftime('%Y-%m-%d %H:%M'), r.username or '', r.action or '', r.page or '', r.details or '', r.ip_address or ''])
    elif report_type == 'faults':
        title = 'Заявки о неисправности'
        headers = ['ID', 'Дата', 'Заголовок', 'Станок', 'Приоритет', 'Статус', 'Репортер']
        q = FaultReport.query.filter(FaultReport.created_at >= d_from, FaultReport.created_at < d_to)
        if user_id: q = q.filter_by(reporter_id=int(user_id))
        for f in q.order_by(FaultReport.created_at.desc()).all():
            rows.append([str(f.id), f.created_at.strftime('%Y-%m-%d %H:%M'), f.title or '', f.machine.name if f.machine else '', f.priority or '', f.status or '', f.reporter.display_name if f.reporter else ''])
    elif report_type == 'warehouse':
        title = 'Движение склада'
        headers = ['Дата', 'Товар', 'Тип', 'Количество', 'Комментарий']
        q = VoorraadMutatie.query.filter(VoorraadMutatie.aangemaakt >= d_from, VoorraadMutatie.aangemaakt < d_to)
        for m in q.order_by(VoorraadMutatie.aangemaakt.desc()).limit(1000).all():
            rows.append([m.aangemaakt.strftime('%Y-%m-%d %H:%M'), m.item.naam if m.item else '', m.type or '', str(m.hoeveelheid), m.opmerking or ''])
    elif report_type == 'errors':
        title = 'Ошибки и предупреждения'
        headers = ['Дата', 'Уровень', 'Категория', 'Сообщение', 'Источник', 'Пользователь']
        q = SystemLog.query.filter(SystemLog.created_at >= d_from, SystemLog.created_at < d_to)
        if user_id: q = q.filter_by(user_id=int(user_id))
        for r in q.order_by(SystemLog.created_at.desc()).limit(1000).all():
            rows.append([r.created_at.strftime('%Y-%m-%d %H:%M'), r.level or '', r.category or '', r.message or '', r.source or '', r.user.display_name if r.user else ''])
    elif report_type == 'users':
        title = 'Журнал аудита'
        headers = ['Дата', 'Пользователь', 'Действие', 'Тип', 'Детали', 'IP']
        q = AuditLog.query.filter(AuditLog.created_at >= d_from, AuditLog.created_at < d_to)
        if user_id: q = q.filter_by(user_id=int(user_id))
        for r in q.order_by(AuditLog.created_at.desc()).limit(1000).all():
            rows.append([r.created_at.strftime('%Y-%m-%d %H:%M'), r.user.display_name if r.user else '', r.action or '', r.entity_type or '', r.details or '', r.ip_address or ''])
    elif report_type == 'responsible':
        title = 'Отчёт по ответственным'
        headers = ['Ответственный', 'Должность', 'Телефон', 'Внутр. номер', 'Email', 'Всего заявок', 'Открытых', 'Решённых']
        persons = Verantwoordelijke.query.filter(Verantwoordelijke.is_active == True).order_by(Verantwoordelijke.naam).all()
        for p in persons:
            fault_count = FaultReport.query.filter(
                FaultReport.reporter_id.in_(db.session.query(User.id).filter(User.person_id == p.id)),
                FaultReport.created_at >= d_from, FaultReport.created_at < d_to
            ).count()
            open_count = FaultReport.query.filter(
                FaultReport.reporter_id.in_(db.session.query(User.id).filter(User.person_id == p.id)),
                FaultReport.created_at >= d_from, FaultReport.created_at < d_to,
                FaultReport.status.in_(['open', 'accepted', 'in_progress'])
            ).count()
            resolved_count = FaultReport.query.filter(
                FaultReport.reporter_id.in_(db.session.query(User.id).filter(User.person_id == p.id)),
                FaultReport.created_at >= d_from, FaultReport.created_at < d_to,
                FaultReport.status == 'resolved'
            ).count()
            rows.append([p.naam or '', p.position or '', p.telefoon or '', p.internal_phone or '', p.email or '', str(fault_count), str(open_count), str(resolved_count)])
    elif report_type == 'machines':
        title = 'Отчёт по станкам'
        headers = ['Станок', 'Тип', 'Серийный номер', 'Отдел', 'Заявок за период', 'Всего заявок', 'Открытых', 'Критичных']
        machines_q = Machine.query.order_by(Machine.name)
        if section_id: machines_q = machines_q.filter_by(section_id=int(section_id))
        for m in machines_q.all():
            period_count = FaultReport.query.filter(FaultReport.machine_id == m.id, FaultReport.created_at >= d_from, FaultReport.created_at < d_to).count()
            total_count = FaultReport.query.filter(FaultReport.machine_id == m.id).count()
            open_count = FaultReport.query.filter(FaultReport.machine_id == m.id, FaultReport.status.in_(['open', 'accepted', 'in_progress'])).count()
            critical_count = FaultReport.query.filter(FaultReport.machine_id == m.id, FaultReport.priority == 'critical').count()
            rows.append([m.name or '', m.machine_type or '', m.serial_number or '', m.section.name if m.section else '', str(period_count), str(total_count), str(open_count), str(critical_count)])
    
    period = f'{date_from} — {date_to}'
    filename = f'report_{report_type}_{date_from}_{date_to}'
    
    # === CSV ===
    if format_type == 'csv':
        import csv, io
        output = io.StringIO()
        writer = csv.writer(output)
        writer.writerow(headers)
        writer.writerows(rows)
        output.seek(0)
        from flask import Response
        return Response(output.getvalue(), mimetype='text/csv',
            headers={'Content-Disposition': f'attachment;filename={filename}.csv'})
    
    # === EXCEL ===
    elif format_type == 'excel':
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        wb = Workbook()
        ws = wb.active
        ws.title = title[:31]
        
        # Title
        ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=len(headers))
        ws['A1'] = title
        ws['A1'].font = Font(bold=True, size=14)
        ws['A2'] = f'Период: {period}'
        ws['A2'].font = Font(size=10, color='666666')
        
        # Headers
        header_fill = PatternFill(start_color='2C3E50', end_color='2C3E50', fill_type='solid')
        header_font = Font(bold=True, color='FFFFFF', size=11)
        thin_border = Border(
            left=Side(style='thin'), right=Side(style='thin'),
            top=Side(style='thin'), bottom=Side(style='thin'))
        
        for col, h in enumerate(headers, 1):
            cell = ws.cell(row=4, column=col, value=h)
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(horizontal='center')
            cell.border = thin_border
        
        # Data
        for row_idx, row_data in enumerate(rows, 5):
            for col_idx, val in enumerate(row_data, 1):
                cell = ws.cell(row=row_idx, column=col_idx, value=val)
                cell.border = thin_border
                cell.alignment = Alignment(wrap_text=True)
        
        # Auto-width
        for col in range(1, len(headers) + 1):
            max_len = len(headers[col-1])
            for row in range(5, len(rows) + 5):
                cell_val = str(ws.cell(row=row, column=col).value or '')
                max_len = max(max_len, min(len(cell_val), 40))
            ws.column_dimensions[ws.cell(row=4, column=col).column_letter].width = max_len + 4
        
        buf = io.BytesIO()
        wb.save(buf)
        buf.seek(0)
        return send_file(buf, mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            download_name=f'{filename}.xlsx', as_attachment=True)
    
    # === WORD ===
    elif format_type == 'word':
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.table import WD_TABLE_ALIGNMENT
        doc = Document()
        
        # Title
        doc.add_heading(title, level=1)
        doc.add_paragraph(f'Период: {period}')
        doc.add_paragraph(f'Сгенерировано: {datetime.utcnow().strftime("%d.%m.%Y %H:%M")}')
        
        # Table
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Light Grid Accent 1'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Headers
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(9)
        
        # Data
        for row_data in rows:
            row = table.add_row()
            for i, val in enumerate(row_data):
                row.cells[i].text = str(val)
                for paragraph in row.cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(8)
        
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return send_file(buf, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            download_name=f'{filename}.docx', as_attachment=True)
    
    # === PDF ===
    elif format_type == 'pdf':
        # Generate HTML table for PDF
        html = f'''<!DOCTYPE html>
<html><head><meta charset="utf-8">
<style>
body {{ font-family: Arial; font-size: 10px; padding: 20px; }}
h1 {{ font-size: 16px; margin-bottom: 5px; }}
p {{ color: #666; font-size: 11px; margin-bottom: 15px; }}
table {{ width: 100%; border-collapse: collapse; font-size: 9px; }}
th {{ background: #2c3e50; color: white; padding: 6px 8px; text-align: left; font-weight: bold; }}
td {{ padding: 5px 8px; border-bottom: 1px solid #eee; }}
tr:nth-child(even) {{ background: #f9f9f9; }}
.footer {{ margin-top: 20px; font-size: 8px; color: #999; text-align: center; }}
</style></head><body>
<h1>{title}</h1>
<p>Период: {period} | Сгенерировано: {datetime.utcnow().strftime("%d.%m.%Y %H:%M")}</p>
<table><tr>'''
        for h in headers:
            html += f'<th>{h}</th>'
        html += '</tr>'
        for row in rows:
            html += '<tr>'
            for val in row:
                html += f'<td>{val}</td>'
            html += '</tr>'
        html += f'</table><div class="footer">CRM Мастерская — {title} — {period}</div></body></html>'
        
        try:
            import pdfkit
            pdf = pdfkit.from_string(html, False)
            from flask import Response
            return Response(pdf, mimetype='application/pdf',
                headers={'Content-Disposition': f'attachment;filename={filename}.pdf'})
        except Exception:
            # Fallback: return HTML that can be printed as PDF
            from flask import Response
            return Response(html, mimetype='text/html',
                headers={'Content-Disposition': f'attachment;filename={filename}.html'})
    
    return jsonify({'error': 'Unknown format'}), 400

@app.route('/reports/period', methods=['GET', 'POST'])
@login_required
@role_required('admin', 'director', 'technician')
def report_period():
    if request.method == 'POST':
        d_from = datetime.strptime(request.form['date_from'], '%Y-%m-%d')
        d_to = datetime.strptime(request.form['date_to'], '%Y-%m-%d') + timedelta(days=1)
    else:
        d_to = datetime.utcnow()
        d_from = d_to - timedelta(days=30)
    orders = Opdracht.query.filter(Opdracht.aangemaakt >= d_from, Opdracht.aangemaakt < d_to).order_by(Opdracht.aangemaakt.desc()).all()
    omzet = sum(o.totaal for o in orders if o.status == 'afgeleverd')
    voltooid = len([o for o in orders if o.status == 'afgeleverd'])
    geannuleerd = len([o for o in orders if o.status == 'geannuleerd'])
    # Pre-fetch all monteurs in one query
    monteur_ids = list(set(o.monteur_id for o in orders if o.monteur_id))
    monteurs_map = {m.id: m.naam for m in Monteur.query.filter(Monteur.id.in_(monteur_ids)).all()} if monteur_ids else {}
    ws = {}
    for o in orders:
        if o.monteur_id:
            if o.monteur_id not in ws:
                ws[o.monteur_id] = {'naam': monteurs_map.get(o.monteur_id, 'Onbekend'), 'orders': 0, 'omzet': 0, 'voltooid': 0}
            ws[o.monteur_id]['orders'] += 1
            if o.status == 'afgeleverd':
                ws[o.monteur_id]['omzet'] += o.totaal
                ws[o.monteur_id]['voltooid'] += 1
    return render_template('report_period.html', orders=orders,
                         date_from=d_from.strftime('%Y-%m-%d'),
                         date_to=(d_to - timedelta(days=1)).strftime('%Y-%m-%d'),
                         total_revenue=omzet, total_orders=len(orders),
                         completed=voltooid, cancelled=geannuleerd, worker_stats=ws)

@app.route('/reports/worker/<int:worker_id>')
@login_required
@role_required('admin', 'director', 'technician')
def report_worker(worker_id):
    w = Monteur.query.get_or_404(worker_id)
    df = request.args.get('date_from', (datetime.utcnow() - timedelta(days=30)).strftime('%Y-%m-%d'))
    dt = request.args.get('date_to', datetime.utcnow().strftime('%Y-%m-%d'))
    orders = Opdracht.query.filter(Opdracht.monteur_id == worker_id,
        Opdracht.aangemaakt >= datetime.strptime(df, '%Y-%m-%d'),
        Opdracht.aangemaakt < datetime.strptime(dt, '%Y-%m-%d') + timedelta(days=1)
    ).order_by(Opdracht.aangemaakt.desc()).all()
    omzet = sum(o.totaal for o in orders if o.status == 'afgeleverd')
    voltooid = len([o for o in orders if o.status == 'afgeleverd'])
    avg = 0
    gereed = [o for o in orders if o.gereed and o.gestart]
    if gereed: avg = sum((o.gereed - o.gestart).total_seconds()/3600 for o in gereed) / len(gereed)
    return render_template('report_worker.html', worker=w, orders=orders, date_from=df, date_to=dt,
                         total_revenue=omzet, completed=voltooid, avg_time=round(avg,1))

@app.route('/api/stats')
@login_required
def api_stats():
    days = []
    for i in range(29, -1, -1):
        dag = datetime.utcnow().date() - timedelta(days=i)
        c = Opdracht.query.filter(Opdracht.aangemaakt >= datetime.combine(dag, datetime.min.time()),
            Opdracht.aangemaakt < datetime.combine(dag + timedelta(days=1), datetime.min.time())).count()
        days.append({'date': dag.strftime('%d.%m'), 'count': c})
    statuses = {}
    for s in ['aangenomen','diagnose','in behandeling','gereed','afgeleverd','geannuleerd']:
        statuses[s] = Opdracht.query.filter_by(status=s).count()
    wdata = []
    for w in Monteur.query.filter_by(actief=True).all():
        cnt = Opdracht.query.filter_by(monteur_id=w.id, status='afgeleverd').count()
        rev = db.session.query(db.func.sum(Opdracht.totaal)).filter_by(monteur_id=w.id, status='afgeleverd').scalar() or 0
        wdata.append({'id': w.id, 'name': w.naam, 'orders': cnt, 'revenue': float(rev)})
    return jsonify({'days': days, 'statuses': statuses, 'workers': wdata})

# ============================================================
# ROUTES — QR CODE (existing)
# ============================================================

@app.route('/qr/scan')
@login_required
def qr_scan():
    return render_template('qr_scan.html')

@app.route('/qr/generate/<int:order_id>')
@login_required
def qr_generate(order_id):
    order = Opdracht.query.get_or_404(order_id)
    data = {
        'type': 'opdracht',
        'id': order.id,
        'nummer': order.nummer,
        'apparaat': order.apparaat,
        'model': order.model or '',
        'serienummer': order.serienummer or '',
        'klant': order.verantwoordelijke.naam,
        'status': order.status
    }
    qr = qrcode.QRCode(version=1, box_size=10, border=4)
    qr.add_data(json.dumps(data, ensure_ascii=False))
    qr.make(fit=True)
    img = qr.make_image(fill_color="black", back_color="white")
    buf = io.BytesIO()
    img.save(buf, format='PNG')
    buf.seek(0)
    return send_file(buf, mimetype='image/png', download_name=f'QR_{order.nummer}.png')

@app.route('/api/machine/<int:machine_id>/faults')
@login_required
def api_machine_faults(machine_id):
    """Return faults and stats for a machine (used by floor plan)"""
    m = Machine.query.get_or_404(machine_id)
    faults = FaultReport.query.filter_by(machine_id=m.id).order_by(FaultReport.created_at.desc()).limit(10).all()
    total = FaultReport.query.filter_by(machine_id=m.id).count()
    open_count = FaultReport.query.filter(FaultReport.machine_id == m.id, FaultReport.status.in_(['open', 'accepted', 'in_progress'])).count()
    resolved = FaultReport.query.filter(FaultReport.machine_id == m.id, FaultReport.status == 'resolved').count()
    critical = FaultReport.query.filter(FaultReport.machine_id == m.id, FaultReport.priority == 'critical').count()
    
    faults_data = [{
        'id': f.id,
        'title': f.title or '',
        'status': f.status or '',
        'priority': f.priority or '',
        'date': f.created_at.strftime('%d-%m-%Y'),
    } for f in faults]
    
    return jsonify({
        'total': total,
        'open': open_count,
        'resolved': resolved,
        'critical': critical,
        'faults': faults_data
    })

@app.route('/faults/<int:fault_id>/delete', methods=['POST'])
@login_required
@role_required('admin')
def fault_delete(fault_id):
    """Delete a fault report"""
    f = FaultReport.query.get_or_404(fault_id)
    title = f.title
    db.session.delete(f)
    db.session.commit()
    log_audit('delete', 'fault', fault_id, title)
    if request.is_json:
        return jsonify({'ok': True})
    flash(_('Fault deleted'), 'success')
    return redirect(url_for('faults_list'))

@app.route('/api/machines/<int:machine_id>/qr')
@login_required
def machine_qr(machine_id):
    m = Machine.query.get_or_404(machine_id)
    data = {
        'type': 'machine',
        'id': m.id,
        'name': m.name,
        'serial': m.serial_number or '',
        'manufacturer': m.manufacturer or '',
        'location': m.installation_location or '',
        'section': m.section.name if m.section else ''
    }
    qr = qrcode.QRCode(version=1, box_size=10, border=4)
    qr.add_data(json.dumps(data, ensure_ascii=False))
    qr.make(fit=True)
    img = qr.make_image(fill_color="black", back_color="white")
    buf = io.BytesIO()
    img.save(buf, format='PNG')
    buf.seek(0)
    return send_file(buf, mimetype='image/png', download_name=f'QR_{m.name}.png')

@app.route('/api/machines/<int:machine_id>/barcode')
@login_required
def machine_barcode(machine_id):
    from PIL import Image, ImageDraw, ImageFont
    m = Machine.query.get_or_404(machine_id)
    code_str = f"M{m.id:05d}"
    
    try:
        import barcode
        from barcode.writer import ImageWriter
        code128 = barcode.get('code128', code_str, writer=ImageWriter())
        buf = io.BytesIO()
        code128.write(buf, options={'module_width': 0.3, 'module_height': 8, 'font_size': 8, 'text_distance': 2, 'quiet_zone': 2})
        buf.seek(0)
        return send_file(buf, mimetype='image/png', download_name=f'BAR_{m.name}.png')
    except ImportError:
        pass
    
    # Fallback: generate with Pillow
    def code128_encode(text):
        chars = ' !"#$%&\'()*+,-./0123456789:;<=>?@ABCDEFGHIJKLMNOPQRSTUVWXYZ[\\]^_`abcdefghijklmnopqrstuvwxyz{|}~'
        codes = [104]
        checksum = 104
        for i, c in enumerate(text):
            if c in chars:
                val = chars.index(c) + 32
                codes.append(val)
                checksum += val * (i + 1)
        codes.append(checksum % 103)
        codes.append(106)
        patterns = [
            '11011001100','11001101100','11001100110','10010011000','10010001100',
            '10001001100','10011001000','10011000100','10001100100','11001001000',
            '11001000100','11000100100','10110011100','10011011100','10011001110',
            '10111001100','10011101100','10011100110','11001110010','11001011100',
            '11001001110','11011100100','11001110100','11101101110','11101001100',
            '11100101100','11100100110','11101100100','11100110100','11100110010',
            '11011011000','11011000110','11000110110','10100011000','10001011000',
            '10001000110','10110001000','10001101000','10001100010','11010001000',
            '11000101000','11000100010','10110111000','10110001110','10001101110',
            '10111011000','10111000110','10001110110','11101110110','11010001110',
            '11000101110','11011101000','11011100010','11011101110','11101011000',
            '11101000110','11100010110','11101101000','11101100010','11100011010',
            '11101111010','11001000010','11110001010','10100110000','10100001100',
            '10010110000','10010000110','10000101100','10000100110','10110010000',
            '10110000100','10011010000','10011000010','10000110100','10000110010',
            '11000010010','11001010000','11110111010','11000010100','10001111010',
            '10100111100','10010111100','10010011110','10111100100','10011110100',
            '10011110010','11110100100','11110010100','11110010010','11011011110',
            '11011110110','11110110110','10101111000','10100011110','10001011110',
            '10111101000','10111100010','11110101000','11110100010','10111011110',
            '10111101110','11101011110','11110101110','11010000100','11010010000',
            '11010011100','1100011101011'
        ]
        bars = []
        for c in codes:
            if c < len(patterns):
                bars.append(patterns[c])
        return bars
    
    bar_width = 2
    height = 60
    text_height = 16
    bars = code128_encode(code_str)
    total_width = sum(len(b) for b in bars) * bar_width + 20
    img = Image.new('RGB', (total_width, height + text_height + 4), 'white')
    draw = ImageDraw.Draw(img)
    x = 10
    for bar_pattern in bars:
        for bit in bar_pattern:
            if bit == '1':
                draw.rectangle([x, 0, x + bar_width - 1, height - 1], fill='black')
            x += bar_width
    try:
        font = ImageFont.truetype("arial.ttf", 12)
    except:
        font = ImageFont.load_default()
    bbox = draw.textbbox((0, 0), code_str, font=font)
    tw = bbox[2] - bbox[0]
    draw.text(((total_width - tw) // 2, height + 2), code_str, fill='black', font=font)
    buf = io.BytesIO()
    img.save(buf, format='PNG')
    buf.seek(0)
    return send_file(buf, mimetype='image/png', download_name=f'BAR_{m.name}.png')

@app.route('/machines/<int:machine_id>/qr-label')
@login_required
def machine_qr_label(machine_id):
    m = Machine.query.get_or_404(machine_id)
    return render_template('machine_qr_label.html', machine=m)

@app.route('/machines/qr-labels')
@login_required
def machines_qr_labels():
    ids_str = request.args.get('ids', '')
    if not ids_str:
        flash(_('Select machines first'), 'error')
        return redirect(url_for('machines_list'))
    ids = [int(x) for x in ids_str.split(',') if x.strip().isdigit()]
    machines = Machine.query.filter(Machine.id.in_(ids)).order_by(Machine.name).all()
    return render_template('machines_qr_labels.html', machines=machines)

@app.route('/qr/product', methods=['POST'])
@login_required
def qr_product():
    data = request.get_json()
    if not data:
        return jsonify({'error': 'Geen data'}), 400
    klant = None
    if data.get('klant'):
        klant = Verantwoordelijke.query.filter(Verantwoordelijke.naam.ilike(f"%{data['klant']}%")).first()
    if not klant and data.get('telefoon'):
        klant = Verantwoordelijke.query.filter(Verantwoordelijke.telefoon.like(f"%{data['telefoon']}%")).first()
    result = {
        'apparaat': data.get('apparaat', ''),
        'model': data.get('model', ''),
        'serienummer': data.get('serienummer', ''),
        'probleem': data.get('probleem', ''),
        'klant_id': klant.id if klant else None,
        'klant_naam': klant.naam if klant else data.get('klant', '')
    }
    return jsonify(result)

@app.route('/api/qr/lookup', methods=['POST'])
@login_required
def qr_lookup():
    data = request.get_json()
    code = data.get('code', '')
    try:
        parsed = json.loads(code)
        if parsed.get('type') == 'opdracht':
            order = Opdracht.query.get(parsed.get('id'))
            if order:
                return jsonify({
                    'found': True, 'type': 'opdracht',
                    'id': order.id, 'nummer': order.nummer,
                    'apparaat': order.apparaat, 'model': order.model,
                    'klant': order.verantwoordelijke.naam, 'status': order.status,
                    'totaal': order.totaal
                })
    except (json.JSONDecodeError, AttributeError):
        pass
    order = Opdracht.query.filter_by(serienummer=code).order_by(Opdracht.aangemaakt.desc()).first()
    if order:
        return jsonify({
            'found': True, 'type': 'serienummer',
            'id': order.id, 'nummer': order.nummer,
            'apparaat': order.apparaat, 'model': order.model,
            'klant': order.verantwoordelijke.naam, 'status': order.status
        })
    klant = Verantwoordelijke.query.filter(Verantwoordelijke.telefoon.like(f"%{code}%")).first()
    if klant:
        orders = Opdracht.query.filter_by(responsible_id=klant.id).order_by(Opdracht.aangemaakt.desc()).limit(5).all()
        return jsonify({
            'found': True, 'type': 'klant',
            'klant': klant.naam, 'telefoon': klant.telefoon,
            'orders': [{'id': o.id, 'nummer': o.nummer, 'apparaat': o.apparaat, 'status': o.status} for o in orders]
        })
    return jsonify({'found': False, 'code': code})

# ============================================================
# ROUTES — WAREHOUSE QR & PARTS SEARCH
# ============================================================

@app.route('/api/warehouse/search')
@login_required
def warehouse_search():
    q = request.args.get('q', '')
    if not q:
        items = VoorraadItem.query.order_by(VoorraadItem.naam).limit(50).all()
    else:
        items = VoorraadItem.query.filter(
            (VoorraadItem.naam.ilike(f'%{q}%')) | 
            (VoorraadItem.categorie.ilike(f'%{q}%'))
        ).order_by(VoorraadItem.naam).all()
    return jsonify([{
        'id': i.id, 'name': i.naam, 'category': i.categorie or '',
        'quantity': i.hoeveelheid, 'unit': i.eenheid, 'price': i.prijs,
        'available': i.hoeveelheid > 0
    } for i in items])

@app.route('/api/warehouse/qr/<int:item_id>')
@login_required
def warehouse_qr(item_id):
    item = VoorraadItem.query.get_or_404(item_id)
    data = {
        'type': 'warehouse_item',
        'id': item.id,
        'name': item.naam,
        'category': item.categorie or '',
        'location': item.locatie or '',
        'quantity': item.hoeveelheid,
        'unit': item.eenheid
    }
    qr = qrcode.QRCode(version=1, box_size=10, border=4)
    qr.add_data(json.dumps(data, ensure_ascii=False))
    qr.make(fit=True)
    img = qr.make_image(fill_color="black", back_color="white")
    buf = io.BytesIO()
    img.save(buf, format='PNG')
    buf.seek(0)
    return send_file(buf, mimetype='image/png', download_name=f'QR_{item.naam}.png')

@app.route('/api/warehouse/barcode/<int:item_id>')
@login_required
def warehouse_barcode(item_id):
    from PIL import Image, ImageDraw, ImageFont
    item = VoorraadItem.query.get_or_404(item_id)
    code_str = item.supplier_part_number if item.supplier_part_number else f"W{item.id:05d}"
    
    try:
        import barcode
        from barcode.writer import ImageWriter
        code128 = barcode.get('code128', code_str, writer=ImageWriter())
        buf = io.BytesIO()
        code128.write(buf, options={'module_width': 0.3, 'module_height': 8, 'font_size': 8, 'text_distance': 2, 'quiet_zone': 2})
        buf.seek(0)
        return send_file(buf, mimetype='image/png', download_name=f'BAR_{item.naam}.png')
    except ImportError:
        # Fallback: generate barcode with Pillow using Code128B encoding
        pass
    
    # Code128 encoding table (subset for alphanumeric)
    def code128_encode(text):
        # Code128B character set
        chars = ' !"#$%&\'()*+,-./0123456789:;<=>?@ABCDEFGHIJKLMNOPQRSTUVWXYZ[\\]^_`abcdefghijklmnopqrstuvwxyz{|}~'
        # Start code B = 104, Stop = 106
        codes = [104]  # Start B
        checksum = 104
        for i, c in enumerate(text):
            if c in chars:
                val = chars.index(c) + 32
                codes.append(val)
                checksum += val * (i + 1)
            else:
                codes.append(0)  # fallback
        codes.append(checksum % 103)
        codes.append(106)  # Stop
        
        # Code128 bar patterns (widths of bars and spaces)
        patterns = [
            '11011001100','11001101100','11001100110','10010011000','10010001100',
            '10001001100','10011001000','10011000100','10001100100','11001001000',
            '11001000100','11000100100','10110011100','10011011100','10011001110',
            '10111001100','10011101100','10011100110','11001110010','11001011100',
            '11001001110','11011100100','11001110100','11101101110','11101001100',
            '11100101100','11100100110','11101100100','11100110100','11100110010',
            '11011011000','11011000110','11000110110','10100011000','10001011000',
            '10001000110','10110001000','10001101000','10001100010','11010001000',
            '11000101000','11000100010','10110111000','10110001110','10001101110',
            '10111011000','10111000110','10001110110','11101110110','11010001110',
            '11000101110','11011101000','11011100010','11011101110','11101011000',
            '11101000110','11100010110','11101101000','11101100010','11100011010',
            '11101111010','11001000010','11110001010','10100110000','10100001100',
            '10010110000','10010000110','10000101100','10000100110','10110010000',
            '10110000100','10011010000','10011000010','10000110100','10000110010',
            '11000010010','11001010000','11110111010','11000010100','10001111010',
            '10100111100','10010111100','10010011110','10111100100','10011110100',
            '10011110010','11110100100','11110010100','11110010010','11011011110',
            '11011110110','11110110110','10101111000','10100011110','10001011110',
            '10111101000','10111100010','11110101000','11110100010','10111011110',
            '10111101110','11101011110','11110101110','11010000100','11010010000',
            '11010011100','1100011101011'
        ]
        bars = []
        for c in codes:
            if c < len(patterns):
                bars.append(patterns[c])
        return bars
    
    # Generate image
    bar_width = 2
    height = 60
    text_height = 16
    total_height = height + text_height + 4
    
    bars = code128_encode(code_str)
    total_width = sum(len(b) for b in bars) * bar_width + 20  # margins
    
    img = Image.new('RGB', (total_width, total_height), 'white')
    draw = ImageDraw.Draw(img)
    
    x = 10
    for bar_pattern in bars:
        for i, bit in enumerate(bar_pattern):
            if bit == '1':
                draw.rectangle([x, 0, x + bar_width - 1, height - 1], fill='black')
            x += bar_width
    
    # Draw text
    try:
        font = ImageFont.truetype("arial.ttf", 12)
    except:
        font = ImageFont.load_default()
    bbox = draw.textbbox((0, 0), code_str, font=font)
    text_width = bbox[2] - bbox[0]
    text_x = (total_width - text_width) // 2
    draw.text((text_x, height + 2), code_str, fill='black', font=font)
    
    buf = io.BytesIO()
    img.save(buf, format='PNG')
    buf.seek(0)
    return send_file(buf, mimetype='image/png', download_name=f'BAR_{item.naam}.png')

@app.route('/warehouse/movements')
@login_required
@role_required('admin', 'director', 'technician')
def warehouse_movements():
    item_id = request.args.get('item', '')
    move_type = request.args.get('type', '')
    q = VoorraadMutatie.query
    if item_id:
        q = q.filter_by(item_id=int(item_id))
    if move_type:
        q = q.filter_by(type=move_type)
    movements = q.order_by(VoorraadMutatie.aangemaakt.desc()).limit(200).all()
    items = VoorraadItem.query.order_by(VoorraadItem.naam).all()
    return render_template('warehouse_movements.html', movements=movements, items=items,
                         item_filter=int(item_id) if item_id else None, type_filter=move_type)

@app.route('/warehouse/report')
@login_required
@role_required('admin', 'director', 'technician')
def warehouse_search_report():
    from sqlalchemy import func, and_, or_
    # Filters
    q_text = request.args.get('q', '').strip()
    spn = request.args.get('spn', '').strip()
    locatie = request.args.get('locatie', '').strip()
    categorie = request.args.get('categorie', '').strip()
    move_type = request.args.get('move_type', '').strip()
    date_from = request.args.get('date_from', '')
    date_to = request.args.get('date_to', '')
    group_id = request.args.get('group', '')
    report_mode = request.args.get('report', '')

    # Item search
    item_q = VoorraadItem.query
    if q_text:
        item_q = item_q.filter(or_(
            VoorraadItem.naam.ilike(f'%{q_text}%'),
            VoorraadItem.description.ilike(f'%{q_text}%'),
            VoorraadItem.supplier_part_number.ilike(f'%{q_text}%')
        ))
    if spn:
        item_q = item_q.filter(VoorraadItem.supplier_part_number.ilike(f'%{spn}%'))
    if locatie:
        item_q = item_q.filter(VoorraadItem.locatie.ilike(f'%{locatie}%'))
    if categorie:
        item_q = item_q.filter_by(categorie=categorie)
    if group_id:
        item_q = item_q.filter_by(group_id=int(group_id))
    items = item_q.order_by(VoorraadItem.naam).all()

    # Movement search
    mov_q = VoorraadMutatie.query.join(VoorraadItem)
    if q_text:
        mov_q = mov_q.filter(or_(
            VoorraadItem.naam.ilike(f'%{q_text}%'),
            VoorraadItem.supplier_part_number.ilike(f'%{q_text}%'),
            VoorraadMutatie.opmerking.ilike(f'%{q_text}%')
        ))
    if spn:
        mov_q = mov_q.filter(VoorraadItem.supplier_part_number.ilike(f'%{spn}%'))
    if locatie:
        mov_q = mov_q.filter(VoorraadItem.locatie.ilike(f'%{locatie}%'))
    if categorie:
        mov_q = mov_q.filter(VoorraadItem.categorie == categorie)
    if group_id:
        mov_q = mov_q.filter(VoorraadItem.group_id == int(group_id))
    if move_type:
        mov_q = mov_q.filter(VoorraadMutatie.type == move_type)
    if date_from:
        try:
            dt = datetime.strptime(date_from, '%Y-%m-%d')
            mov_q = mov_q.filter(VoorraadMutatie.aangemaakt >= dt)
        except ValueError:
            flash(_('Invalid date format'), 'error')
    if date_to:
        try:
            dt = datetime.strptime(date_to, '%Y-%m-%d') + timedelta(days=1)
            mov_q = mov_q.filter(VoorraadMutatie.aangemaakt < dt)
        except ValueError:
            flash(_('Invalid date format'), 'error')
    movements = mov_q.order_by(VoorraadMutatie.aangemaakt.desc()).limit(500).all()

    # Report data
    total_in = sum(m.hoeveelheid for m in movements if m.type == 'inkomend')
    total_out = sum(m.hoeveelheid for m in movements if m.type == 'uitgaand')
    total_in_value = sum(m.hoeveelheid * m.item.prijs for m in movements if m.type == 'inkomend')
    total_out_value = sum(m.hoeveelheid * m.item.prijs for m in movements if m.type == 'uitgaand')

    # Grouped report by item
    item_report = {}
    for m in movements:
        key = m.item_id
        if key not in item_report:
            item_report[key] = {'item': m.item, 'in_qty': 0, 'out_qty': 0, 'in_val': 0, 'out_val': 0}
        if m.type == 'inkomend':
            item_report[key]['in_qty'] += m.hoeveelheid
            item_report[key]['in_val'] += m.hoeveelheid * m.item.prijs
        else:
            item_report[key]['out_qty'] += m.hoeveelheid
            item_report[key]['out_val'] += m.hoeveelheid * m.item.prijs
    item_report = sorted(item_report.values(), key=lambda x: x['item'].naam)

    cats = [c[0] for c in db.session.query(VoorraadItem.categorie).distinct().all() if c[0]]
    groups = WarehouseGroup.query.order_by(WarehouseGroup.name).all()

    return render_template('warehouse_search.html',
        items=items, movements=movements, item_report=item_report,
        total_in=total_in, total_out=total_out,
        total_in_value=total_in_value, total_out_value=total_out_value,
        cats=cats, groups=groups,
        f_q=q_text, f_spn=spn, f_locatie=locatie, f_categorie=categorie,
        f_move_type=move_type, f_date_from=date_from, f_date_to=date_to,
        f_group=group_id, report_mode=report_mode)

@app.route('/warehouse/labels')
@login_required
def warehouse_labels():
    ids = request.args.get('ids', '')
    if ids:
        item_ids = [int(x) for x in ids.split(',') if x.strip()]
        items = VoorraadItem.query.filter(VoorraadItem.id.in_(item_ids)).all()
    else:
        items = VoorraadItem.query.order_by(VoorraadItem.naam).all()
    return render_template('warehouse_labels.html', items=items)

@app.route('/api/warehouse/scan', methods=['POST'])
@login_required
def warehouse_scan():
    data = request.get_json()
    code = data.get('code', '')
    try:
        parsed = json.loads(code)
        if parsed.get('type') == 'warehouse_item':
            item = VoorraadItem.query.get(parsed.get('id'))
            if item:
                return jsonify({
                    'found': True, 'id': item.id, 'name': item.naam,
                    'category': item.categorie, 'quantity': item.hoeveelheid,
                    'unit': item.eenheid, 'price': item.prijs,
                    'available': item.hoeveelheid > 0
                })
    except (json.JSONDecodeError, AttributeError):
        pass
    # Search by supplier part number (barcode), then by name
    item = VoorraadItem.query.filter(
        db.or_(
            VoorraadItem.supplier_part_number == code,
            VoorraadItem.naam.ilike(f'%{code}%')
        )
    ).first()
    if item:
        return jsonify({
            'found': True, 'id': item.id, 'name': item.naam,
            'category': item.categorie, 'quantity': item.hoeveelheid,
            'unit': item.eenheid, 'price': item.prijs,
            'available': item.hoeveelheid > 0
        })
    return jsonify({'found': False, 'code': code})

# ============================================================
# ROUTES — PURCHASE REQUESTS
# ============================================================

@app.route('/purchase-requests')
@login_required
@role_required('admin', 'director', 'technician')
def purchase_requests_list():
    if current_user.has_role('admin', 'director'):
        requests = PurchaseRequest.query.order_by(PurchaseRequest.created_at.desc()).all()
    else:
        requests = PurchaseRequest.query.filter_by(requester_id=current_user.id).order_by(PurchaseRequest.created_at.desc()).all()
    return render_template('purchase_requests.html', requests=requests)

@app.route('/purchase-requests/new', methods=['GET', 'POST'])
@login_required
@role_required('admin', 'director', 'technician')
def purchase_request_new():
    if request.method == 'POST':
        pr = PurchaseRequest(
            fault_id=request.form.get('fault_id') or None,
            machine_id=int(request.form['machine_id']),
            requester_id=current_user.id,
            machine_serial=request.form.get('machine_serial', ''),
            fault_number=request.form.get('fault_number', ''),
            fault_description=request.form.get('fault_description', ''),
            part_name=request.form['part_name'],
            part_catalog=request.form.get('part_catalog', ''),
            quantity=float(request.form.get('quantity', 1)),
            unit=request.form.get('unit', 'st'),
            urgency=request.form.get('urgency', 'normal'),
            reason=request.form.get('reason', '')
        )
        db.session.add(pr)
        db.session.commit()
        
        log_audit('create', 'purchase_request', pr.id, f'{pr.part_name} x{pr.quantity} — {pr.machine.name} (срочность: {pr.urgency})')
        add_work_report(f'🛒 Новая заявка: {pr.part_name} x{pr.quantity} — {pr.machine.name} (срочность: {pr.urgency})')
        
        admins = User.query.filter(User.role.in_(['admin', 'director']), User.is_active_user == True).all()
        for admin in admins:
            create_notification(
                admin.id,
                _('New purchase request'),
                f"{pr.part_name} x{pr.quantity} — {pr.machine.name}",
                'fault',
                url_for('purchase_request_detail', request_id=pr.id)
            )
        
        flash(_('Purchase request created'), 'success')
        return redirect(url_for('purchase_requests_list'))
    
    if current_user.has_role('admin', 'director', 'technician'):
        machines = Machine.query.all()
    else:
        machines = current_user.assigned_machines
    faults = FaultReport.query.filter(
        (FaultReport.reporter_id == current_user.id) | (FaultReport.technician_id == current_user.id)
    ).order_by(FaultReport.created_at.desc()).limit(20).all()
    return render_template('purchase_request_form.html', machines=machines, faults=faults)

@app.route('/purchase-requests/<int:request_id>')
@login_required
@role_required('admin', 'director', 'technician')
def purchase_request_detail(request_id):
    pr = PurchaseRequest.query.get_or_404(request_id)
    return render_template('purchase_request_detail.html', pr=pr)

@app.route('/purchase-requests/<int:request_id>/approve', methods=['POST'])
@login_required
@role_required('admin', 'director')
def purchase_request_approve(request_id):
    pr = PurchaseRequest.query.get_or_404(request_id)
    pr.status = 'approved'
    pr.reviewed_at = datetime.utcnow()
    pr.reviewer_id = current_user.id
    db.session.commit()
    
    log_audit('approve', 'purchase_request', pr.id, f'{pr.part_name} x{pr.quantity} — {pr.machine.name}')
    add_work_report(f'✅ Заявка одобрена: {pr.part_name} x{pr.quantity} — {pr.machine.name}')
    
    create_notification(
        pr.requester_id,
        _('Purchase request approved'),
        f"{pr.part_name} x{pr.quantity} — {pr.machine.name}",
        'info',
        url_for('purchase_request_detail', request_id=pr.id)
    )
    
    flash(_('Purchase request approved'), 'success')
    return redirect(url_for('purchase_request_detail', request_id=pr.id))

@app.route('/purchase-requests/<int:request_id>/reject', methods=['POST'])
@login_required
@role_required('admin', 'director')
def purchase_request_reject(request_id):
    pr = PurchaseRequest.query.get_or_404(request_id)
    pr.status = 'rejected'
    pr.reviewed_at = datetime.utcnow()
    pr.reviewer_id = current_user.id
    db.session.commit()
    
    log_audit('reject', 'purchase_request', pr.id, f'{pr.part_name} x{pr.quantity} — {pr.machine.name}')
    add_work_report(f'❌ Заявка отклонена: {pr.part_name} x{pr.quantity} — {pr.machine.name}')
    
    create_notification(
        pr.requester_id,
        _('Purchase request rejected'),
        f"{pr.part_name} x{pr.quantity} — {pr.machine.name}",
        'warning',
        url_for('purchase_request_detail', request_id=pr.id)
    )
    
    flash(_('Purchase request rejected'), 'error')
    return redirect(url_for('purchase_request_detail', request_id=pr.id))

# ============================================================
# ROUTES — WORK SCHEDULE & TIME TRACKING
# ============================================================

@app.route('/schedule')
@login_required
@role_required('admin', 'director', 'technician')
def schedule_list():
    if current_user.has_role('admin', 'director'):
        users = User.query.filter(User.is_active_user == True, User.role.in_(['technician', 'user'])).all()
    else:
        users = [current_user]
    schedules = WorkSchedule.query.filter(WorkSchedule.user_id.in_([u.id for u in users])).all()
    return render_template('schedule.html', users=users, schedules=schedules)

@app.route('/schedule/<int:user_id>', methods=['GET', 'POST'])
@login_required
def schedule_user(user_id):
    if not current_user.has_role('admin', 'director') and current_user.id != user_id:
        flash(_('Access denied'), 'error')
        return redirect(url_for('schedule_list'))
    user = User.query.get_or_404(user_id)
    if request.method == 'POST':
        work_days = ','.join(request.form.getlist('work_days'))
        s = WorkSchedule(
            user_id=user.id,
            name=request.form['name'],
            shift_start=request.form['shift_start'],
            shift_end=request.form['shift_end'],
            break_minutes=int(request.form.get('break_minutes', 60)),
            work_days=work_days or '1,2,3,4,5'
        )
        db.session.add(s)
        db.session.commit()
        flash(_('Schedule created'), 'success')
    schedules = WorkSchedule.query.filter_by(user_id=user.id).all()
    return render_template('schedule_user.html', user=user, schedules=schedules)

@app.route('/schedule/<int:user_id>/delete', methods=['POST'])
@login_required
@role_required('admin')
def schedule_delete(user_id):
    user = User.query.get_or_404(user_id)
    deleted = WorkSchedule.query.filter_by(user_id=user.id).delete()
    db.session.commit()
    flash(_('Schedule deleted for') + ' ' + (user.display_name or user.username) + f' ({deleted})', 'success')
    return redirect(url_for('schedule_list'))

# Belgian public holidays (fixed + Easter-based)
BELGIAN_HOLIDAYS_FIXED = {
    (1, 1): "Nieuwjaar / Jour de l'An",
    (5, 1): "Dag van de Arbeid / Fête du Travail",
    (7, 21): "Nationale Feestdag / Fête nationale",
    (8, 15): "O-L-V-Hemelvaart / Assomption",
    (11, 1): "Allerheiligen / Toussaint",
    (11, 11): "Wapenstilstand / Armistice",
    (12, 25): "Kerstmis / Noël",
}

def get_easter(year):
    """Calculate Easter Sunday using Meeus/Jones/Butcher algorithm."""
    a = year % 19
    b = year // 100
    c = year % 100
    d = b // 4
    e = b % 4
    f = (b + 8) // 25
    g = (b - f + 1) // 3
    h = (19 * a + b - d - g + 15) % 30
    i = c // 4
    k = c % 4
    l = (32 + 2 * e + 2 * i - h - k) % 7
    m = (a + 11 * h + 22 * l) // 451
    month = (h + l - 7 * m + 114) // 31
    day = ((h + l - 7 * m + 114) % 31) + 1
    return datetime(year, month, day).date()

def get_belgian_holidays(year):
    """Return dict of {date: name} for Belgian public holidays."""
    holidays = {}
    for (m, d), name in BELGIAN_HOLIDAYS_FIXED.items():
        holidays[datetime(year, m, d).date()] = name
    easter = get_easter(year)
    holidays[easter - timedelta(days=2)] = "Goede Vrijdag / Vendredi saint"
    holidays[easter] = "Pasen / Pâques"
    holidays[easter + timedelta(days=1)] = "Paasmaandag / Lundi de Pâques"
    holidays[easter + timedelta(days=39)] = "Hemelvaart / Ascension"
    holidays[easter + timedelta(days=50)] = "Pinkstermaandag / Lundi de Pentecôte"
    return holidays

@app.route('/schedule/monthly')
@login_required
@role_required('admin', 'director')
def schedule_monthly():
    year = int(request.args.get('year', datetime.utcnow().year))
    month = int(request.args.get('month', datetime.utcnow().month))
    filter_user = request.args.get('user', '')
    if month < 1: month = 12; year -= 1
    if month > 12: month = 1; year += 1

    all_users = User.query.filter(User.is_active_user == True, User.role.in_(['technician'])).order_by(User.display_name).all()
    if filter_user:
        users = [u for u in all_users if str(u.id) == filter_user]
    else:
        users = all_users
    first_day = datetime(year, month, 1).date()
    if month == 12:
        last_day = datetime(year + 1, 1, 1).date() - timedelta(days=1)
    else:
        last_day = datetime(year, month + 1, 1).date() - timedelta(days=1)

    # Get all weekend shifts for this month
    shifts = WeekendShift.query.filter(
        WeekendShift.date >= first_day,
        WeekendShift.date <= last_day
    ).all()
    shift_map = {}
    for s in shifts:
        key = (s.user_id, s.date)
        shift_map[key] = s

    # Get work schedules for each user
    user_schedules = {}
    all_schedules = WorkSchedule.query.filter(WorkSchedule.user_id.in_([u.id for u in users]), WorkSchedule.is_active == True).all()
    for s in all_schedules:
        if s.user_id not in user_schedules:
            user_schedules[s.user_id] = s

    # Get Belgian holidays
    holidays = get_belgian_holidays(year)

    # Build days list
    days = []
    current = first_day
    while current <= last_day:
        days.append({
            'date': current,
            'day': current.day,
            'weekday': current.weekday(),  # 0=Mon, 6=Sun
            'is_weekend': current.weekday() >= 5,
            'is_holiday': current in holidays,
            'holiday_name': holidays.get(current, ''),
        })
        current += timedelta(days=1)

    return render_template('schedule_monthly.html',
        users=users, days=days, year=year, month=month,
        shift_map=shift_map, holidays=holidays, user_schedules=user_schedules,
        all_users=all_users, filter_user=filter_user)

@app.route('/schedule/monthly/shift', methods=['POST'])
@login_required
@role_required('admin', 'director')
def schedule_monthly_shift():
    data = request.get_json()
    user_id = data.get('user_id')
    date_str = data.get('date')
    action = data.get('action')  # 'add' or 'remove'
    shift_type = data.get('shift_type', 'full')

    date = datetime.strptime(date_str, '%Y-%m-%d').date()
    existing = WeekendShift.query.filter_by(user_id=user_id, date=date).first()

    if action == 'add':
        if existing:
            existing.shift_type = shift_type
        else:
            s = WeekendShift(user_id=user_id, date=date, shift_type=shift_type, created_by=current_user.id)
            db.session.add(s)
    elif action == 'remove' and existing:
        db.session.delete(existing)

    db.session.commit()
    return jsonify({'ok': True})

@app.route('/schedule/monthly/delete', methods=['POST'])
@login_required
@role_required('admin', 'director')
def schedule_monthly_delete():
    data = request.get_json()
    user_id = data.get('user_id')
    year = data.get('year')
    month = data.get('month')
    first_day = datetime(year, month, 1).date()
    if month == 12:
        last_day = datetime(year + 1, 1, 1).date() - timedelta(days=1)
    else:
        last_day = datetime(year, month + 1, 1).date() - timedelta(days=1)
    deleted = WeekendShift.query.filter(
        WeekendShift.user_id == user_id,
        WeekendShift.date >= first_day,
        WeekendShift.date <= last_day
    ).delete()
    db.session.commit()
    return jsonify({'ok': True, 'deleted': deleted})

@app.route('/time-tracking')
@login_required
def time_tracking():
    if current_user.has_role('admin', 'director'):
        users = User.query.filter(User.is_active_user == True, User.role.in_(['technician', 'user'])).all()
    else:
        users = [current_user]
    
    today = datetime.utcnow().date()
    month_start = today.replace(day=1)
    
    entries = TimeEntry.query.filter(
        TimeEntry.user_id.in_([u.id for u in users]),
        TimeEntry.date >= month_start,
        TimeEntry.date <= today
    ).order_by(TimeEntry.date.desc()).all()
    
    return render_template('time_tracking.html', users=users, entries=entries, today=today, month_start=month_start)

@app.route('/time-tracking/clock-in', methods=['POST'])
@login_required
def clock_in():
    today = datetime.utcnow().date()
    existing = TimeEntry.query.filter_by(user_id=current_user.id, date=today).first()
    if existing and existing.clock_in:
        flash(_('Already clocked in today'), 'error')
        return redirect(url_for('time_tracking'))
    
    if existing:
        existing.clock_in = datetime.utcnow()
        existing.status = 'present'
    else:
        entry = TimeEntry(
            user_id=current_user.id,
            date=today,
            clock_in=datetime.utcnow(),
            status='present'
        )
        db.session.add(entry)
    db.session.commit()
    flash(_('Clocked in at') + ' ' + datetime.utcnow().strftime('%H:%M'), 'success')
    return redirect(url_for('time_tracking'))

@app.route('/time-tracking/clock-out', methods=['POST'])
@login_required
def clock_out():
    today = datetime.utcnow().date()
    entry = TimeEntry.query.filter_by(user_id=current_user.id, date=today).first()
    if not entry or not entry.clock_in:
        flash(_('Not clocked in today'), 'error')
        return redirect(url_for('time_tracking'))
    if entry.clock_out:
        flash(_('Already clocked out today'), 'error')
        return redirect(url_for('time_tracking'))
    
    entry.clock_out = datetime.utcnow()
    delta = entry.clock_out - entry.clock_in
    hours = delta.total_seconds() / 3600
    entry.hours_worked = round(hours - (entry.break_minutes / 60), 2)
    
    # Calculate overtime (standard 8h)
    if entry.hours_worked > 8:
        entry.overtime_hours = round(entry.hours_worked - 8, 2)
    
    db.session.commit()
    flash(_('Clocked out at') + ' ' + entry.clock_out.strftime('%H:%M') + '. ' + _('Hours worked') + ': ' + str(entry.hours_worked), 'success')
    return redirect(url_for('time_tracking'))

@app.route('/time-tracking/manual', methods=['POST'])
@login_required
@role_required('admin', 'director')
def time_tracking_manual():
    user_id = int(request.form['user_id'])
    date = datetime.strptime(request.form['date'], '%Y-%m-%d').date()
    status = request.form.get('status', 'present')
    
    entry = TimeEntry.query.filter_by(user_id=user_id, date=date).first()
    if not entry:
        entry = TimeEntry(user_id=user_id, date=date, status=status)
        db.session.add(entry)
    
    entry.status = status
    entry.notes = request.form.get('notes', '')
    
    if status == 'present':
        entry.clock_in = datetime.combine(date, datetime.strptime(request.form['clock_in'], '%H:%M').time())
        entry.clock_out = datetime.combine(date, datetime.strptime(request.form['clock_out'], '%H:%M').time())
        delta = entry.clock_out - entry.clock_in
        hours = delta.total_seconds() / 3600
        entry.break_minutes = int(request.form.get('break_minutes', 60))
        entry.hours_worked = round(hours - (entry.break_minutes / 60), 2)
        if entry.hours_worked > 8:
            entry.overtime_hours = round(entry.hours_worked - 8, 2)
    
    db.session.commit()
    flash(_('Time entry saved'), 'success')
    return redirect(url_for('time_tracking'))

@app.route('/vacations')
@login_required
@role_required('admin', 'director', 'technician')
def vacations_list():
    if current_user.has_role('admin', 'director'):
        vacations = Vacation.query.order_by(Vacation.created_at.desc()).all()
    else:
        vacations = Vacation.query.filter_by(user_id=current_user.id).order_by(Vacation.created_at.desc()).all()
    return render_template('vacations.html', vacations=vacations)

@app.route('/vacations/new', methods=['GET', 'POST'])
@login_required
@role_required('admin', 'director', 'technician')
def vacation_new():
    if request.method == 'POST':
        d_from = datetime.strptime(request.form['date_from'], '%Y-%m-%d').date()
        d_to = datetime.strptime(request.form['date_to'], '%Y-%m-%d').date()
        days = (d_to - d_from).days + 1
        v = Vacation(
            user_id=current_user.id,
            vacation_type=request.form['vacation_type'],
            date_from=d_from,
            date_to=d_to,
            days_count=days,
            reason=request.form.get('reason', '')
        )
        db.session.add(v)
        db.session.commit()
        
        admins = User.query.filter(User.role.in_(['admin', 'director']), User.is_active_user == True).all()
        for admin in admins:
            create_notification(
                admin.id,
                _('New vacation request'),
                f"{current_user.display_name}: {v.vacation_type} {v.date_from} - {v.date_to}",
                'info',
                url_for('vacations_list')
            )
        
        flash(_('Vacation request submitted'), 'success')
        return redirect(url_for('vacations_list'))
    return render_template('vacation_form.html')

@app.route('/vacations/<int:vacation_id>/approve', methods=['POST'])
@login_required
@role_required('admin', 'director')
def vacation_approve(vacation_id):
    v = Vacation.query.get_or_404(vacation_id)
    v.status = 'approved'
    v.approved_by = current_user.id
    db.session.commit()
    create_notification(v.user_id, _('Vacation approved'), f"{v.vacation_type} {v.date_from} - {v.date_to}", 'info')
    flash(_('Vacation approved'), 'success')
    return redirect(url_for('vacations_list'))

@app.route('/vacations/<int:vacation_id>/reject', methods=['POST'])
@login_required
@role_required('admin', 'director')
def vacation_reject(vacation_id):
    v = Vacation.query.get_or_404(vacation_id)
    v.status = 'rejected'
    v.approved_by = current_user.id
    db.session.commit()
    create_notification(v.user_id, _('Vacation rejected'), f"{v.vacation_type} {v.date_from} - {v.date_to}", 'warning')
    flash(_('Vacation rejected'), 'error')
    return redirect(url_for('vacations_list'))

@app.route('/time-report/<int:user_id>')
@login_required
def time_report(user_id):
    if not current_user.has_role('admin', 'director') and current_user.id != user_id:
        flash(_('Access denied'), 'error')
        return redirect(url_for('time_tracking'))
    
    user = User.query.get_or_404(user_id)
    month = request.args.get('month', datetime.utcnow().strftime('%Y-%m'))
    year, mon = map(int, month.split('-'))
    start = datetime(year, mon, 1).date()
    if mon == 12:
        end = datetime(year + 1, 1, 1).date()
    else:
        end = datetime(year, mon + 1, 1).date()
    
    entries = TimeEntry.query.filter(
        TimeEntry.user_id == user_id,
        TimeEntry.date >= start,
        TimeEntry.date < end
    ).order_by(TimeEntry.date).all()
    
    total_hours = sum(e.hours_worked for e in entries)
    total_overtime = sum(e.overtime_hours for e in entries)
    days_present = len([e for e in entries if e.status == 'present'])
    days_absent = len([e for e in entries if e.status in ['absent', 'sick']])
    
    vacations = Vacation.query.filter(
        Vacation.user_id == user_id,
        Vacation.status == 'approved',
        Vacation.date_from < end,
        Vacation.date_to >= start
    ).all()
    
    return render_template('time_report.html', user=user, entries=entries, month=month,
                         total_hours=total_hours, total_overtime=total_overtime,
                         days_present=days_present, days_absent=days_absent, vacations=vacations)

# ============================================================
# ROUTES — TRANSLATION API
# ============================================================

@app.route('/api/translate', methods=['POST'])
@login_required
def api_translate():
    data = request.get_json()
    if not data or 'text' not in data:
        return jsonify({'error': 'No text provided'}), 400
    text = data['text']
    target = data.get('target', g.lang)
    translated = translate_text(text, target)
    return jsonify({'translated': translated, 'original': text})

# ============================================================
# AUTOMATION — AUTO-NOTIFICATIONS & REMINDERS
# ============================================================

@app.route('/api/automation/check', methods=['POST'])
@login_required
@role_required('admin')
def automation_check():
    """Run all automated checks and send notifications"""
    today = datetime.utcnow().date()
    soon = today + timedelta(days=14)
    results = {'maintenance': 0, 'low_stock': 0, 'contracts': 0, 'cylinders': 0}

    # Check maintenance due
    parts = MachinePart.query.filter(MachinePart.next_replacement <= soon).all()
    for p in parts:
        if p.responsible_user_id:
            days = (p.next_replacement - today).days
            title = f"{'OVERDUE' if days < 0 else 'Upcoming'} replacement: {p.name}"
            msg = f"{p.machine.name}: {p.name} - {'overdue by ' + str(-days) + ' days' if days < 0 else 'due in ' + str(days) + ' days'}"
            existing = Notification.query.filter_by(user_id=p.responsible_user_id, is_read=False, title=title).first()
            if not existing:
                create_notification(p.responsible_user_id, title, msg, 'warning', url_for('machine_parts', machine_id=p.machine_id))
                results['maintenance'] += 1

    # Check low stock
    low_items = VoorraadItem.query.filter(VoorraadItem.hoeveelheid <= VoorraadItem.minimum).all()
    if low_items:
        admins = User.query.filter(User.role.in_(['admin', 'director']), User.is_active_user == True).all()
        for admin in admins:
            title = f"Low stock: {len(low_items)} items"
            msg = ', '.join([f"{i.naam} ({i.hoeveelheid}/{i.minimum})" for i in low_items[:5]])
            existing = Notification.query.filter_by(user_id=admin.id, is_read=False, title=title).first()
            if not existing:
                create_notification(admin.id, title, msg, 'warning', url_for('warehouse_list'))
                results['low_stock'] += 1

    # Check expiring contracts
    contractors = Contractor.query.filter(Contractor.contract_end <= soon, Contractor.contract_end >= today, Contractor.is_active == True).all()
    for c in contractors:
        days = (c.contract_end - today).days
        admins = User.query.filter(User.role.in_(['admin', 'director']), User.is_active_user == True).all()
        for admin in admins:
            title = f"Contract expiring: {c.company_name}"
            msg = f"Contract with {c.company_name} expires in {days} days"
            existing = Notification.query.filter_by(user_id=admin.id, is_read=False, title=title).first()
            if not existing:
                create_notification(admin.id, title, msg, 'warning', url_for('contractor_detail', contractor_id=c.id))
                results['contracts'] += 1

    # Check cylinder stock
    n2_total = GasCylinder.query.filter(GasCylinder.gas_type == 'nitrogen', GasCylinder.status.in_(['full', 'in_use'])).count()
    co2_total = GasCylinder.query.filter(GasCylinder.gas_type == 'co2', GasCylinder.status.in_(['full', 'in_use'])).count()
    admins = User.query.filter(User.role.in_(['admin', 'director']), User.is_active_user == True).all()
    if n2_total <= 3:
        for admin in admins:
            existing = Notification.query.filter_by(user_id=admin.id, is_read=False, title='Low stock: N₂').first()
            if not existing:
                create_notification(admin.id, 'Low stock: N₂', f'Only {n2_total} nitrogen cylinders remaining', 'warning', url_for('cylinders_dashboard'))
                results['cylinders'] += 1
    if co2_total <= 1:
        for admin in admins:
            existing = Notification.query.filter_by(user_id=admin.id, is_read=False, title='Low stock: CO₂').first()
            if not existing:
                create_notification(admin.id, 'Low stock: CO₂', f'Only {co2_total} CO2 cylinders remaining', 'warning', url_for('cylinders_dashboard'))
                results['cylinders'] += 1

    log_audit('automation_check', details=json.dumps(results))
    return jsonify({'ok': True, 'results': results})

# ============================================================
# EXPORT — CSV/EXCEL REPORTS
# ============================================================

@app.route('/export/machines')
@login_required
@role_required('admin', 'director')
def export_machines():
    import csv
    import io
    machines = Machine.query.all()
    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow(['ID', 'Name', 'Type', 'Manufacturer', 'Serial', 'Status', 'Section', 'Contractor'])
    for m in machines:
        writer.writerow([m.id, m.name, m.machine_type, m.manufacturer, m.serial_number,
                         m.status, m.section.name if m.section else '',
                         m.contractor_rel.company_name if m.contractor_rel else ''])
    output.seek(0)
    return send_file(io.BytesIO(output.getvalue().encode('utf-8-sig')),
                     mimetype='text/csv', as_attachment=True,
                     download_name=f'machines_{datetime.utcnow().strftime("%Y%m%d")}.csv')

@app.route('/export/warehouse')
@login_required
@role_required('admin', 'director')
def export_warehouse():
    import csv
    import io
    items = VoorraadItem.query.order_by(VoorraadItem.naam).all()
    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow(['ID', 'Name', 'Category', 'Group', 'Quantity', 'Unit', 'Min', 'Price', 'Location'])
    for i in items:
        writer.writerow([i.id, i.naam, i.categorie, i.group.name if i.group else '',
                         i.hoeveelheid, i.eenheid, i.minimum, i.prijs, i.locatie])
    output.seek(0)
    return send_file(io.BytesIO(output.getvalue().encode('utf-8-sig')),
                     mimetype='text/csv', as_attachment=True,
                     download_name=f'warehouse_{datetime.utcnow().strftime("%Y%m%d")}.csv')

@app.route('/export/movements')
@login_required
@role_required('admin', 'director')
def export_movements():
    import csv
    import io
    item_id = request.args.get('item', '')
    move_type = request.args.get('type', '')
    q = VoorraadMutatie.query
    if item_id:
        q = q.filter_by(item_id=int(item_id))
    if move_type:
        q = q.filter_by(type=move_type)
    movements = q.order_by(VoorraadMutatie.aangemaakt.desc()).all()
    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow(['Date', 'Type', 'Item', 'Quantity', 'Unit', 'Order', 'Comment'])
    for m in movements:
        writer.writerow([
            m.aangemaakt.strftime('%Y-%m-%d %H:%M'),
            m.type,
            m.item.naam if m.item else '',
            m.hoeveelheid,
            m.item.eenheid if m.item else '',
            m.opdracht.nummer if m.opdracht else '',
            m.opmerking or ''
        ])
    output.seek(0)
    return send_file(io.BytesIO(output.getvalue().encode('utf-8-sig')),
                     mimetype='text/csv', as_attachment=True,
                     download_name=f'movements_{datetime.utcnow().strftime("%Y%m%d")}.csv')

@app.route('/export/maintenance')
@login_required
@role_required('admin', 'director')
def export_maintenance():
    import csv
    import io
    records = MaintenanceRecord.query.order_by(MaintenanceRecord.date_performed.desc()).all()
    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow(['ID', 'Machine', 'Type', 'Description', 'Date', 'Cost', 'Next', 'Performed By'])
    for r in records:
        writer.writerow([r.id, r.machine.name, r.maintenance_type, r.description,
                         r.date_performed.strftime('%Y-%m-%d'), r.cost,
                         r.next_maintenance.strftime('%Y-%m-%d') if r.next_maintenance else '',
                         r.performer.display_name if r.performer else ''])
    output.seek(0)
    return send_file(io.BytesIO(output.getvalue().encode('utf-8-sig')),
                     mimetype='text/csv', as_attachment=True,
                     download_name=f'maintenance_{datetime.utcnow().strftime("%Y%m%d")}.csv')

@app.route('/export/faults')
@login_required
@role_required('admin', 'director')
def export_faults():
    import csv
    import io
    faults = FaultReport.query.order_by(FaultReport.created_at.desc()).all()
    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow(['ID', 'Title', 'Machine', 'Priority', 'Status', 'Reporter', 'Technician', 'Created', 'Resolved'])
    for f in faults:
        tech = User.query.get(f.technician_id) if f.technician_id else None
        writer.writerow([f.id, f.title, f.machine.name, f.priority, f.status,
                         f.reporter.display_name if f.reporter else '',
                         tech.display_name if tech else '',
                         f.created_at.strftime('%Y-%m-%d'),
                         f.resolved_at.strftime('%Y-%m-%d') if f.resolved_at else ''])
    output.seek(0)
    return send_file(io.BytesIO(output.getvalue().encode('utf-8-sig')),
                     mimetype='text/csv', as_attachment=True,
                     download_name=f'faults_{datetime.utcnow().strftime("%Y%m%d")}.csv')

# ============================================================
# CONTRACTORS ROUTES
# ============================================================

@app.route('/contractors')
@login_required
@role_required('admin', 'director')
def contractors_list():
    contractors = Contractor.query.filter_by(is_active=True).order_by(Contractor.company_name).all()
    return render_template('contractors.html', contractors=contractors)

@app.route('/contractors/new', methods=['GET', 'POST'])
@login_required
@role_required('admin')
def contractor_new():
    if request.method == 'POST':
        c = Contractor(
            company_name=request.form['company_name'],
            contact_person=request.form.get('contact_person', ''),
            contact_position=request.form.get('contact_position', ''),
            phone=request.form.get('phone', ''),
            phone2=request.form.get('phone2', ''),
            email=request.form.get('email', ''),
            website=request.form.get('website', ''),
            address=request.form.get('address', ''),
            postcode=request.form.get('postcode', ''),
            city=request.form.get('city', ''),
            country=request.form.get('country', 'Nederland'),
            kvk_number=request.form.get('kvk_number', ''),
            btw_number=request.form.get('btw_number', ''),
            iban=request.form.get('iban', ''),
            service_type=request.form.get('service_type', ''),
            contract_number=request.form.get('contract_number', ''),
            contract_start=datetime.strptime(request.form['contract_start'], '%Y-%m-%d').date() if request.form.get('contract_start') else None,
            contract_end=datetime.strptime(request.form['contract_end'], '%Y-%m-%d').date() if request.form.get('contract_end') else None,
            notes=request.form.get('notes', '')
        )
        db.session.add(c)
        db.session.commit()
        log_audit('create', 'contractor', c.id, c.company_name)
        flash(_('Contractor added'), 'success')
        return redirect(url_for('contractor_detail', contractor_id=c.id))
    return render_template('contractor_form.html', contractor=None)

@app.route('/contractors/<int:contractor_id>')
@login_required
@role_required('admin', 'director')
def contractor_detail(contractor_id):
    c = Contractor.query.get_or_404(contractor_id)
    today = datetime.utcnow().date()
    return render_template('contractor_detail.html', contractor=c, today=today)

@app.route('/contractors/<int:contractor_id>/edit', methods=['GET', 'POST'])
@login_required
@role_required('admin')
def contractor_edit(contractor_id):
    c = Contractor.query.get_or_404(contractor_id)
    if request.method == 'POST':
        c.company_name = request.form['company_name']
        c.contact_person = request.form.get('contact_person', '')
        c.contact_position = request.form.get('contact_position', '')
        c.phone = request.form.get('phone', '')
        c.phone2 = request.form.get('phone2', '')
        c.email = request.form.get('email', '')
        c.website = request.form.get('website', '')
        c.address = request.form.get('address', '')
        c.postcode = request.form.get('postcode', '')
        c.city = request.form.get('city', '')
        c.country = request.form.get('country', 'Nederland')
        c.kvk_number = request.form.get('kvk_number', '')
        c.btw_number = request.form.get('btw_number', '')
        c.iban = request.form.get('iban', '')
        c.service_type = request.form.get('service_type', '')
        c.contract_number = request.form.get('contract_number', '')
        c.contract_start = datetime.strptime(request.form['contract_start'], '%Y-%m-%d').date() if request.form.get('contract_start') else None
        c.contract_end = datetime.strptime(request.form['contract_end'], '%Y-%m-%d').date() if request.form.get('contract_end') else None
        c.notes = request.form.get('notes', '')
        db.session.commit()
        log_audit('update', 'contractor', c.id, c.company_name)
        flash(_('Contractor updated'), 'success')
        return redirect(url_for('contractor_detail', contractor_id=c.id))
    return render_template('contractor_form.html', contractor=c)

@app.route('/contractors/<int:contractor_id>/delete', methods=['POST'])
@login_required
@role_required('admin')
def contractor_delete(contractor_id):
    c = Contractor.query.get_or_404(contractor_id)
    c.is_active = False
    db.session.commit()
    log_audit('delete', 'contractor', c.id, c.company_name)
    flash(_('Contractor deactivated'), 'success')
    return redirect(url_for('contractors_list'))

@app.route('/contractors/<int:contractor_id>/add-employee', methods=['POST'])
@login_required
@role_required('admin')
def contractor_add_employee(contractor_id):
    c = Contractor.query.get_or_404(contractor_id)
    emp = ContractorEmployee(
        contractor_id=c.id,
        name=request.form['name'],
        position=request.form.get('position', ''),
        phone=request.form.get('phone', ''),
        email=request.form.get('email', ''),
        notes=request.form.get('notes', '')
    )
    db.session.add(emp)
    db.session.commit()
    flash(_('Employee added'), 'success')
    return redirect(url_for('contractor_detail', contractor_id=c.id))

@app.route('/contractors/employee/<int:emp_id>/delete', methods=['POST'])
@login_required
@role_required('admin')
def contractor_delete_employee(emp_id):
    emp = ContractorEmployee.query.get_or_404(emp_id)
    cid = emp.contractor_id
    db.session.delete(emp)
    db.session.commit()
    flash(_('Employee removed'), 'success')
    return redirect(url_for('contractor_detail', contractor_id=cid))

# ============================================================
# WORK REPORT — STORING
# ============================================================

@app.route('/work-report')
@login_required
@role_required('admin', 'director', 'technician')
def work_report_page():
    entries = WorkReportEntry.query.order_by(WorkReportEntry.created_at.desc()).limit(500).all()
    return render_template('work_report_page.html', entries=entries)

@app.route('/work-report/add', methods=['POST'])
@login_required
@role_required('admin', 'director', 'technician')
def work_report_add():
    entry_text = request.form.get('entry', '').strip()
    if not entry_text:
        flash(_('Entry cannot be empty'), 'error')
        return redirect(url_for('work_report_page'))
    entry = WorkReportEntry(
        user_id=current_user.id,
        entry=entry_text
    )
    db.session.add(entry)
    db.session.commit()
    flash(_('Entry added'), 'success')
    return redirect(url_for('work_report_page'))

@app.route('/work-report/delete/<int:entry_id>', methods=['POST'])
@login_required
@role_required('admin')
def work_report_delete(entry_id):
    entry = WorkReportEntry.query.get_or_404(entry_id)
    db.session.delete(entry)
    db.session.commit()
    flash(_('Entry deleted'), 'success')
    return redirect(url_for('work_report_page'))

def add_work_report(entry_text):
    """Helper function to add entry to Work Report from anywhere in the app"""
    try:
        entry = WorkReportEntry(
            user_id=current_user.id if current_user.is_authenticated else None,
            entry=entry_text
        )
        db.session.add(entry)
        db.session.commit()
    except Exception:
        db.session.rollback()

# ============================================================
# TOOL WEAR TRACKING
# ============================================================

@app.route('/knife-warning-preview')
@login_required
def knife_warning_preview():
    return render_template('knife_warning_preview.html')

@app.route('/tool-wear')
@login_required
def tool_wear_page():
    tools = ToolWear.query.order_by(ToolWear.machine_name).all()
    # Default machines if none exist
    default_machines = ['VULBUS 1', 'VULBUS 2', 'Lift Seydelmann', 'Станок BOLDT']
    if not tools:
        for m in default_machines:
            t = ToolWear(machine_name=m, tool_name='Ножи / Фреза', cycle_days=14, last_replaced=datetime.utcnow().date())
            db.session.add(t)
        db.session.commit()
        tools = ToolWear.query.order_by(ToolWear.machine_name).all()
    
    # Auto-calculate wear based on days since last replacement
    today = datetime.utcnow().date()
    for t in tools:
        cycle = t.cycle_days or 14
        if t.last_replaced:
            days = (today - t.last_replaced).days
            t.wear_percent = min(100.0, round((days / cycle) * 100, 1))
        else:
            t.wear_percent = 100.0
    
    return render_template('tool_wear.html', tools=tools, today=today)

@app.route('/tool-wear/add', methods=['POST'])
@login_required
@role_required('admin', 'technician')
def tool_wear_add():
    machine_name = request.form.get('machine_name', '').strip()
    tool_name = request.form.get('tool_name', '').strip() or 'Ножи / Фреза'
    cycle_days = int(request.form.get('cycle_days', 14))
    if machine_name:
        t = ToolWear(machine_name=machine_name, tool_name=tool_name, cycle_days=cycle_days, last_replaced=datetime.utcnow().date())
        db.session.add(t)
        db.session.commit()
        flash(_('Tool added'), 'success')
    return redirect(url_for('tool_wear_page'))

@app.route('/tool-wear/update/<int:tool_id>', methods=['POST'])
@login_required
@role_required('admin', 'technician')
def tool_wear_update(tool_id):
    tool = ToolWear.query.get_or_404(tool_id)
    tool.machine_name = request.form.get('machine_name', tool.machine_name).strip()
    tool.tool_name = request.form.get('tool_name', tool.tool_name).strip()
    tool.cycle_days = int(request.form.get('cycle_days', tool.cycle_days or 14))
    date_str = request.form.get('last_replaced')
    if date_str:
        tool.last_replaced = datetime.strptime(date_str, '%Y-%m-%d').date()
    tool.notes = request.form.get('notes', tool.notes)
    tool.updated_by = current_user.id
    db.session.commit()
    flash(_('Tool updated'), 'success')
    return redirect(url_for('tool_wear_page'))

@app.route('/tool-wear/delete/<int:tool_id>', methods=['POST'])
@login_required
@role_required('admin')
def tool_wear_delete(tool_id):
    tool = ToolWear.query.get_or_404(tool_id)
    db.session.delete(tool)
    db.session.commit()
    flash(_('Tool deleted'), 'success')
    return redirect(url_for('tool_wear_page'))

@app.route('/tool-wear/reset/<int:tool_id>', methods=['POST'])
@login_required
@role_required('admin', 'technician')
def tool_wear_reset(tool_id):
    tool = ToolWear.query.get_or_404(tool_id)
    tool.wear_percent = 0
    tool.last_replaced = datetime.utcnow().date()
    tool.updated_by = current_user.id
    db.session.commit()
    add_work_report(f'🔪 Замена инструмента: {tool.machine_name} — {tool.tool_name} (износ сброшен)')
    flash(_('Tool replaced, wear reset to 0%'), 'success')
    return redirect(url_for('tool_wear_page'))

@app.route('/api/tool-wear/warnings')
@login_required
def api_tool_wear_warnings():
    """Return tools with wear >= 80% for popup warning."""
    today = datetime.utcnow().date()
    all_tools = ToolWear.query.all()
    warnings = []
    for t in all_tools:
        cycle = t.cycle_days or 14
        if t.last_replaced:
            days = (today - t.last_replaced).days
            wear = min(100.0, round((days / cycle) * 100, 1))
        else:
            wear = 100.0
        if wear >= 80:
            warnings.append({
                'machine': t.machine_name,
                'tool': t.tool_name,
                'wear': wear,
                'last_replaced': t.last_replaced.strftime('%d.%m.%Y') if t.last_replaced else '—'
            })
    return jsonify({'warnings': warnings})

# ============================================================
# MONTHLY ARCHIVE
# ============================================================

@app.route('/archive')
@login_required
@role_required('admin', 'director', 'technician')
def archive_page():
    archives = MonthlyArchive.query.order_by(MonthlyArchive.archive_month.desc(), MonthlyArchive.section).all()
    # Group by month
    months = {}
    for a in archives:
        if a.archive_month not in months:
            months[a.archive_month] = []
        months[a.archive_month].append(a)
    return render_template('archive.html', months=months)

@app.route('/archive/create', methods=['POST'])
@login_required
@role_required('admin')
def archive_create():
    """Archive current month's data and reset for new month."""
    now = datetime.utcnow()
    prev_month = (now.replace(day=1) - timedelta(days=1))
    month_str = prev_month.strftime('%Y-%m')
    
    # Check if already archived
    existing = MonthlyArchive.query.filter_by(archive_month=month_str).first()
    if existing:
        flash(_('This month is already archived'), 'warning')
        return redirect(url_for('archive_page'))
    
    d_from = prev_month.replace(day=1)
    d_to = now.replace(day=1)
    
    # Archive faults
    faults = FaultReport.query.filter(FaultReport.created_at >= d_from, FaultReport.created_at < d_to).all()
    faults_data = [{'id': f.id, 'title': f.title, 'machine': f.machine.name if f.machine else '', 'priority': f.priority, 'status': f.status, 'created_at': f.created_at.strftime('%Y-%m-%d %H:%M')} for f in faults]
    db.session.add(MonthlyArchive(archive_month=month_str, section='faults', data_json=json.dumps(faults_data, ensure_ascii=False), created_by=current_user.id))
    
    # Archive orders
    orders = Opdracht.query.filter(Opdracht.aangemaakt >= d_from, Opdracht.aangemaakt < d_to).all()
    orders_data = [{'id': o.id, 'nummer': o.nummer, 'client': o.verantwoordelijke.naam if o.verantwoordelijke else '', 'status': o.status, 'total': o.totaal, 'created_at': o.aangemaakt.strftime('%Y-%m-%d %H:%M')} for o in orders]
    db.session.add(MonthlyArchive(archive_month=month_str, section='orders', data_json=json.dumps(orders_data, ensure_ascii=False), created_by=current_user.id))
    
    # Archive work reports
    reports = WorkReport.query.filter(WorkReport.created_at >= d_from, WorkReport.created_at < d_to).all()
    reports_data = [{'id': r.id, 'fault_id': r.fault_id, 'hours': r.time_spent_hours, 'description': r.description or '', 'created_at': r.created_at.strftime('%Y-%m-%d %H:%M')} for r in reports]
    db.session.add(MonthlyArchive(archive_month=month_str, section='work_reports', data_json=json.dumps(reports_data, ensure_ascii=False), created_by=current_user.id))
    
    # Archive time entries
    entries = TimeEntry.query.filter(TimeEntry.date >= d_from.date(), TimeEntry.date < d_to.date()).all()
    entries_data = [{'id': e.id, 'user_id': e.user_id, 'date': e.date.strftime('%Y-%m-%d'), 'hours': e.hours or 0, 'notes': e.notes or ''} for e in entries]
    db.session.add(MonthlyArchive(archive_month=month_str, section='time_entries', data_json=json.dumps(entries_data, ensure_ascii=False), created_by=current_user.id))
    
    db.session.commit()
    flash(_('Month archived successfully'), 'success')
    return redirect(url_for('archive_page'))

@app.route('/archive/<month>/<section>')
@login_required
@role_required('admin', 'director', 'technician')
def archive_detail(month, section):
    archive = MonthlyArchive.query.filter_by(archive_month=month, section=section).first_or_404()
    data = json.loads(archive.data_json)
    return render_template('archive_detail.html', archive=archive, data=data, month=month, section=section)

# ============================================================
# GLOBAL SEARCH
# ============================================================

@app.route('/api/search')
@login_required
def api_search():
    q = request.args.get('q', '').strip()
    if not q or len(q) < 2:
        return jsonify({'results': []})
    
    results = []
    limit = 20
    
    # Search machines
    machines = Machine.query.filter(
        (Machine.name.ilike(f'%{q}%')) | 
        (Machine.serial_number.ilike(f'%{q}%')) |
        (Machine.description.ilike(f'%{q}%'))
    ).limit(5).all()
    for m in machines:
        results.append({
            'type': 'machine',
            'icon': '⚙️',
            'title': m.name,
            'subtitle': f'{m.machine_type or ""} {m.serial_number or ""}'.strip(),
            'url': f'/machines/{m.id}',
            'status': m.status
        })
    
    # Search faults
    faults = FaultReport.query.filter(
        (FaultReport.title.ilike(f'%{q}%')) | 
        (FaultReport.description.ilike(f'%{q}%'))
    ).limit(5).all()
    for f in faults:
        results.append({
            'type': 'fault',
            'icon': '⚠️',
            'title': f.title,
            'subtitle': f'{f.machine.name} - {f.priority}',
            'url': f'/faults/{f.id}',
            'status': f.status
        })
    
    # Search work orders
    orders = Opdracht.query.filter(
        (Opdracht.nummer.ilike(f'%{q}%')) | 
        (Opdracht.apparaat.ilike(f'%{q}%')) |
        (Opdracht.model.ilike(f'%{q}%')) |
        (Opdracht.serienummer.ilike(f'%{q}%'))
    ).limit(5).all()
    for o in orders:
        results.append({
            'type': 'order',
            'icon': '📋',
            'title': f'{o.nummer} - {o.apparaat}',
            'subtitle': f'{o.model or ""} | {o.verantwoordelijke.naam}',
            'url': f'/orders/{o.id}',
            'status': o.status
        })
    
    # Search warehouse
    items = VoorraadItem.query.filter(
        (VoorraadItem.naam.ilike(f'%{q}%')) | 
        (VoorraadItem.categorie.ilike(f'%{q}%')) |
        (VoorraadItem.locatie.ilike(f'%{q}%'))
    ).limit(5).all()
    for i in items:
        results.append({
            'type': 'warehouse',
            'icon': '📦',
            'title': i.naam,
            'subtitle': f'{i.categorie or ""} | {i.hoeveelheid} {i.eenheid}',
            'url': f'/warehouse/{i.id}/edit',
            'status': 'low' if i.hoeveelheid <= i.minimum else 'ok'
        })
    
    # Search clients/responsible
    clients = Verantwoordelijke.query.filter(
        (Verantwoordelijke.naam.ilike(f'%{q}%')) | 
        (Verantwoordelijke.company.ilike(f'%{q}%')) |
        (Verantwoordelijke.telefoon.ilike(f'%{q}%')) |
        (Verantwoordelijke.email.ilike(f'%{q}%'))
    ).limit(5).all()
    for c in clients:
        results.append({
            'type': 'client',
            'icon': '👤',
            'title': c.naam,
            'subtitle': f'{c.company or ""} {c.telefoon or ""}'.strip(),
            'url': f'/responsible/{c.id}',
            'status': 'active'
        })
    
    # Search workers
    workers = Monteur.query.filter(
        (Monteur.naam.ilike(f'%{q}%')) | 
        (Monteur.specialisatie.ilike(f'%{q}%'))
    ).limit(3).all()
    for w in workers:
        results.append({
            'type': 'worker',
            'icon': '🔧',
            'title': w.naam,
            'subtitle': w.specialisatie or '',
            'url': f'/workers/{w.id}/edit',
            'status': 'active' if w.actief else 'inactive'
        })
    
    # Search contractors
    contractors = Contractor.query.filter(
        (Contractor.company_name.ilike(f'%{q}%')) | 
        (Contractor.service_type.ilike(f'%{q}%'))
    ).limit(3).all()
    for c in contractors:
        results.append({
            'type': 'contractor',
            'icon': '🏢',
            'title': c.company_name,
            'subtitle': c.service_type or '',
            'url': f'/contractors/{c.id}',
            'status': 'active' if c.is_active else 'inactive'
        })
    
    # Search TWO
    twos = TechnicalWorkOrder.query.filter(
        (TechnicalWorkOrder.number.ilike(f'%{q}%')) | 
        (TechnicalWorkOrder.description.ilike(f'%{q}%'))
    ).limit(3).all()
    for t in twos:
        results.append({
            'type': 'two',
            'icon': '🔧',
            'title': t.number,
            'subtitle': t.description[:50] if t.description else '',
            'url': f'/two/{t.id}',
            'status': t.status
        })
    
    # Search users (admin only)
    if current_user.has_role('admin'):
        users = User.query.filter(
            (User.username.ilike(f'%{q}%')) | 
            (User.display_name.ilike(f'%{q}%')) |
            (User.first_name.ilike(f'%{q}%')) |
            (User.last_name.ilike(f'%{q}%'))
        ).limit(3).all()
        for u in users:
            results.append({
                'type': 'user',
                'icon': '👥',
                'title': u.display_name or u.username,
                'subtitle': f'{u.role} | {u.username}',
                'url': f'/users/{u.id}',
                'status': 'active' if u.is_active_user else 'inactive'
            })
    
    return jsonify({'results': results[:limit], 'total': len(results)})

@app.route('/search')
@login_required
def search_page():
    q = request.args.get('q', '').strip()
    return render_template('search.html', query=q)

# ============================================================
# START
# ============================================================

if __name__ == '__main__':
    with app.app_context():
        db.create_all()
        
        # Migrations
        import sqlite3
        db_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'instance', 'werkplaats.db')
        if os.path.exists(db_path):
            conn = sqlite3.connect(db_path)
            cur = conn.cursor()
            
            # Create section_responsible table if missing
            cur.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='section_responsible'")
            if not cur.fetchone():
                cur.execute("CREATE TABLE section_responsible (section_id INTEGER REFERENCES factory_section(id), person_id INTEGER REFERENCES client(id), PRIMARY KEY (section_id, person_id))")
                # Migrate data from responsible_person_id column if it exists
                cur.execute("PRAGMA table_info(factory_section)")
                cols = [c[1] for c in cur.fetchall()]
                if 'responsible_person_id' in cols:
                    cur.execute("INSERT INTO section_responsible (section_id, person_id) SELECT id, responsible_person_id FROM factory_section WHERE responsible_person_id IS NOT NULL")
                conn.commit()
                print("Migration: created section_responsible table")
            
            # Add received_at to gas_cylinder if missing
            cur.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='gas_cylinder'")
            if cur.fetchone():
                cur.execute("PRAGMA table_info(gas_cylinder)")
                cols = [c[1] for c in cur.fetchall()]
                if 'received_at' not in cols:
                    cur.execute("ALTER TABLE gas_cylinder ADD COLUMN received_at DATETIME")
                    conn.commit()
                    print("Migration: added received_at to gas_cylinder")
            
            # Add responsible_user_id to machine_part if missing
            cur.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='machine_part'")
            if cur.fetchone():
                cur.execute("PRAGMA table_info(machine_part)")
                cols = [c[1] for c in cur.fetchall()]
                if 'responsible_user_id' not in cols:
                    cur.execute("ALTER TABLE machine_part ADD COLUMN responsible_user_id INTEGER REFERENCES user(id)")
                    conn.commit()
                    print("Migration: added responsible_user_id to machine_part")

            # Create fault_video table if missing
            cur.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='fault_video'")
            if not cur.fetchone():
                cur.execute("CREATE TABLE fault_video (id INTEGER PRIMARY KEY, fault_id INTEGER NOT NULL REFERENCES fault_report(id), filename TEXT NOT NULL, description TEXT, uploaded_at DATETIME)")
                conn.commit()
                print("Migration: created fault_video table")

            conn.close()

        
        # Create admin user if no users exist
        if User.query.count() == 0:
            import secrets as _secrets
            creds = []
            admin_pw = _secrets.token_urlsafe(12)
            admin = User(username='admin', display_name='Administrator', role='admin')
            admin.set_password(admin_pw)
            db.session.add(admin)
            creds.append(('admin', admin_pw))

            tech_pw = _secrets.token_urlsafe(12)
            tech = User(username='tech', display_name='Sergei Petrov', role='technician')
            tech.set_password(tech_pw)
            db.session.add(tech)
            creds.append(('tech', tech_pw))

            user_pw = _secrets.token_urlsafe(12)
            user = User(username='user', display_name='Jan de Vries', role='user')
            user.set_password(user_pw)
            db.session.add(user)
            creds.append(('user', user_pw))

            director_pw = _secrets.token_urlsafe(12)
            director = User(username='director', display_name='Director', role='director')
            director.set_password(director_pw)
            db.session.add(director)
            creds.append(('director', director_pw))
            
            # Create demo sections
            sections = [
                FactorySection(name='CNC Machining', description='CNC milling and turning operations', section_type='workshop', color='#3498db', floor_x=10, floor_y=10, width=40, height=35, responsible_user_id=2),
                FactorySection(name='Welding & Assembly', description='Welding stations and assembly area', section_type='workshop', color='#e67e22', floor_x=55, floor_y=10, width=35, height=35, responsible_user_id=2),
                FactorySection(name='Finishing', description='Grinding, polishing and surface treatment', section_type='workshop', color='#27ae60', floor_x=10, floor_y=50, width=30, height=30, responsible_user_id=2),
                FactorySection(name='Quality Control', description='Inspection and quality assurance', section_type='office', color='#9b59b6', floor_x=45, floor_y=50, width=25, height=30, responsible_user_id=4),
                FactorySection(name='Storage', description='Raw materials and finished goods storage', section_type='storage', color='#95a5a6', floor_x=75, floor_y=50, width=20, height=30, responsible_user_id=None),
            ]
            db.session.add_all(sections)
            db.session.commit()
            
            # Create demo machines
            machines = [
                Machine(name='CNC Mill X500', description='5-axis CNC milling machine', serial_number='CNC-2024-001', machine_type='Milling', floor_x=20, floor_y=25, status='active', section_id=1),
                Machine(name='Lathe LT200', description='CNC lathe for precision turning', serial_number='LT-2024-002', machine_type='Turning', floor_x=35, floor_y=25, status='active', section_id=1),
                Machine(name='Press HP100', description='Hydraulic press 100 ton', serial_number='HP-2024-003', machine_type='Pressing', floor_x=70, floor_y=25, status='active', section_id=2),
                Machine(name='Welder WS300', description='MIG/MAG welding station', serial_number='WS-2024-004', machine_type='Welding', floor_x=80, floor_y=25, status='active', section_id=2),
                Machine(name='Grinder GR50', description='Surface grinder', serial_number='GR-2024-005', machine_type='Grinding', floor_x=20, floor_y=65, status='maintenance', section_id=3),
                Machine(name='Drill DP20', description='Radial drill press', serial_number='DP-2024-006', machine_type='Drilling', floor_x=55, floor_y=65, status='active', section_id=4),
            ]
            db.session.add_all(machines)
            db.session.commit()
            
            # Assign machines to user
            user.assigned_machines = [machines[0], machines[1], machines[2]]
            db.session.commit()
            
            # Demo data
            demo = [
                Verantwoordelijke(naam='Jan de Vries', telefoon='+31 6 12345678', email='jan@mail.nl'),
                Verantwoordelijke(naam='Maria Bakker', telefoon='+31 6 23456789'),
                Verantwoordelijke(naam='Pieter Jansen', telefoon='+31 6 34567890', adres='Hoofdstraat 10, Amsterdam'),
            ]
            demo_w = [
                Monteur(naam='Sergei Petrov', specialisatie='Elektronica', tarief_per_uur=45),
                Monteur(naam='Alex de Boer', specialisatie='Mechanica', tarief_per_uur=40),
                Monteur(naam='Dmitri Smit', specialisatie='Algemeen onderhoud', tarief_per_uur=42),
            ]
            demo_i = [
                VoorraadItem(naam='Soldeerpasta', categorie='Verbruiksartikelen', eenheid='st', hoeveelheid=10, minimum=3, prijs=12),
                VoorraadItem(naam='Koperdraad 0.5mm', categorie='Materialen', eenheid='m', hoeveelheid=50, minimum=10, prijs=2),
                VoorraadItem(naam='Lager 608ZZ', categorie='Onderdelen', eenheid='st', hoeveelheid=20, minimum=5, prijs=6),
                VoorraadItem(naam='Lijm SuperMoment', categorie='Verbruiksartikelen', eenheid='st', hoeveelheid=5, minimum=2, prijs=8),
                VoorraadItem(naam='V-snaar', categorie='Onderdelen', eenheid='st', hoeveelheid=3, minimum=2, prijs=25),
            ]
            db.session.add_all(demo + demo_w + demo_i)
            db.session.commit()
            
            demo_o = [
                Opdracht(nummer='WO-20260809-0001', responsible_id=1, monteur_id=1, apparaat='Smartphone', model='iPhone 12',
                         probleem='Gebroken scherm, touchscreen werkt niet', status='afgeleverd',
                         arbeidskosten=250, onderdelenkosten=150, totaal=400,
                         gestart=datetime.utcnow()-timedelta(days=5), gereed=datetime.utcnow()-timedelta(days=3),
                         afgeleverd=datetime.utcnow()-timedelta(days=2)),
                Opdracht(nummer='WO-20260809-0002', responsible_id=2, monteur_id=2, apparaat='Laptop', model='ASUS X515',
                         probleem='Start niet op, laadt niet', status='in behandeling',
                         arbeidskosten=150, onderdelenkosten=0, totaal=150,
                         gestart=datetime.utcnow()-timedelta(days=1)),
                Opdracht(nummer='WO-20260809-0003', responsible_id=3, monteur_id=1, apparaat='Tablet', model='Samsung Tab A',
                         probleem='WiFi werkt niet, traag', status='aangenomen', arbeidskosten=0, onderdelenkosten=0, totaal=0),
            ]
            db.session.add_all(demo_o)
            db.session.commit()
            
            # Demo fault report
            fault = FaultReport(
                title='CNC Mill spindle vibration',
                description='Excessive vibration during high-speed operation. Possible bearing wear.',
                priority='high',
                machine_id=1,
                reporter_id=3  # user
            )
            db.session.add(fault)
            db.session.commit()
            
            print("\n" + "=" * 50)
            print("Demo data loaded! Generated credentials:")
            print("=" * 50)
            for username, pw in creds:
                print(f"  {username} / {pw}")
            print("=" * 50)
            print("Save these passwords! They are shown only once.\n")
    
    app.run(debug=True, host='0.0.0.0', port=5000)

# Chat routes removed - not needed currently
